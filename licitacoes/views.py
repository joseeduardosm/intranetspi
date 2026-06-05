# Criado por José Eduardo Santana Martins em 04/06/2026
# Objetivo: Controlar telas, permissões, CRUDs e exportações dos fluxos de licitações.

from copy import copy
from io import BytesIO

from django.conf import settings
from django.contrib import messages
from django.contrib.auth import get_user_model
from django.contrib.auth.mixins import LoginRequiredMixin
from django.contrib.auth.mixins import UserPassesTestMixin
from django.db.models import Max, Q
from django.http import Http404, HttpResponse
from django.shortcuts import get_object_or_404, redirect
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from django.utils.text import slugify
from django.views import View
from django.views.generic import CreateView, DeleteView, DetailView, ListView, TemplateView, UpdateView

from .forms import (
    DfdCreateForm,
    DfdItemTabelaForm,
    DfdSecaoForm,
    EtpTicCreateForm,
    EtpTicSecaoForm,
    FornecedorForm,
    ItemEtpMoveForm,
    ItemEtpTicForm,
    ItemMoveForm,
    ItemTRForm,
    PesquisaPrecoCreateForm,
    PesquisaPrecoFornecedorForm,
    PesquisaPrecoOrcamentoForm,
    SessaoEtpTicForm,
    SessaoTRForm,
    TabelaItemLinhaForm,
    TermoReferenciaForm,
)
from .models import (
    Dfd,
    DfdItemTabela,
    EtpTic,
    Fornecedor,
    ItemEtpTic,
    ItemTR,
    PesquisaPreco,
    PesquisaPrecoContato,
    PesquisaPrecoFornecedor,
    PesquisaPrecoFornecedorNota,
    PesquisaPrecoItemValor,
    SessaoEtpTic,
    SessaoTR,
    TabelaItemLinha,
    TermoReferencia,
)
from .services import (
    DFD_SECOES,
    DFD_SECOES_MAP,
    ETP_TIC_SECOES,
    ETP_TIC_SECOES_MAP,
    build_etp_tree,
    build_item_rows,
    clear_etp_item_children,
    clear_item_children,
    item_parent_for_tipo,
    build_termo_tree,
    duplicate_dfd,
    duplicate_etp,
    duplicate_etp_item,
    duplicate_item,
    duplicate_termo,
    dfd_status_por_secao,
    etp_status_por_secao,
    move_etp_item,
    move_item,
    next_ordem_etp_item,
    next_ordem_etp_sessao,
    next_ordem_item,
    next_ordem_sessao,
    normalize_etp_items,
    normalize_etp_sessoes,
    normalize_items,
    normalize_sessoes,
    parse_bulk_item_markers,
    pesquisa_preco_context,
    pesquisa_preco_itens,
    quantidade_text,
    red_mark_segments,
    replace_item_with_marker_nodes,
    render_dfd_sections,
    render_etp_sections,
)
from usuarios.services import SYSTEM_USERNAMES


User = get_user_model()


# Helpers de auditoria e marcação mantêm autoria/atualização sem repetir lógica nas views.
def _audit_user(request):
    return request.user if request.user.is_authenticated else None


def _touch_instance(instance, request):
    instance.atualizado_por = _audit_user(request)
    instance.save(update_fields=['atualizado_por', 'atualizado_em'])


def _touch_etp(etp, request):
    _touch_instance(etp, request)


def _touch_termo(termo, request):
    _touch_instance(termo, request)


def parser_legend(max_hash_level):
    """Retorna a legenda dos marcadores aceitos no editor de itens em lote."""

    return [
        {'marker': '@', 'label': 'subseção'},
        {'marker': '#', 'label': f'item/subitem até {max_hash_level} níveis'},
        {'marker': '**', 'label': 'inciso'},
        {'marker': '$$', 'label': 'alínea'},
    ]


ITEM_RED_MODE_FIELD = 'modo_destaque_texto'
ITEM_RED_MODE_COMMON = 'comum'
ITEM_RED_MODE_ALL = 'todo_vermelho'
ITEM_RED_MODE_ALL_WITH_CHILDREN = 'todo_vermelho_com_filhos'
ITEM_RED_MODE_EXCEPT_SUBSECTIONS = 'vermelho_sem_subsecoes'
ITEM_RED_MODE_OPTIONS = [
    (ITEM_RED_MODE_COMMON, 'texto comum'),
    (ITEM_RED_MODE_ALL, 'todo o texto vermelho'),
    (ITEM_RED_MODE_ALL_WITH_CHILDREN, 'todo o texto em vermelho incluindo filhos'),
    (ITEM_RED_MODE_EXCEPT_SUBSECTIONS, 'somente em vermelho, ignorar subseções'),
]
SESSAO_CHILDREN_RED_FIELD = 'filhos_em_vermelho'


def item_red_mode(request):
    """Lê o modo de destaque em vermelho escolhido no formulário."""

    value = request.POST.get(ITEM_RED_MODE_FIELD, ITEM_RED_MODE_COMMON)
    valid_values = {option[0] for option in ITEM_RED_MODE_OPTIONS}
    return value if value in valid_values else ITEM_RED_MODE_COMMON


def red_wrap_text(texto):
    texto = (texto or '').strip()
    if not texto:
        return texto
    if texto.startswith('*') and texto.endswith('*'):
        return texto
    return f'*{texto}*'


def apply_item_red_mode_to_text(texto, mode):
    if mode == ITEM_RED_MODE_COMMON:
        return texto
    return red_wrap_text(texto)


def apply_item_red_mode_to_nodes(nodes, mode):
    """Aplica destaque em vermelho aos nós parseados conforme o modo selecionado."""

    if mode == ITEM_RED_MODE_COMMON:
        return nodes
    for node in nodes:
        if mode in (ITEM_RED_MODE_ALL, ITEM_RED_MODE_ALL_WITH_CHILDREN) or node['tipo'] != 'SUBSECAO':
            node['texto'] = red_wrap_text(node['texto'])
        apply_item_red_mode_to_nodes(node['filhos'], mode)
    return nodes


def apply_red_to_item_descendants(item):
    for child in item.filhos.all():
        child.texto = red_wrap_text(child.texto)
        child.save(update_fields=['texto'])
        apply_red_to_item_descendants(child)


def apply_red_to_session_items(sessao):
    for item in sessao.itens.all():
        item.texto = red_wrap_text(item.texto)
        item.save(update_fields=['texto'])


def parser_context(context, request, max_hash_level):
    """Adiciona instruções do parser e modo de destaque ao contexto do formulário."""

    context['parser_legend'] = parser_legend(max_hash_level)
    context['item_red_mode_field'] = ITEM_RED_MODE_FIELD
    context['item_red_mode_options'] = ITEM_RED_MODE_OPTIONS
    context['item_red_mode_value'] = item_red_mode(request) if request.method == 'POST' else ITEM_RED_MODE_COMMON
    return context


def session_children_red_context(context, request):
    context['session_children_red_field'] = SESSAO_CHILDREN_RED_FIELD
    context['session_children_red_checked'] = request.POST.get(SESSAO_CHILDREN_RED_FIELD) == '1'
    return context


def item_focus_url(item):
    return f"{reverse('licitacoes:tr_detail', args=[item.sessao.termo.pk])}#item-{item.pk}"


def sessao_focus_url(sessao):
    return f"{reverse('licitacoes:tr_detail', args=[sessao.termo.pk])}#sessao-{sessao.pk}"


def item_delete_return_url(item):
    """Escolhe o melhor ponto de retorno depois da exclusão de um item do TR."""

    siblings = list(item.sessao.itens.filter(parent_id=item.parent_id).exclude(pk=item.pk).order_by('ordem', 'id'))
    next_item = next((sibling for sibling in siblings if (sibling.ordem, sibling.id) > (item.ordem, item.id)), None)
    if next_item:
        return item_focus_url(next_item)

    previous_items = [sibling for sibling in siblings if (sibling.ordem, sibling.id) < (item.ordem, item.id)]
    if previous_items:
        return item_focus_url(previous_items[-1])

    if item.parent:
        return item_focus_url(item.parent)

    return sessao_focus_url(item.sessao)


def item_indice(item):
    for row in build_item_rows(item.sessao):
        if row['item'].id == item.id:
            return row['indice']
    return ''


def tabela_item_url(item):
    return f"{reverse('licitacoes:tr_detail', args=[item.sessao.termo.pk])}#tabela-itens-{item.pk}"


def get_item_tabela_1_1(sessao_pk, item_pk, user):
    item = get_object_or_404(
        ItemTR,
        pk=item_pk,
        sessao__termo__sessoes__id=sessao_pk,
        sessao__termo__in=owned_queryset(TermoReferencia, user),
    )
    if item_indice(item) != '1.1':
        raise Http404('Tabela disponivel apenas para o item 1.1.')
    return item


def create_bulk_marker_items(model, sessao, parent, nodes, tipo_enum, next_order):
    first_item = None
    created_count = 0
    start_ordem = next_order(sessao, parent)

    def create_nodes(current_nodes, current_parent, start_order=1):
        nonlocal first_item, created_count
        for offset, node in enumerate(current_nodes):
            item = model.objects.create(
                sessao=sessao,
                parent=current_parent,
                tipo=getattr(tipo_enum, node['tipo']),
                texto=node['texto'],
                ordem=start_order + offset,
            )
            created_count += 1
            if first_item is None:
                first_item = item
            create_nodes(node['filhos'], item)

    create_nodes(nodes, parent, start_ordem)
    return first_item, created_count


def starts_with_item_marker(texto):
    first_line = next((line.strip() for line in (texto or '').splitlines() if line.strip()), '')
    return first_line.startswith(('@', '#', '**', '$$'))


def etp_item_focus_url(item):
    return f"{reverse('licitacoes:etp_detail', args=[item.sessao.etp.pk])}#item-etp-{item.pk}"


def etp_sessao_focus_url(sessao):
    return f"{reverse('licitacoes:etp_detail', args=[sessao.etp.pk])}#sessao-etp-{sessao.pk}"


def etp_item_delete_return_url(item):
    """Escolhe o melhor ponto de retorno depois da exclusão de um item do ETP TIC."""

    siblings = list(item.sessao.itens.filter(parent_id=item.parent_id).exclude(pk=item.pk).order_by('ordem', 'id'))
    next_item = next((sibling for sibling in siblings if (sibling.ordem, sibling.id) > (item.ordem, item.id)), None)
    if next_item:
        return etp_item_focus_url(next_item)

    previous_items = [sibling for sibling in siblings if (sibling.ordem, sibling.id) < (item.ordem, item.id)]
    if previous_items:
        return etp_item_focus_url(previous_items[-1])

    if item.parent:
        return etp_item_focus_url(item.parent)

    return etp_sessao_focus_url(item.sessao)


class SuperuserRequiredMixin(UserPassesTestMixin):
    """Exige usuário autenticado para acessar os fluxos internos de licitações."""

    def test_func(self):
        return self.request.user.is_authenticated


def is_system_admin(user):
    """Centraliza a regra de administrador do sistema."""

    return bool(user and user.is_authenticated and (user.is_superuser or user.is_staff))


class AdminRequiredMixin(UserPassesTestMixin):
    """Restringe cadastros administrativos a staff ou superusuários."""

    def test_func(self):
        return is_system_admin(self.request.user)


def owned_queryset(model, user):
    """Filtra documentos do usuário, compartilhados com ele ou todos para administradores."""

    queryset = model.objects.all()
    if is_system_admin(user):
        return queryset
    return queryset.filter(Q(criado_por=user) | Q(compartilhado_com=user)).distinct()


def owned_delete_queryset(model, user):
    """Permite exclusão por administradores ou pelo criador do documento."""

    queryset = model.objects.all()
    if is_system_admin(user):
        return queryset
    return queryset.filter(criado_por=user)


def owned_object_or_404(model, user, **kwargs):
    return get_object_or_404(owned_queryset(model, user), **kwargs)


def user_can_share_document(user, document):
    return user.is_authenticated and (is_system_admin(user) or document.criado_por_id == user.id)


def shareable_users_for(user):
    return (
        User.objects.filter(is_active=True, perfil__isnull=False)
        .exclude(pk=user.pk)
        .exclude(username__in=SYSTEM_USERNAMES)
        .order_by('perfil__nome_completo', 'username')
    )


def assign_owner(instance, request):
    """Preenche autoria inicial e atualização quando o modelo possui esses campos."""

    user = _audit_user(request)
    if not instance.criado_por_id:
        instance.criado_por = user
    if hasattr(instance, 'atualizado_por'):
        instance.atualizado_por = user
    return instance


class LicitacoesHomeView(SuperuserRequiredMixin, TemplateView):
    """Exibe a página inicial com os módulos documentais de licitações."""

    template_name = 'licitacoes/home.html'


class LicitacoesAdminContextMixin:
    """Disponibiliza no template se o usuário pode executar ações administrativas."""

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['is_system_admin'] = is_system_admin(self.request.user)
        return context


class EtpTicListView(LicitacoesAdminContextMixin, SuperuserRequiredMixin, ListView):
    """Lista ETPs TIC acessíveis ao usuário e prepara opções de compartilhamento."""

    model = EtpTic
    template_name = 'licitacoes/etp_list.html'
    context_object_name = 'etps'

    def get_queryset(self):
        return owned_queryset(EtpTic, self.request.user).select_related('criado_por', 'atualizado_por')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['shareable_users'] = shareable_users_for(self.request.user)
        return context


class EtpTicCreateView(SuperuserRequiredMixin, CreateView):
    """Cria um ETP TIC novo já configurado para o editor dinâmico."""

    model = EtpTic
    form_class = EtpTicCreateForm
    template_name = 'licitacoes/form.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo ETP TIC'
        context['voltar_url'] = reverse('licitacoes:etp_list')
        return context

    def form_valid(self, form):
        self.object = form.save(commit=False)
        assign_owner(self.object, self.request)
        self.object.declaracao_viabilidade = EtpTic.DECLARACAO_PADRAO
        self.object.usa_editor_dinamico = True
        self.object.save()
        return redirect('licitacoes:etp_detail', pk=self.object.pk)


class EtpTicEditView(SuperuserRequiredMixin, UpdateView):
    """Edita metadados ou a seção atual de um ETP TIC."""

    model = EtpTic
    context_object_name = 'etp'

    def get_queryset(self):
        return owned_queryset(EtpTic, self.request.user)

    def _edita_secao(self):
        return 'secao' in self.request.GET

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        if not self.object.usa_editor_dinamico:
            messages.info(request, 'ETP TIC legado disponivel somente para visualizacao.')
            return redirect('licitacoes:etp_preview', pk=self.object.pk)
        return super().dispatch(request, *args, **kwargs)

    def get_template_names(self):
        if self._edita_secao() and not self.object.usa_editor_dinamico:
            return ['licitacoes/etp_edit.html']
        return ['licitacoes/form.html']

    def get_form_class(self):
        if self._edita_secao() and not self.object.usa_editor_dinamico:
            return EtpTicSecaoForm
        return EtpTicCreateForm

    def _secao_numero(self):
        try:
            numero = int(self.request.GET.get('secao') or self.object.secao_atual or 1)
        except (TypeError, ValueError):
            numero = 1
        return numero if numero in ETP_TIC_SECOES_MAP else 1

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self._edita_secao():
            kwargs['section_fields'] = ETP_TIC_SECOES_MAP[self._secao_numero()]['campos']
        return kwargs

    def form_valid(self, form):
        if not self._edita_secao():
            self.object = form.save(commit=False)
            self.object.atualizado_por = _audit_user(self.request)
            self.object.save()
            return redirect('licitacoes:etp_preview', pk=self.object.pk)

        etp = form.save(commit=False)
        secao = self._secao_numero()
        acao = self.request.POST.get('_acao', 'salvar')
        if acao == 'proximo':
            etp.secao_atual = min(secao + 1, 18)
        elif acao == 'anterior':
            etp.secao_atual = max(secao - 1, 1)
        else:
            etp.secao_atual = secao
        etp.atualizado_por = _audit_user(self.request)
        etp.save()
        return redirect(f"{reverse('licitacoes:etp_edit', args=[etp.pk])}?secao={etp.secao_atual}")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if not self._edita_secao() or self.object.usa_editor_dinamico:
            context['titulo'] = 'Editar ETP TIC'
            context['voltar_url'] = reverse('licitacoes:etp_detail', args=[self.object.pk])
            return context

        secao_numero = self._secao_numero()
        status = etp_status_por_secao(self.object)
        context['secao_numero'] = secao_numero
        context['secao_atual'] = ETP_TIC_SECOES_MAP[secao_numero]
        context['secoes'] = [{**s, 'status': status[s['numero']]} for s in ETP_TIC_SECOES]
        return context


class EtpTicDetailView(SuperuserRequiredMixin, DetailView):
    model = EtpTic
    template_name = 'licitacoes/etp_detail.html'
    context_object_name = 'etp'

    def get_queryset(self):
        return owned_queryset(EtpTic, self.request.user)

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        if not self.object.usa_editor_dinamico:
            return redirect('licitacoes:etp_preview', pk=self.object.pk)
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['tree'] = build_etp_tree(self.object)
        return context


class EtpTicPreviewView(SuperuserRequiredMixin, DetailView):
    model = EtpTic
    template_name = 'licitacoes/etp_preview.html'
    context_object_name = 'etp'

    def get_queryset(self):
        return owned_queryset(EtpTic, self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if self.object.usa_editor_dinamico:
            context['tree'] = build_etp_tree(self.object)
        else:
            context['secoes_render'] = render_etp_sections(self.object)
        return context


class EtpTicConcluirView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        etp = owned_object_or_404(EtpTic, request.user, pk=pk)
        etp.status = EtpTic.Status.CONCLUIDO
        etp.atualizado_por = _audit_user(request)
        etp.save(update_fields=['status', 'atualizado_por', 'atualizado_em'])
        messages.success(request, 'ETP TIC concluido.')
        return redirect('licitacoes:etp_preview', pk=pk)


class EtpTicDeleteView(SuperuserRequiredMixin, DeleteView):
    model = EtpTic
    template_name = 'licitacoes/confirm_delete.html'
    success_url = reverse_lazy('licitacoes:etp_list')

    def get_queryset(self):
        return owned_delete_queryset(EtpTic, self.request.user)


class EtpTicExportDocxView(SuperuserRequiredMixin, View):
    def get(self, request, pk):
        etp = owned_object_or_404(EtpTic, request.user, pk=pk)
        if etp.usa_editor_dinamico:
            return _etp_dynamic_docx_response(etp)
        return _docx_response(f'ETP TIC - {etp.nome}', render_etp_sections(etp), f'etp_tic_{slugify(etp.nome) or etp.pk}.docx')


class EtpTicDuplicateView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        etp = owned_object_or_404(EtpTic, request.user, pk=pk)
        duplicate = duplicate_etp(etp)
        duplicate.criado_por = _audit_user(request)
        duplicate.save(update_fields=['criado_por', 'atualizado_em'])
        _touch_etp(duplicate, request)
        messages.success(request, 'ETP TIC duplicado.')
        if duplicate.usa_editor_dinamico:
            return redirect('licitacoes:etp_detail', pk=duplicate.pk)
        return redirect('licitacoes:etp_preview', pk=duplicate.pk)


class EtpTicShareView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        etp = owned_object_or_404(EtpTic, request.user, pk=pk)
        if not user_can_share_document(request.user, etp):
            raise Http404
        user = get_object_or_404(User, pk=request.POST.get('user_id'), is_active=True)
        if user.pk == request.user.pk or user.username in SYSTEM_USERNAMES:
            messages.error(request, 'Selecione um usuario valido para compartilhar.')
        else:
            etp.compartilhado_com.add(user)
            messages.success(request, f'ETP TIC compartilhado com {user.get_full_name() or user.username}.')
        return redirect('licitacoes:etp_list')


class SessaoEtpCreateView(SuperuserRequiredMixin, CreateView):
    model = SessaoEtpTic
    form_class = SessaoEtpTicForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.etp = owned_object_or_404(EtpTic, request.user, pk=kwargs['etp_pk'], usa_editor_dinamico=True)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.etp = self.etp
        self.object.ordem = next_ordem_etp_sessao(self.etp)
        self.object.save()
        _touch_etp(self.etp, self.request)
        return redirect(etp_sessao_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova sessao do ETP TIC'
        context['voltar_url'] = reverse('licitacoes:etp_detail', args=[self.etp.pk])
        return context


class SessaoEtpUpdateView(SuperuserRequiredMixin, UpdateView):
    model = SessaoEtpTic
    form_class = SessaoEtpTicForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(SessaoEtpTic, pk=kwargs['pk'], etp_id=kwargs['etp_pk'], etp__usa_editor_dinamico=True, etp__in=owned_queryset(EtpTic, request.user))
        self.etp = self.object.etp
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def form_valid(self, form):
        response = super().form_valid(form)
        if self.request.POST.get(SESSAO_CHILDREN_RED_FIELD) == '1':
            apply_red_to_session_items(self.object)
        return response

    def get_success_url(self):
        _touch_etp(self.etp, self.request)
        return etp_sessao_focus_url(self.object)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar sessao do ETP TIC'
        context['voltar_url'] = reverse('licitacoes:etp_detail', args=[self.etp.pk])
        session_children_red_context(context, self.request)
        return context


class SessaoEtpDeleteView(SuperuserRequiredMixin, DeleteView):
    model = SessaoEtpTic
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return SessaoEtpTic.objects.filter(etp_id=self.kwargs['etp_pk'], etp__usa_editor_dinamico=True, etp__in=owned_queryset(EtpTic, self.request.user))

    def get_success_url(self):
        etp = self.object.etp
        normalize_etp_sessoes(etp)
        _touch_etp(etp, self.request)
        return reverse('licitacoes:etp_detail', args=[etp.pk])


class SessaoEtpClearItemsView(SuperuserRequiredMixin, DeleteView):
    model = SessaoEtpTic
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return SessaoEtpTic.objects.filter(etp_id=self.kwargs['etp_pk'], etp__usa_editor_dinamico=True, etp__in=owned_queryset(EtpTic, self.request.user))

    def form_valid(self, form):
        removed = self.object.itens.count()
        self.object.itens.all().delete()
        _touch_etp(self.object.etp, self.request)
        messages.success(self.request, f'{removed} item(ns) removido(s) da sessao.')
        return redirect(etp_sessao_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Limpar filhos da sessao'
        context['descricao'] = 'Esta acao exclui todos os itens e subitens desta sessao e nao pode ser desfeita.'
        context['pergunta'] = f'Confirma a exclusao de todos os itens da sessao "{self.object}"?'
        context['botao_confirmar'] = 'Limpar filhos'
        return context


class ItemEtpCreateView(SuperuserRequiredMixin, CreateView):
    model = ItemEtpTic
    form_class = ItemEtpTicForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.sessao = get_object_or_404(SessaoEtpTic, pk=kwargs['sessao_pk'], etp__usa_editor_dinamico=True, etp__in=owned_queryset(EtpTic, request.user))
        self.parent = None
        if kwargs.get('parent_pk'):
            self.parent = get_object_or_404(ItemEtpTic, pk=kwargs['parent_pk'], sessao__etp=self.sessao.etp)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        texto = self.request.POST.get('texto', '')
        mode = item_red_mode(self.request)
        grupos_marcados = parse_bulk_item_markers(texto, max_hash_level=5)
        if grupos_marcados:
            apply_item_red_mode_to_nodes(grupos_marcados, mode)
            sessao = self.parent.sessao if self.parent else self.sessao
            first_item, created_count = create_bulk_marker_items(
                ItemEtpTic,
                sessao,
                self.parent,
                grupos_marcados,
                ItemEtpTic.Tipo,
                next_ordem_etp_item,
            )
            _touch_etp(first_item.sessao.etp, self.request)
            messages.success(self.request, f'{created_count} item(ns) criado(s).')
            return redirect(etp_item_focus_url(first_item))

        self.object = form.save(commit=False)
        self.object.texto = apply_item_red_mode_to_text(self.object.texto, mode)
        self.object.sessao = self.parent.sessao if self.parent else self.sessao
        self.object.parent = self.parent
        self.object.ordem = next_ordem_etp_item(self.object.sessao, self.parent)
        self.object.save()
        _touch_etp(self.object.sessao.etp, self.request)
        return redirect(etp_item_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo item do ETP TIC'
        context['ctrl_enter_submit'] = True
        parser_context(context, self.request, 5)
        if self.parent:
            context['voltar_url'] = etp_item_focus_url(self.parent)
        else:
            context['voltar_url'] = reverse('licitacoes:etp_detail', args=[self.sessao.etp.pk])
        return context


class ItemEtpUpdateView(SuperuserRequiredMixin, UpdateView):
    model = ItemEtpTic
    form_class = ItemEtpTicForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(ItemEtpTic, pk=kwargs['pk'], sessao__etp__sessoes__id=kwargs['sessao_pk'], sessao__etp__usa_editor_dinamico=True, sessao__etp__in=owned_queryset(EtpTic, request.user))
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def form_valid(self, form):
        texto = self.request.POST.get('texto', '')
        mode = item_red_mode(self.request)
        grupos_marcados = parse_bulk_item_markers(texto, max_hash_level=5) if starts_with_item_marker(texto) else []
        if grupos_marcados:
            apply_item_red_mode_to_nodes(grupos_marcados, mode)
            self.object, updated_count = replace_item_with_marker_nodes(self.object, grupos_marcados, ItemEtpTic.Tipo)
            _touch_etp(self.object.sessao.etp, self.request)
            messages.success(self.request, f'{updated_count} item(ns) atualizado(s).')
            return redirect(etp_item_focus_url(self.object))

        self.object = form.save(commit=False)
        self.object.texto = apply_item_red_mode_to_text(self.object.texto, mode)
        self.object.save()
        if mode == ITEM_RED_MODE_ALL_WITH_CHILDREN:
            apply_red_to_item_descendants(self.object)
        _touch_etp(self.object.sessao.etp, self.request)
        return redirect(etp_item_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar item do ETP TIC'
        context['ctrl_enter_submit'] = True
        parser_context(context, self.request, 5)
        context['voltar_url'] = etp_item_focus_url(self.object)
        return context


class ItemEtpDeleteView(SuperuserRequiredMixin, DeleteView):
    model = ItemEtpTic
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return ItemEtpTic.objects.filter(sessao__etp__sessoes__id=self.kwargs['sessao_pk'], sessao__etp__usa_editor_dinamico=True, sessao__etp__in=owned_queryset(EtpTic, self.request.user))

    def form_valid(self, form):
        sessao = self.object.sessao
        parent_id = self.object.parent_id
        return_url = etp_item_delete_return_url(self.object)
        self.object.delete()
        normalize_etp_sessoes(sessao.etp)
        normalize_etp_items(sessao, parent_id)
        _touch_etp(sessao.etp, self.request)
        return redirect(return_url)


class ItemEtpClearChildrenView(SuperuserRequiredMixin, DeleteView):
    model = ItemEtpTic
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return ItemEtpTic.objects.filter(sessao__etp__sessoes__id=self.kwargs['sessao_pk'], sessao__etp__usa_editor_dinamico=True, sessao__etp__in=owned_queryset(EtpTic, self.request.user))

    def form_valid(self, form):
        removed = clear_etp_item_children(self.object)
        _touch_etp(self.object.sessao.etp, self.request)
        messages.success(self.request, f'{removed} subitem(ns) removido(s).')
        return redirect(etp_item_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Limpar filhos'
        context['descricao'] = 'Esta acao exclui todos os subitens deste item e nao pode ser desfeita.'
        context['pergunta'] = f'Confirma a exclusao de todos os subitens de "{self.object}"?'
        context['botao_confirmar'] = 'Limpar filhos'
        return context


class ItemEtpMoveView(SuperuserRequiredMixin, View):
    template_name = 'licitacoes/etp_item_move.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_object_or_404(ItemEtpTic, pk=kwargs['pk'], sessao__etp__sessoes__id=kwargs['sessao_pk'], sessao__etp__usa_editor_dinamico=True, sessao__etp__in=owned_queryset(EtpTic, request.user))
        self.etp = self.item.sessao.etp
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        form = ItemEtpMoveForm(etp=self.etp, item=self.item)
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'etp': self.etp, 'modo': 'mover'})

    def post(self, request, *args, **kwargs):
        form = ItemEtpMoveForm(request.POST, etp=self.etp, item=self.item)
        if form.is_valid():
            target_token = form.cleaned_data['target']
            action = form.cleaned_data['action']
            child_position = form.cleaned_data.get('child_position')
            target = None
            target_sessao = None
            if target_token.startswith('item:'):
                target = get_object_or_404(ItemEtpTic, pk=int(target_token.split(':', 1)[1]), sessao__etp=self.etp)
            else:
                target_sessao = get_object_or_404(SessaoEtpTic, pk=int(target_token.split(':', 1)[1]), etp=self.etp)
                action = 'child'
            try:
                move_etp_item(self.item, target, action, target_sessao=target_sessao, child_position=child_position)
                _touch_etp(self.etp, request)
                messages.success(request, 'Item movido e estrutura renumerada.')
                return redirect(etp_item_focus_url(self.item))
            except ValueError as exc:
                form.add_error(None, str(exc))
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'etp': self.etp, 'modo': 'mover'})


class ItemEtpDuplicateView(SuperuserRequiredMixin, View):
    template_name = 'licitacoes/etp_item_move.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_object_or_404(ItemEtpTic, pk=kwargs['pk'], sessao__etp__sessoes__id=kwargs['sessao_pk'], sessao__etp__usa_editor_dinamico=True, sessao__etp__in=owned_queryset(EtpTic, request.user))
        self.etp = self.item.sessao.etp
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        form = ItemEtpMoveForm(etp=self.etp, item=self.item, action_label='Duplicar')
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'etp': self.etp, 'modo': 'duplicar'})

    def post(self, request, *args, **kwargs):
        form = ItemEtpMoveForm(request.POST, etp=self.etp, item=self.item, action_label='Duplicar')
        if form.is_valid():
            target_token = form.cleaned_data['target']
            action = form.cleaned_data['action']
            child_position = form.cleaned_data.get('child_position')
            target = None
            target_sessao = None
            if target_token.startswith('item:'):
                target = get_object_or_404(ItemEtpTic, pk=int(target_token.split(':', 1)[1]), sessao__etp=self.etp)
            else:
                target_sessao = get_object_or_404(SessaoEtpTic, pk=int(target_token.split(':', 1)[1]), etp=self.etp)
                action = 'child'
            try:
                duplicate = duplicate_etp_item(self.item, target, action, target_sessao=target_sessao, child_position=child_position)
                _touch_etp(self.etp, request)
                messages.success(request, 'Item duplicado com subitens e estrutura renumerada.')
                return redirect(etp_item_focus_url(duplicate))
            except ValueError as exc:
                form.add_error(None, str(exc))
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'etp': self.etp, 'modo': 'duplicar'})


class DfdListView(LicitacoesAdminContextMixin, SuperuserRequiredMixin, ListView):
    """Lista DFDs acessíveis ao usuário e prepara compartilhamento."""

    model = Dfd
    template_name = 'licitacoes/dfd_list.html'
    context_object_name = 'dfds'

    def get_queryset(self):
        return owned_queryset(Dfd, self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['shareable_users'] = shareable_users_for(self.request.user)
        return context


class DfdCreateView(SuperuserRequiredMixin, CreateView):
    """Cria um DFD e direciona para o preenchimento por seção."""

    model = Dfd
    form_class = DfdCreateForm
    template_name = 'licitacoes/form.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo DFD'
        context['voltar_url'] = reverse('licitacoes:dfd_list')
        return context

    def form_valid(self, form):
        self.object = form.save(commit=False)
        assign_owner(self.object, self.request)
        self.object.save()
        return redirect(f"{reverse('licitacoes:dfd_edit', args=[self.object.pk])}?secao=1")


class DfdEditView(SuperuserRequiredMixin, UpdateView):
    """Edita seções do DFD e controla avanço, conclusão e retorno."""

    model = Dfd
    context_object_name = 'dfd'

    def get_queryset(self):
        return owned_queryset(Dfd, self.request.user)

    def _edita_secao(self):
        return 'secao' in self.request.GET

    def get_template_names(self):
        if self._edita_secao():
            return ['licitacoes/dfd_edit.html']
        return ['licitacoes/form.html']

    def get_form_class(self):
        if self._edita_secao():
            return DfdSecaoForm
        return DfdCreateForm

    def _secao_numero(self):
        try:
            numero = int(self.request.GET.get('secao') or self.object.secao_atual or 1)
        except (TypeError, ValueError):
            numero = 1
        return numero if numero in DFD_SECOES_MAP else 1

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        if self._edita_secao():
            kwargs['section_fields'] = DFD_SECOES_MAP[self._secao_numero()]['campos']
        return kwargs

    def form_valid(self, form):
        if not self._edita_secao():
            self.object = form.save()
            return redirect('licitacoes:dfd_preview', pk=self.object.pk)

        dfd = form.save(commit=False)
        secao = self._secao_numero()
        acao = self.request.POST.get('_acao', 'salvar')
        if acao == 'proximo':
            dfd.secao_atual = min(secao + 1, len(DFD_SECOES))
        elif acao == 'anterior':
            dfd.secao_atual = max(secao - 1, 1)
        elif acao == 'concluir':
            dfd.secao_atual = secao
            dfd.status = Dfd.Status.CONCLUIDO
        else:
            dfd.secao_atual = secao
        dfd.save()
        if acao == 'concluir':
            messages.success(self.request, 'DFD concluido.')
            return redirect('licitacoes:dfd_preview', pk=dfd.pk)
        return redirect(f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao={dfd.secao_atual}")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        if not self._edita_secao():
            context['titulo'] = 'Editar DFD'
            context['voltar_url'] = reverse('licitacoes:dfd_preview', args=[self.object.pk])
            return context

        secao_numero = self._secao_numero()
        status = dfd_status_por_secao(self.object)
        context['secao_numero'] = secao_numero
        context['secao_atual'] = DFD_SECOES_MAP[secao_numero]
        context['secoes'] = [{**s, 'status': status[s['numero']]} for s in DFD_SECOES]
        context['itens_tabela'] = self.object.itens_tabela.order_by('ordem', 'id')
        return context


class DfdPreviewView(SuperuserRequiredMixin, DetailView):
    model = Dfd
    template_name = 'licitacoes/dfd_preview.html'
    context_object_name = 'dfd'

    def get_queryset(self):
        return owned_queryset(Dfd, self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['secoes_render'] = render_dfd_sections(self.object)
        return context


class DfdConcluirView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        dfd = owned_object_or_404(Dfd, request.user, pk=pk)
        dfd.status = Dfd.Status.CONCLUIDO
        dfd.save(update_fields=['status', 'atualizado_em'])
        messages.success(request, 'DFD concluido.')
        return redirect('licitacoes:dfd_preview', pk=pk)


class DfdDeleteView(SuperuserRequiredMixin, DeleteView):
    model = Dfd
    template_name = 'licitacoes/confirm_delete.html'
    success_url = reverse_lazy('licitacoes:dfd_list')

    def get_queryset(self):
        return owned_delete_queryset(Dfd, self.request.user)


class DfdDuplicateView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        dfd = owned_object_or_404(Dfd, request.user, pk=pk)
        duplicate = duplicate_dfd(dfd)
        duplicate.criado_por = _audit_user(request)
        duplicate.save(update_fields=['criado_por', 'atualizado_em'])
        messages.success(request, 'DFD duplicado.')
        return redirect(f"{reverse('licitacoes:dfd_edit', args=[duplicate.pk])}?secao={duplicate.secao_atual}")


class DfdShareView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        dfd = owned_object_or_404(Dfd, request.user, pk=pk)
        if not user_can_share_document(request.user, dfd):
            raise Http404
        user = get_object_or_404(User, pk=request.POST.get('user_id'), is_active=True)
        if user.pk == request.user.pk or user.username in SYSTEM_USERNAMES:
            messages.error(request, 'Selecione um usuario valido para compartilhar.')
        else:
            dfd.compartilhado_com.add(user)
            messages.success(request, f'DFD compartilhado com {user.get_full_name() or user.username}.')
        return redirect('licitacoes:dfd_list')


class DfdItemTabelaCreateView(SuperuserRequiredMixin, CreateView):
    model = DfdItemTabela
    form_class = DfdItemTabelaForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.dfd = owned_object_or_404(Dfd, request.user, pk=kwargs['dfd_pk'])
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.dfd = self.dfd
        self.object.ordem = (self.dfd.itens_tabela.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        self.object.save()
        return redirect(f"{reverse('licitacoes:dfd_edit', args=[self.dfd.pk])}?secao=2#dfd-tabela")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo item da tabela do DFD'
        context['voltar_url'] = f"{reverse('licitacoes:dfd_edit', args=[self.dfd.pk])}?secao=2#dfd-tabela"
        return context


class DfdItemTabelaUpdateView(SuperuserRequiredMixin, UpdateView):
    model = DfdItemTabela
    form_class = DfdItemTabelaForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.dfd = owned_object_or_404(Dfd, request.user, pk=kwargs['dfd_pk'])
        self.object = get_object_or_404(DfdItemTabela, pk=kwargs['pk'], dfd=self.dfd)
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def get_success_url(self):
        return f"{reverse('licitacoes:dfd_edit', args=[self.dfd.pk])}?secao=2#dfd-tabela"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar item da tabela do DFD'
        context['voltar_url'] = self.get_success_url()
        return context


class DfdItemTabelaDeleteView(SuperuserRequiredMixin, DeleteView):
    model = DfdItemTabela
    template_name = 'licitacoes/confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.dfd = owned_object_or_404(Dfd, request.user, pk=kwargs['dfd_pk'])
        self.object = get_object_or_404(DfdItemTabela, pk=kwargs['pk'], dfd=self.dfd)
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def form_valid(self, form):
        self.object.delete()
        for idx, item in enumerate(self.dfd.itens_tabela.order_by('ordem', 'id'), start=1):
            if item.ordem != idx:
                item.ordem = idx
                item.save(update_fields=['ordem'])
        return redirect(f"{reverse('licitacoes:dfd_edit', args=[self.dfd.pk])}?secao=2#dfd-tabela")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir item da tabela do DFD'
        context['voltar_url'] = f"{reverse('licitacoes:dfd_edit', args=[self.dfd.pk])}?secao=2#dfd-tabela"
        return context


class DfdExportDocxView(SuperuserRequiredMixin, View):
    def get(self, request, pk):
        dfd = owned_object_or_404(Dfd, request.user, pk=pk)
        return _dfd_docx_response(dfd)


class TermoListView(LicitacoesAdminContextMixin, SuperuserRequiredMixin, ListView):
    """Lista Termos de Referência acessíveis ao usuário."""

    model = TermoReferencia
    template_name = 'licitacoes/tr_list.html'
    context_object_name = 'termos'

    def get_queryset(self):
        return owned_queryset(TermoReferencia, self.request.user).select_related('criado_por', 'atualizado_por')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['shareable_users'] = shareable_users_for(self.request.user)
        return context


class TermoCreateView(SuperuserRequiredMixin, CreateView):
    """Cria um Termo de Referência e atribui autoria."""

    model = TermoReferencia
    form_class = TermoReferenciaForm
    template_name = 'licitacoes/form.html'

    def get_success_url(self):
        return reverse('licitacoes:tr_detail', args=[self.object.pk])

    def form_valid(self, form):
        self.object = form.save(commit=False)
        assign_owner(self.object, self.request)
        self.object.save()
        return redirect(self.get_success_url())

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo TR'
        context['voltar_url'] = reverse('licitacoes:tr_list')
        return context


class TermoUpdateView(SuperuserRequiredMixin, UpdateView):
    model = TermoReferencia
    form_class = TermoReferenciaForm
    template_name = 'licitacoes/form.html'

    def get_queryset(self):
        return owned_queryset(TermoReferencia, self.request.user)

    def get_success_url(self):
        return reverse('licitacoes:tr_detail', args=[self.object.pk])

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.atualizado_por = _audit_user(self.request)
        self.object.save()
        return redirect(self.get_success_url())

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar TR'
        context['voltar_url'] = reverse('licitacoes:tr_detail', args=[self.object.pk])
        return context


class TermoDeleteView(SuperuserRequiredMixin, DeleteView):
    model = TermoReferencia
    template_name = 'licitacoes/confirm_delete.html'
    success_url = reverse_lazy('licitacoes:tr_list')

    def get_queryset(self):
        return owned_delete_queryset(TermoReferencia, self.request.user)


class TermoDuplicateView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        termo = owned_object_or_404(TermoReferencia, request.user, pk=pk)
        duplicate = duplicate_termo(termo)
        duplicate.criado_por = _audit_user(request)
        duplicate.save(update_fields=['criado_por', 'atualizado_em'])
        _touch_termo(duplicate, request)
        messages.success(request, 'TR duplicado.')
        return redirect('licitacoes:tr_detail', pk=duplicate.pk)


class TermoShareView(SuperuserRequiredMixin, View):
    def post(self, request, pk):
        termo = owned_object_or_404(TermoReferencia, request.user, pk=pk)
        if not user_can_share_document(request.user, termo):
            raise Http404
        user = get_object_or_404(User, pk=request.POST.get('user_id'), is_active=True)
        if user.pk == request.user.pk or user.username in SYSTEM_USERNAMES:
            messages.error(request, 'Selecione um usuario valido para compartilhar.')
        else:
            termo.compartilhado_com.add(user)
            messages.success(request, f'TR compartilhado com {user.get_full_name() or user.username}.')
        return redirect('licitacoes:tr_list')


class TermoDetailView(SuperuserRequiredMixin, DetailView):
    """Exibe o TR em árvore com seções, itens e tabelas vinculadas."""

    model = TermoReferencia
    template_name = 'licitacoes/tr_detail.html'
    context_object_name = 'termo'

    def get_queryset(self):
        return owned_queryset(TermoReferencia, self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['tree'] = build_termo_tree(self.object)
        return context


class PesquisaPrecoOpenView(SuperuserRequiredMixin, View):
    """Abre a pesquisa existente do TR ou redireciona para criação."""

    def get(self, request, pk):
        termo = owned_object_or_404(TermoReferencia, request.user, pk=pk)
        pesquisa = getattr(termo, 'pesquisa_preco', None)
        if pesquisa:
            return redirect('licitacoes:pesquisa_preco_detail', termo_pk=termo.pk)
        return redirect('licitacoes:pesquisa_preco_create', termo_pk=termo.pk)


class PesquisaPrecoCreateView(SuperuserRequiredMixin, CreateView):
    """Cria a pesquisa de preço associada a um TR."""

    model = PesquisaPreco
    form_class = PesquisaPrecoCreateForm
    template_name = 'licitacoes/pesquisa_preco_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.termo = owned_object_or_404(TermoReferencia, request.user, pk=kwargs['termo_pk'])
        if hasattr(self.termo, 'pesquisa_preco'):
            return redirect('licitacoes:pesquisa_preco_detail', termo_pk=self.termo.pk)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.termo = self.termo
        if self.object.tipo == PesquisaPreco.Tipo.AQUISICAO:
            self.object.vigencia_meses = None
        self.object.save()
        _touch_termo(self.termo, self.request)
        messages.success(self.request, 'Pesquisa de Preço criada.')
        return redirect('licitacoes:pesquisa_preco_detail', termo_pk=self.termo.pk)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova Pesquisa de Preço'
        context['voltar_url'] = reverse('licitacoes:tr_detail', args=[self.termo.pk])
        return context


class PesquisaPrecoDetailView(SuperuserRequiredMixin, DetailView):
    model = PesquisaPreco
    template_name = 'licitacoes/pesquisa_preco_detail.html'
    context_object_name = 'pesquisa'

    def get_object(self, queryset=None):
        return get_object_or_404(PesquisaPreco.objects.select_related('termo'), termo_id=self.kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, self.request.user))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context.update(pesquisa_preco_context(self.object))
        context['fornecedor_form'] = PesquisaPrecoFornecedorForm(pesquisa=self.object)
        context['is_servico'] = self.object.tipo == PesquisaPreco.Tipo.SERVICO
        return context


class PesquisaPrecoDeleteView(SuperuserRequiredMixin, DeleteView):
    model = PesquisaPreco
    template_name = 'licitacoes/confirm_delete.html'

    def get_object(self, queryset=None):
        return get_object_or_404(PesquisaPreco.objects.select_related('termo'), termo_id=self.kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, self.request.user))

    def get_success_url(self):
        termo = self.object.termo
        _touch_termo(termo, self.request)
        messages.success(self.request, 'Pesquisa de Preço excluída.')
        return reverse('licitacoes:tr_detail', args=[termo.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir Pesquisa de Preço'
        context['descricao'] = 'Esta ação remove fornecedores vinculados, contatos e valores lançados nesta pesquisa.'
        context['pergunta'] = f'Confirma a exclusão da Pesquisa de Preço do TR "{self.object.termo.nome}"?'
        return context


class PesquisaPrecoFornecedorAddView(SuperuserRequiredMixin, View):
    def post(self, request, termo_pk):
        pesquisa = get_object_or_404(PesquisaPreco, termo_id=termo_pk, termo__in=owned_queryset(TermoReferencia, request.user))
        form = PesquisaPrecoFornecedorForm(request.POST, pesquisa=pesquisa)
        if form.is_valid():
            PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=form.cleaned_data['fornecedor'])
            _touch_termo(pesquisa.termo, request)
            messages.success(request, 'Fornecedor adicionado à pesquisa.')
        else:
            messages.error(request, 'Selecione um fornecedor válido.')
        return redirect('licitacoes:pesquisa_preco_detail', termo_pk=termo_pk)


class PesquisaPrecoFornecedorRemoveView(SuperuserRequiredMixin, DeleteView):
    model = PesquisaPrecoFornecedor
    template_name = 'licitacoes/confirm_delete.html'

    def get_object(self, queryset=None):
        return get_object_or_404(
            PesquisaPrecoFornecedor.objects.select_related('pesquisa__termo', 'fornecedor'),
            pk=self.kwargs['pk'],
            pesquisa__termo_id=self.kwargs['termo_pk'],
            pesquisa__termo__in=owned_queryset(TermoReferencia, self.request.user),
        )

    def get_success_url(self):
        termo = self.object.pesquisa.termo
        _touch_termo(termo, self.request)
        messages.success(self.request, 'Fornecedor removido desta pesquisa.')
        return reverse('licitacoes:pesquisa_preco_detail', args=[termo.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Remover fornecedor da pesquisa'
        context['descricao'] = 'Esta ação remove o vínculo, os contatos e os valores deste fornecedor apenas nesta pesquisa.'
        context['pergunta'] = f'Remover "{self.object.fornecedor.razao_social}" desta pesquisa?'
        context['botao_confirmar'] = 'Remover'
        return context


class PesquisaPrecoFornecedorCreateView(SuperuserRequiredMixin, CreateView):
    model = Fornecedor
    form_class = FornecedorForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.pesquisa = get_object_or_404(PesquisaPreco, termo_id=kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, request.user))
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save()
        PesquisaPrecoFornecedor.objects.get_or_create(pesquisa=self.pesquisa, fornecedor=self.object)
        _touch_termo(self.pesquisa.termo, self.request)
        messages.success(self.request, 'Fornecedor cadastrado e adicionado à pesquisa.')
        return redirect('licitacoes:pesquisa_preco_detail', termo_pk=self.pesquisa.termo_id)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo fornecedor'
        context['voltar_url'] = reverse('licitacoes:pesquisa_preco_detail', args=[self.pesquisa.termo_id])
        return context


class PesquisaPrecoFornecedorUpdateView(SuperuserRequiredMixin, UpdateView):
    model = Fornecedor
    form_class = FornecedorForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.pesquisa = get_object_or_404(PesquisaPreco, termo_id=kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, request.user))
        self.object = get_object_or_404(Fornecedor, pk=kwargs['pk'], pesquisas_preco__pesquisa=self.pesquisa)
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def get_success_url(self):
        _touch_termo(self.pesquisa.termo, self.request)
        messages.success(self.request, 'Fornecedor atualizado.')
        return reverse('licitacoes:pesquisa_preco_detail', args=[self.pesquisa.termo_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar fornecedor'
        context['voltar_url'] = reverse('licitacoes:pesquisa_preco_detail', args=[self.pesquisa.termo_id])
        return context


class PesquisaPrecoAtualizarContatoView(SuperuserRequiredMixin, View):
    def post(self, request, termo_pk, pk):
        pesquisa_fornecedor = get_object_or_404(PesquisaPrecoFornecedor, pk=pk, pesquisa__termo_id=termo_pk, pesquisa__termo__in=owned_queryset(TermoReferencia, request.user))
        PesquisaPrecoContato.objects.create(
            pesquisa_fornecedor=pesquisa_fornecedor,
            data_contato=timezone.localdate(),
        )
        _touch_termo(pesquisa_fornecedor.pesquisa.termo, request)
        messages.success(request, 'Último contato atualizado.')
        return redirect('licitacoes:pesquisa_preco_detail', termo_pk=termo_pk)


class PesquisaPrecoFornecedorNotaCreateView(SuperuserRequiredMixin, View):
    def post(self, request, termo_pk, pk):
        pesquisa_fornecedor = get_object_or_404(
            PesquisaPrecoFornecedor,
            pk=pk,
            pesquisa__termo_id=termo_pk,
            pesquisa__termo__in=owned_queryset(TermoReferencia, request.user),
        )
        texto = (request.POST.get('texto') or '').strip()
        if not texto:
            messages.error(request, 'Informe o texto da nota.')
            return redirect('licitacoes:pesquisa_preco_detail', termo_pk=termo_pk)
        PesquisaPrecoFornecedorNota.objects.create(
            pesquisa_fornecedor=pesquisa_fornecedor,
            texto=texto,
            criado_por=_audit_user(request),
        )
        _touch_termo(pesquisa_fornecedor.pesquisa.termo, request)
        messages.success(request, 'Nota registrada.')
        return redirect('licitacoes:pesquisa_preco_detail', termo_pk=termo_pk)


class PesquisaPrecoOrcamentoView(SuperuserRequiredMixin, View):
    template_name = 'licitacoes/pesquisa_preco_orcamento_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.pesquisa_fornecedor = get_object_or_404(
            PesquisaPrecoFornecedor.objects.select_related('pesquisa__termo', 'fornecedor'),
            pk=kwargs['pk'],
            pesquisa__termo_id=kwargs['termo_pk'],
            pesquisa__termo__in=owned_queryset(TermoReferencia, request.user),
        )
        self.pesquisa = self.pesquisa_fornecedor.pesquisa
        self.itens = list(pesquisa_preco_itens(self.pesquisa))
        return super().dispatch(request, *args, **kwargs)

    def _initial_valores(self):
        return {
            valor.item_id: valor.preco_unitario
            for valor in self.pesquisa_fornecedor.valores.all()
        }

    def get(self, request, *args, **kwargs):
        form = PesquisaPrecoOrcamentoForm(
            initial={
                'data_resposta': self.pesquisa_fornecedor.data_resposta,
                'validade_orcamento_dias': self.pesquisa_fornecedor.validade_orcamento_dias,
            },
            itens=self.itens,
            valores=self._initial_valores(),
            has_document=bool(self.pesquisa_fornecedor.documento_fornecedor),
        )
        return _render(request, self.template_name, {
            'form': form,
            'pesquisa': self.pesquisa,
            'pesquisa_fornecedor': self.pesquisa_fornecedor,
            'itens': self.itens,
        })

    def post(self, request, *args, **kwargs):
        form = PesquisaPrecoOrcamentoForm(
            request.POST,
            request.FILES,
            itens=self.itens,
            valores=self._initial_valores(),
            has_document=bool(self.pesquisa_fornecedor.documento_fornecedor),
        )
        if form.is_valid():
            self.pesquisa_fornecedor.data_resposta = form.cleaned_data['data_resposta']
            self.pesquisa_fornecedor.validade_orcamento_dias = form.cleaned_data['validade_orcamento_dias']
            update_fields = ['data_resposta', 'validade_orcamento_dias', 'atualizado_em']
            if form.cleaned_data.get('documento_fornecedor'):
                self.pesquisa_fornecedor.documento_fornecedor = form.cleaned_data['documento_fornecedor']
                update_fields.append('documento_fornecedor')
            self.pesquisa_fornecedor.save(update_fields=update_fields)
            for item in self.itens:
                PesquisaPrecoItemValor.objects.update_or_create(
                    pesquisa_fornecedor=self.pesquisa_fornecedor,
                    item=item,
                    defaults={'preco_unitario': form.cleaned_data[PesquisaPrecoOrcamentoForm.item_field_name(item)]},
                )
            _touch_termo(self.pesquisa.termo, request)
            messages.success(request, 'Orçamento salvo.')
            return redirect('licitacoes:pesquisa_preco_detail', termo_pk=self.pesquisa.termo_id)
        return _render(request, self.template_name, {
            'form': form,
            'pesquisa': self.pesquisa,
            'pesquisa_fornecedor': self.pesquisa_fornecedor,
            'itens': self.itens,
        })


class PesquisaPrecoExportXlsxView(SuperuserRequiredMixin, View):
    """Exporta o quadro comparativo da pesquisa de preço em XLSX."""

    def get(self, request, termo_pk):
        pesquisa = get_object_or_404(PesquisaPreco.objects.select_related('termo'), termo_id=termo_pk, termo__in=owned_queryset(TermoReferencia, request.user))
        return _pesquisa_preco_xlsx_response(pesquisa)


class FornecedorListView(AdminRequiredMixin, ListView):
    """Lista fornecedores cadastrados para uso nas pesquisas de preço."""

    model = Fornecedor
    template_name = 'licitacoes/fornecedor_list.html'
    context_object_name = 'fornecedores'


class FornecedorCreateView(AdminRequiredMixin, CreateView):
    """Cadastra fornecedores disponíveis para cotação."""

    model = Fornecedor
    form_class = FornecedorForm
    template_name = 'licitacoes/form.html'
    success_url = reverse_lazy('licitacoes:fornecedor_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo fornecedor'
        context['voltar_url'] = reverse('licitacoes:fornecedor_list')
        return context


class FornecedorUpdateView(AdminRequiredMixin, UpdateView):
    model = Fornecedor
    form_class = FornecedorForm
    template_name = 'licitacoes/form.html'
    success_url = reverse_lazy('licitacoes:fornecedor_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar fornecedor'
        context['voltar_url'] = reverse('licitacoes:fornecedor_list')
        return context


class FornecedorDeleteView(AdminRequiredMixin, DeleteView):
    model = Fornecedor
    template_name = 'licitacoes/confirm_delete.html'
    success_url = reverse_lazy('licitacoes:fornecedor_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir fornecedor'
        context['descricao'] = 'Esta ação exclui o fornecedor do sistema e remove seus vínculos em pesquisas de preço.'
        context['pergunta'] = f'Confirma a exclusão do fornecedor "{self.object.razao_social}"?'
        return context


class SessaoCreateView(SuperuserRequiredMixin, CreateView):
    model = SessaoTR
    form_class = SessaoTRForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.termo = owned_object_or_404(TermoReferencia, request.user, pk=kwargs['termo_pk'])
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.termo = self.termo
        self.object.ordem = next_ordem_sessao(self.termo)
        self.object.save()
        _touch_termo(self.termo, self.request)
        return redirect('licitacoes:tr_detail', pk=self.termo.pk)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova sessao'
        context['voltar_url'] = reverse('licitacoes:tr_detail', args=[self.termo.pk])
        return context


class SessaoUpdateView(SuperuserRequiredMixin, UpdateView):
    model = SessaoTR
    form_class = SessaoTRForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(SessaoTR, pk=kwargs['pk'], termo_id=kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, request.user))
        self.termo = self.object.termo
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def form_valid(self, form):
        response = super().form_valid(form)
        if self.request.POST.get(SESSAO_CHILDREN_RED_FIELD) == '1':
            apply_red_to_session_items(self.object)
        return response

    def get_success_url(self):
        _touch_termo(self.termo, self.request)
        return reverse('licitacoes:tr_detail', args=[self.termo.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar sessao'
        context['voltar_url'] = reverse('licitacoes:tr_detail', args=[self.termo.pk])
        session_children_red_context(context, self.request)
        return context


class SessaoDeleteView(SuperuserRequiredMixin, DeleteView):
    model = SessaoTR
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return SessaoTR.objects.filter(termo_id=self.kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, self.request.user))

    def get_success_url(self):
        termo = self.object.termo
        normalize_sessoes(termo)
        _touch_termo(termo, self.request)
        return reverse('licitacoes:tr_detail', args=[termo.pk])


class SessaoClearItemsView(SuperuserRequiredMixin, DeleteView):
    model = SessaoTR
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return SessaoTR.objects.filter(termo_id=self.kwargs['termo_pk'], termo__in=owned_queryset(TermoReferencia, self.request.user))

    def form_valid(self, form):
        removed = self.object.itens.count()
        self.object.itens.all().delete()
        _touch_termo(self.object.termo, self.request)
        messages.success(self.request, f'{removed} item(ns) removido(s) da sessao.')
        return redirect(sessao_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Limpar filhos da sessao'
        context['descricao'] = 'Esta acao exclui todos os itens e subitens desta sessao e nao pode ser desfeita.'
        context['pergunta'] = f'Confirma a exclusao de todos os itens da sessao "{self.object}"?'
        context['botao_confirmar'] = 'Limpar filhos'
        return context


class ItemCreateView(SuperuserRequiredMixin, CreateView):
    model = ItemTR
    form_class = ItemTRForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.sessao = get_object_or_404(SessaoTR, pk=kwargs['sessao_pk'], termo__in=owned_queryset(TermoReferencia, request.user))
        self.parent = None
        if kwargs.get('parent_pk'):
            self.parent = get_object_or_404(ItemTR, pk=kwargs['parent_pk'], sessao__termo=self.sessao.termo)
        self.tipo = request.GET.get('tipo') or ItemTR.Tipo.NUMERICO
        if self.tipo not in ItemTR.Tipo.values:
            self.tipo = ItemTR.Tipo.NUMERICO
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        texto = self.request.POST.get('texto', '')
        mode = item_red_mode(self.request)
        grupos_marcados = parse_bulk_item_markers(texto, max_hash_level=4)
        if grupos_marcados:
            apply_item_red_mode_to_nodes(grupos_marcados, mode)
            parent = item_parent_for_tipo(self.parent, self.tipo)
            sessao = parent.sessao if parent else self.sessao
            first_item, created_count = create_bulk_marker_items(
                ItemTR,
                sessao,
                parent,
                grupos_marcados,
                ItemTR.Tipo,
                next_ordem_item,
            )
            _touch_termo(first_item.sessao.termo, self.request)
            messages.success(self.request, f'{created_count} item(ns) criado(s).')
            return redirect(item_focus_url(first_item))

        self.object = form.save(commit=False)
        self.object.texto = apply_item_red_mode_to_text(self.object.texto, mode)
        parent = item_parent_for_tipo(self.parent, self.tipo)
        self.object.sessao = parent.sessao if parent else self.sessao
        self.object.parent = parent
        self.object.tipo = self.tipo
        self.object.ordem = next_ordem_item(self.object.sessao, parent)
        self.object.save()
        _touch_termo(self.object.sessao.termo, self.request)
        return redirect(item_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo item do TR'
        context['ctrl_enter_submit'] = True
        parser_context(context, self.request, 4)
        if self.parent:
            context['voltar_url'] = item_focus_url(self.parent)
        else:
            context['voltar_url'] = reverse('licitacoes:tr_detail', args=[self.sessao.termo.pk])
        return context


class ItemUpdateView(SuperuserRequiredMixin, UpdateView):
    model = ItemTR
    form_class = ItemTRForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = get_object_or_404(ItemTR, pk=kwargs['pk'], sessao__termo__sessoes__id=kwargs['sessao_pk'], sessao__termo__in=owned_queryset(TermoReferencia, request.user))
        self.sessao = self.object.sessao
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def form_valid(self, form):
        texto = self.request.POST.get('texto', '')
        mode = item_red_mode(self.request)
        grupos_marcados = parse_bulk_item_markers(texto, max_hash_level=4) if starts_with_item_marker(texto) else []
        if grupos_marcados:
            apply_item_red_mode_to_nodes(grupos_marcados, mode)
            self.object, updated_count = replace_item_with_marker_nodes(self.object, grupos_marcados, ItemTR.Tipo)
            _touch_termo(self.object.sessao.termo, self.request)
            messages.success(self.request, f'{updated_count} item(ns) atualizado(s).')
            return redirect(item_focus_url(self.object))

        self.object = form.save(commit=False)
        self.object.texto = apply_item_red_mode_to_text(self.object.texto, mode)
        self.object.save()
        if mode == ITEM_RED_MODE_ALL_WITH_CHILDREN:
            apply_red_to_item_descendants(self.object)
        _touch_termo(self.object.sessao.termo, self.request)
        return redirect(item_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar item'
        context['ctrl_enter_submit'] = True
        parser_context(context, self.request, 4)
        context['voltar_url'] = item_focus_url(self.object)
        return context


class ItemDeleteView(SuperuserRequiredMixin, DeleteView):
    model = ItemTR
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return ItemTR.objects.filter(sessao__termo__sessoes__id=self.kwargs['sessao_pk'], sessao__termo__in=owned_queryset(TermoReferencia, self.request.user))

    def form_valid(self, form):
        sessao = self.object.sessao
        parent_id = self.object.parent_id
        return_url = item_delete_return_url(self.object)
        self.object.delete()
        normalize_sessoes(sessao.termo)
        normalize_items(sessao, parent_id)
        _touch_termo(sessao.termo, self.request)
        return redirect(return_url)


class ItemClearChildrenView(SuperuserRequiredMixin, DeleteView):
    model = ItemTR
    template_name = 'licitacoes/confirm_delete.html'

    def get_queryset(self):
        return ItemTR.objects.filter(sessao__termo__sessoes__id=self.kwargs['sessao_pk'], sessao__termo__in=owned_queryset(TermoReferencia, self.request.user))

    def form_valid(self, form):
        removed = clear_item_children(self.object)
        _touch_termo(self.object.sessao.termo, self.request)
        messages.success(self.request, f'{removed} subitem(ns) removido(s).')
        return redirect(item_focus_url(self.object))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Limpar filhos'
        context['descricao'] = 'Esta acao exclui todos os subitens deste item e nao pode ser desfeita.'
        context['pergunta'] = f'Confirma a exclusao de todos os subitens de "{self.object}"?'
        context['botao_confirmar'] = 'Limpar filhos'
        return context


class ItemMoveView(SuperuserRequiredMixin, View):
    template_name = 'licitacoes/item_move.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_object_or_404(ItemTR, pk=kwargs['pk'], sessao__termo__sessoes__id=kwargs['sessao_pk'], sessao__termo__in=owned_queryset(TermoReferencia, request.user))
        self.termo = self.item.sessao.termo
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        form = ItemMoveForm(termo=self.termo, item=self.item)
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'termo': self.termo, 'modo': 'mover'})

    def post(self, request, *args, **kwargs):
        form = ItemMoveForm(request.POST, termo=self.termo, item=self.item)
        if form.is_valid():
            target_token = form.cleaned_data['target']
            action = form.cleaned_data['action']
            child_position = form.cleaned_data.get('child_position')
            target = None
            target_sessao = None
            if target_token.startswith('item:'):
                target = get_object_or_404(ItemTR, pk=int(target_token.split(':', 1)[1]), sessao__termo=self.termo)
            else:
                target_sessao = get_object_or_404(SessaoTR, pk=int(target_token.split(':', 1)[1]), termo=self.termo)
                action = 'child'
            try:
                move_item(self.item, target, action, target_sessao=target_sessao, child_position=child_position)
                _touch_termo(self.termo, request)
                messages.success(request, 'Item movido e estrutura renumerada.')
                return redirect(item_focus_url(self.item))
            except ValueError as exc:
                form.add_error(None, str(exc))
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'termo': self.termo, 'modo': 'mover'})


class ItemDuplicateView(SuperuserRequiredMixin, View):
    template_name = 'licitacoes/item_move.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_object_or_404(ItemTR, pk=kwargs['pk'], sessao__termo__sessoes__id=kwargs['sessao_pk'], sessao__termo__in=owned_queryset(TermoReferencia, request.user))
        self.termo = self.item.sessao.termo
        return super().dispatch(request, *args, **kwargs)

    def get(self, request, *args, **kwargs):
        form = ItemMoveForm(termo=self.termo, item=self.item, action_label='Duplicar')
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'termo': self.termo, 'modo': 'duplicar'})

    def post(self, request, *args, **kwargs):
        form = ItemMoveForm(request.POST, termo=self.termo, item=self.item, action_label='Duplicar')
        if form.is_valid():
            target_token = form.cleaned_data['target']
            action = form.cleaned_data['action']
            child_position = form.cleaned_data.get('child_position')
            target = None
            target_sessao = None
            if target_token.startswith('item:'):
                target = get_object_or_404(ItemTR, pk=int(target_token.split(':', 1)[1]), sessao__termo=self.termo)
            else:
                target_sessao = get_object_or_404(SessaoTR, pk=int(target_token.split(':', 1)[1]), termo=self.termo)
                action = 'child'
            try:
                duplicate = duplicate_item(self.item, target, action, target_sessao=target_sessao, child_position=child_position)
                _touch_termo(self.termo, request)
                messages.success(request, 'Item duplicado com subitens e estrutura renumerada.')
                return redirect(item_focus_url(duplicate))
            except ValueError as exc:
                form.add_error(None, str(exc))
        return _render(request, self.template_name, {'form': form, 'item': self.item, 'termo': self.termo, 'modo': 'duplicar'})


class TabelaItemCreateView(SuperuserRequiredMixin, CreateView):
    model = TabelaItemLinha
    form_class = TabelaItemLinhaForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_item_tabela_1_1(kwargs['sessao_pk'], kwargs['item_pk'], request.user)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.item = self.item
        self.object.ordem = (self.item.tabela_linhas.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        self.object.save()
        _touch_termo(self.item.sessao.termo, self.request)
        return redirect(tabela_item_url(self.item))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Novo item da tabela - Item {item_indice(self.item)}'
        context['voltar_url'] = tabela_item_url(self.item)
        return context


class TabelaItemUpdateView(SuperuserRequiredMixin, UpdateView):
    model = TabelaItemLinha
    form_class = TabelaItemLinhaForm
    template_name = 'licitacoes/form.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_item_tabela_1_1(kwargs['sessao_pk'], kwargs['item_pk'], request.user)
        self.object = get_object_or_404(TabelaItemLinha, pk=kwargs['pk'], item=self.item)
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def get_success_url(self):
        _touch_termo(self.item.sessao.termo, self.request)
        return tabela_item_url(self.item)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Editar item da tabela - Item {item_indice(self.item)}'
        context['voltar_url'] = tabela_item_url(self.item)
        return context


class TabelaItemDeleteView(SuperuserRequiredMixin, DeleteView):
    model = TabelaItemLinha
    template_name = 'licitacoes/confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.item = get_item_tabela_1_1(kwargs['sessao_pk'], kwargs['item_pk'], request.user)
        self.object = get_object_or_404(TabelaItemLinha, pk=kwargs['pk'], item=self.item)
        return super().dispatch(request, *args, **kwargs)

    def get_object(self, queryset=None):
        return self.object

    def form_valid(self, form):
        self.object.delete()
        for idx, linha in enumerate(self.item.tabela_linhas.order_by('ordem', 'id'), start=1):
            if linha.ordem != idx:
                linha.ordem = idx
                linha.save(update_fields=['ordem'])
        _touch_termo(self.item.sessao.termo, self.request)
        return redirect(tabela_item_url(self.item))

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir item da tabela'
        context['voltar_url'] = tabela_item_url(self.item)
        return context


class TermoExportDocxView(SuperuserRequiredMixin, View):
    def get(self, request, pk):
        termo = owned_object_or_404(TermoReferencia, request.user, pk=pk)
        from docx import Document
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Cm, Pt, RGBColor

        document = Document()
        normal = document.styles['Normal']
        normal.font.name = 'Verdana'
        normal.font.size = Pt(10)
        section = document.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f'TR - {termo.nome}')
        r.bold = True
        r.font.size = Pt(12)

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        meta.add_run('Processo: ').bold = True
        meta.add_run(termo.numero_processo or '-')
        if termo.link:
            link = document.add_paragraph()
            link.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            link.add_run('Link: ').bold = True
            link.add_run(termo.link)

        def add_marked_runs(paragraph, text):
            for segment, is_red in red_mark_segments(text):
                lines = segment.split('\n')
                for idx, line in enumerate(lines):
                    if idx:
                        paragraph.add_run().add_break()
                    if not line:
                        continue
                    run = paragraph.add_run(line)
                    if is_red:
                        run.font.color.rgb = RGBColor(255, 0, 0)

        for bloco in build_termo_tree(termo):
            h = document.add_paragraph()
            h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = h.add_run(f"{bloco['sessao'].ordem}. {bloco['sessao'].titulo}")
            run.bold = True

            rows = bloco['rows']
            if not rows:
                p = document.add_paragraph(f"{bloco['sessao'].ordem}.1. -")
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for row in rows:
                p = document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if row.get('is_subsecao'):
                    p.add_run(row['item'].texto or '').bold = True
                else:
                    prefix = row['enum_prefix'] or f"{row['indice']}."
                    run_prefix = p.add_run(f'{prefix} ')
                    run_prefix.bold = True
                    add_marked_runs(p, row['item'].texto or '')

                if row['indice'] == '1.1' and row.get('tabela_linhas'):
                    doc_table = document.add_table(rows=1, cols=6)
                    doc_table.style = 'Table Grid'
                    doc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                    headers = ['Item', 'Descricao', 'CATMAT/CATSER', 'Siafisico', 'UF', 'Quantidade']
                    for idx, title in enumerate(headers):
                        cell = doc_table.rows[0].cells[idx]
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        header_run = paragraph.add_run(title)
                        header_run.bold = True
                        header_run.font.size = Pt(8)
                        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                    for linha in row['tabela_linhas']:
                        cells = doc_table.add_row().cells
                        values = [
                            str(linha.ordem),
                            linha.descricao or '-',
                            linha.catmat_catser or '-',
                            linha.siafisico or '-',
                            linha.unidade_fornecimento or '-',
                            quantidade_text(linha.quantidade),
                        ]
                        for idx, value in enumerate(values):
                            if idx == 1:
                                add_marked_runs(cells[idx].paragraphs[0], value)
                            else:
                                cells[idx].text = value
                            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            document.add_paragraph('')

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="tr_{slugify(termo.nome) or termo.pk}.docx"'
        return response


def _docx_response(titulo, secoes, filename):
    """Gera um DOCX simples a partir de seções já renderizadas."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    document = Document()
    normal = document.styles['Normal']
    normal.font.name = 'Verdana'
    normal.font.size = Pt(10)
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(titulo)
    r.bold = True
    r.font.size = Pt(12)

    for secao in secoes:
        h = document.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = h.add_run(f"{secao['numero']}. {secao['titulo']}")
        run.bold = True
        for entrada in secao['entradas'] or [f"{secao['numero']}.1. -"]:
            p = document.add_paragraph(entrada)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_paragraph('')

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


def _etp_dynamic_docx_response(etp):
    """Gera o DOCX do ETP TIC criado pelo editor dinâmico."""

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    document = Document()
    normal = document.styles['Normal']
    normal.font.name = 'Verdana'
    normal.font.size = Pt(10)
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f'ETP TIC - {etp.nome}')
    r.bold = True
    r.font.size = Pt(12)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    meta.add_run('Processo: ').bold = True
    meta.add_run(etp.numero_processo or '-')
    if etp.link:
        link = document.add_paragraph()
        link.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        link.add_run('Link: ').bold = True
        link.add_run(etp.link)

    def add_marked_runs(paragraph, text):
        for segment, is_red in red_mark_segments(text):
            lines = segment.split('\n')
            for idx, line in enumerate(lines):
                if idx:
                    paragraph.add_run().add_break()
                if not line:
                    continue
                run = paragraph.add_run(line)
                if is_red:
                    run.font.color.rgb = RGBColor(255, 0, 0)

    for bloco in build_etp_tree(etp):
        h = document.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = h.add_run(f"{bloco['sessao'].ordem}. {bloco['sessao'].titulo}")
        run.bold = True

        rows = bloco['rows']
        if not rows:
            p = document.add_paragraph(f"{bloco['sessao'].ordem}.1. -")
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for row in rows:
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            if row.get('is_subsecao'):
                p.add_run(row['item'].texto or '').bold = True
            else:
                prefix = row.get('enum_prefix') or f"{row['indice']}."
                run_prefix = p.add_run(f'{prefix} ')
                run_prefix.bold = True
                add_marked_runs(p, row['item'].texto or '')
        document.add_paragraph('')

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="etp_tic_{slugify(etp.nome) or etp.pk}.docx"'
    return response


def _pesquisa_preco_xlsx_response(pesquisa):
    """Preenche o modelo XLSX do quadro comparativo de pesquisa de preço."""

    from openpyxl import load_workbook
    from openpyxl.utils import get_column_letter

    context = pesquisa_preco_context(pesquisa)
    itens = context['itens']
    fornecedores = context['fornecedores']
    medias = context['medias']
    is_servico = pesquisa.tipo == PesquisaPreco.Tipo.SERVICO
    template_path = settings.BASE_DIR / 'docs' / 'Quadro Comparativo Pesquisa de Preços-modelo.xlsx'
    wb = load_workbook(template_path)
    ws = wb['PCs - Tab Alternativa']
    for sheet in list(wb.worksheets):
        if sheet.title != ws.title:
            wb.remove(sheet)

    def copy_cell_style(source, target):
        target._style = copy(source._style)
        if source.has_style:
            target.font = copy(source.font)
            target.fill = copy(source.fill)
            target.border = copy(source.border)
            target.alignment = copy(source.alignment)
            target.number_format = source.number_format
            target.protection = copy(source.protection)

    def copy_row_style(source_row, target_row):
        ws.row_dimensions[target_row].height = ws.row_dimensions[source_row].height
        for col_idx in range(1, ws.max_column + 1):
            copy_cell_style(ws.cell(source_row, col_idx), ws.cell(target_row, col_idx))

    def merge_if_needed(start_row, start_col, end_row, end_col):
        ref = f'{get_column_letter(start_col)}{start_row}:{get_column_letter(end_col)}{end_row}'
        if ref not in [str(item) for item in ws.merged_cells.ranges]:
            ws.merge_cells(ref)

    supplier_count = max(3, len(fornecedores))
    extra_suppliers = max(0, supplier_count - 3)
    if extra_suppliers:
        ws.insert_cols(16, amount=extra_suppliers * 2)
        for supplier_idx in range(3, supplier_count):
            col = 16 + ((supplier_idx - 3) * 2)
            for row_idx in range(1, 21):
                copy_cell_style(ws.cell(row_idx, 14), ws.cell(row_idx, col))
                copy_cell_style(ws.cell(row_idx, 15), ws.cell(row_idx, col + 1))
            ws.column_dimensions[get_column_letter(col)].width = ws.column_dimensions['N'].width
            ws.column_dimensions[get_column_letter(col + 1)].width = ws.column_dimensions['O'].width
            for row_idx in [3, 4, 5, 6, 7, 8, 9, 15, 16, 17, 18]:
                merge_if_needed(row_idx, col, row_idx, col + 1)

    item_count = max(4, len(itens))
    extra_items = max(0, item_count - 4)
    if extra_items:
        ws.insert_rows(15, amount=extra_items)
        for row_idx in range(15, 15 + extra_items):
            copy_row_style(14, row_idx)
            merge_if_needed(row_idx, 5, row_idx, 7)

    total_row = 11 + item_count
    total_contratacao_row = total_row + 1
    validade_row = total_row + 2
    prazo_row = total_row + 3
    data_row = total_row + 5
    avg_col = 16 + (extra_suppliers * 2)
    avg_total_col = avg_col + 1
    supplier_cols = [10 + (idx * 2) for idx in range(supplier_count)]
    actual_supplier_cols = supplier_cols[:len(fornecedores)]

    for row_idx in range(3, data_row + 1):
        for col_idx in range(2, avg_total_col + 1):
            if not isinstance(ws.cell(row_idx, col_idx).__class__.__name__, str):
                continue
            try:
                ws.cell(row_idx, col_idx).value = None
            except AttributeError:
                pass

    ws['B2'] = 'QUADRO COMPARATIVO DE PESQUISA DE PREÇOS'
    ws['B4'] = 'Pesquisa realizada por:'
    ws['B7'] = 'Nome'
    ws['D7'] = pesquisa.pesquisador_nome
    ws['B8'] = 'E-mail'
    ws['D8'] = pesquisa.pesquisador_email
    ws['B9'] = 'Cargo'
    ws['D9'] = pesquisa.pesquisador_cargo
    ws['G4'] = 'Empresa'
    ws['G5'] = 'CNPJ'
    ws['G6'] = 'Telefone'
    ws['G7'] = 'Contato'
    ws['G8'] = 'E-mail do contato'
    ws['G9'] = ''

    ws['B10'] = 'Item'
    ws['C10'] = 'Unidade'
    ws['D10'] = 'Qtd'
    ws['E10'] = 'Descrição'
    ws['H4'] = 'CATMAT/\nCATSER'
    ws['I4'] = 'SIAFISICO'

    for idx, col in enumerate(supplier_cols):
        fornecedor_entry = fornecedores[idx] if idx < len(fornecedores) else None
        ws.cell(3, col, f'EMPRESA {idx + 1}' if fornecedor_entry else '')
        if fornecedor_entry:
            fornecedor = fornecedor_entry['fornecedor']
            ws.cell(4, col, fornecedor.razao_social)
            ws.cell(5, col, fornecedor.cnpj)
            ws.cell(6, col, fornecedor.telefone)
            ws.cell(7, col, fornecedor.contato)
            ws.cell(8, col, fornecedor.email_contato)
            ws.cell(9, col, None)
        else:
            for row_idx in range(4, 10):
                ws.cell(row_idx, col, None)
        ws.cell(10, col, 'Unitário')
        ws.cell(10, col + 1, 'Total')

    ws.cell(3, avg_col, 'PREÇO MÉDIO')
    ws.cell(10, avg_col, 'Unitário')
    ws.cell(10, avg_total_col, 'Total')

    for row_offset in range(item_count):
        row_idx = 11 + row_offset
        item = itens[row_offset] if row_offset < len(itens) else None
        media = medias[row_offset] if row_offset < len(medias) else None
        if item:
            ws.cell(row_idx, 2, item.ordem)
            ws.cell(row_idx, 3, item.unidade_fornecimento)
            ws.cell(row_idx, 4, float(item.quantidade))
            ws.cell(row_idx, 5, item.descricao)
            ws.cell(row_idx, 8, item.catmat_catser)
            ws.cell(row_idx, 9, item.siafisico)
        else:
            for col_idx in range(2, avg_total_col + 1):
                ws.cell(row_idx, col_idx, None)
            continue
        for entry, col in zip(fornecedores, actual_supplier_cols):
            match = next((item_row for item_row in entry['itens'] if item_row['item'].id == item.id), None)
            ws.cell(row_idx, col, float(match['preco_unitario']) if match and match['preco_unitario'] is not None else None)
            ws.cell(row_idx, col + 1, f'={get_column_letter(col)}{row_idx}*D{row_idx}' if match and match['preco_unitario'] is not None else None)
        unit_refs = [f'{get_column_letter(col)}{row_idx}' for col in actual_supplier_cols]
        ws.cell(row_idx, avg_col, f'=AVERAGE({",".join(unit_refs)})' if unit_refs else None)
        ws.cell(row_idx, avg_total_col, f'={get_column_letter(avg_col)}{row_idx}*D{row_idx}' if media and media['preco_medio'] is not None else None)

    item_start = 11
    item_end = 10 + item_count
    ws.cell(total_row, 2, 'TOTAL MENSAL' if is_servico else 'TOTAL')
    for entry, col in zip(fornecedores, actual_supplier_cols):
        ws.cell(total_row, col, f'=SUM({get_column_letter(col + 1)}{item_start}:{get_column_letter(col + 1)}{item_end})')
    ws.cell(total_row, avg_col, f'=SUM({get_column_letter(avg_total_col)}{item_start}:{get_column_letter(avg_total_col)}{item_end})')

    if is_servico:
        ws.cell(total_contratacao_row, 2, f'TOTAL PARA CONTRATAÇÃO ({pesquisa.vigencia_meses} MESES)')
        for entry, col in zip(fornecedores, actual_supplier_cols):
            ws.cell(total_contratacao_row, col, f'={get_column_letter(col)}{total_row}*{pesquisa.vigencia_meses}')
        ws.cell(total_contratacao_row, avg_col, f'={get_column_letter(avg_col)}{total_row}*{pesquisa.vigencia_meses}')
    else:
        for col_idx in range(2, avg_total_col + 1):
            ws.cell(total_contratacao_row, col_idx, None)

    ws.cell(validade_row, 2, 'Validade do Orçamento (em dias)')
    for entry, col in zip(fornecedores, actual_supplier_cols):
        ws.cell(validade_row, col, entry['pesquisa_fornecedor'].validade_orcamento_dias)
    validade_refs = [f'{get_column_letter(col)}{validade_row}' for col in actual_supplier_cols]
    ws.cell(validade_row, avg_col, f'=AVERAGE({",".join(validade_refs)})' if validade_refs else None)

    ws.cell(prazo_row, 2, 'Prazo de Entrega/Execução (em dias)')
    for col in actual_supplier_cols:
        ws.cell(prazo_row, col, None)
    ws.cell(prazo_row, avg_col, None)
    ws.cell(data_row, 2, f'São Paulo, {timezone.localdate():%d/%m/%Y}')

    buffer = BytesIO()
    wb.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="pesquisa_preco_{slugify(pesquisa.termo.nome) or pesquisa.pk}.xlsx"'
    return response


def _dfd_docx_response(dfd):
    """Gera o DOCX do DFD com textos, marcações e tabela de itens."""

    from docx import Document
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    def apply_run_style(run, size=10):
        run.font.name = 'Verdana'
        run.font.size = Pt(size)

    def add_marked_runs(paragraph, text, size=10):
        for segment, is_red in red_mark_segments(text):
            lines = segment.split('\n')
            for idx, line in enumerate(lines):
                if idx:
                    paragraph.add_run().add_break()
                if not line:
                    continue
                run = paragraph.add_run(line)
                apply_run_style(run, size)
                if is_red:
                    run.font.color.rgb = RGBColor(255, 0, 0)

    def apply_cell_font(cell, size=8):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                apply_run_style(run, size)

    document = Document()
    normal = document.styles['Normal']
    normal.font.name = 'Verdana'
    normal.font.size = Pt(10)
    section = document.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f'DFD - {dfd.nome}')
    apply_run_style(r)
    r.bold = True

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, bold in [
        ('Processo: ', True),
        (dfd.numero_processo or '-', False),
        (' | Status: ', True),
        (dfd.get_status_display(), False),
    ]:
        run = meta.add_run(text)
        apply_run_style(run)
        run.bold = bold

    for secao in render_dfd_sections(dfd):
        if secao['numero'] == 1:
            alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif secao['numero'] == 6:
            alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        h = document.add_paragraph()
        h.alignment = alignment
        run = h.add_run(secao['rotulo'])
        apply_run_style(run)
        run.bold = True

        for entrada in secao['entradas'] or ['-']:
            p = document.add_paragraph()
            add_marked_runs(p, entrada)
            p.alignment = alignment

        if secao.get('tabela'):
            table = document.add_table(rows=1, cols=8)
            table.style = 'Table Grid'
            headers = [
                'Item',
                'Especificacao',
                'CATMAT',
                'SIAFISICO',
                'Unidade de medida',
                'Quantidade',
                'Valor Unitario',
                'Valor total',
            ]
            for idx, title in enumerate(headers):
                cell = table.rows[0].cells[idx]
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                header_run = paragraph.add_run(title)
                apply_run_style(header_run, 8)
                header_run.bold = True
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for linha in secao['tabela']:
                cells = table.add_row().cells
                values = [
                    str(linha.ordem),
                    linha.especificacao,
                    linha.catmat or '-',
                    linha.siafisico or '-',
                    linha.unidade_medida or '-',
                    str(linha.quantidade),
                    str(linha.valor_unitario),
                    str(linha.valor_total),
                ]
                for idx, value in enumerate(values):
                    if idx == 1:
                        add_marked_runs(cells[idx].paragraphs[0], value, size=8)
                    else:
                        cells[idx].text = value
                        apply_cell_font(cells[idx], 8)
                    cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for entrada in secao.get('entradas_apos_tabela', []):
            p = document.add_paragraph()
            add_marked_runs(p, entrada)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        document.add_paragraph('')

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="dfd_{slugify(dfd.nome) or dfd.pk}.docx"'
    return response


def _render(request, template_name, context):
    """Importa render localmente para manter os imports pesados próximos ao uso."""

    from django.shortcuts import render

    return render(request, template_name, context)
