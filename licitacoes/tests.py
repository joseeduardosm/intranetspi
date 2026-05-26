from io import BytesIO

from django.contrib.auth import get_user_model
from django.core.files.uploadedfile import SimpleUploadedFile
from django.utils import timezone
from django.test import TestCase
from django.urls import reverse
from openpyxl import load_workbook

from .models import (
    Dfd,
    DfdItemTabela,
    EtpTic,
    Fornecedor,
    ItemEtpTic,
    ItemTR,
    PesquisaPreco,
    PesquisaPrecoContato,
    PesquisaPrecoFornecedor,
    PesquisaPrecoItemValor,
    SessaoEtpTic,
    SessaoTR,
    TabelaItemLinha,
    TermoReferencia,
)
from .forms import FornecedorForm
from .services import (
    build_etp_item_rows,
    build_item_rows,
    clear_etp_item_children,
    clear_item_children,
    duplicate_dfd,
    duplicate_etp,
    duplicate_etp_item,
    duplicate_item,
    duplicate_termo,
    item_parent_for_tipo,
    move_etp_item,
    move_item,
    pesquisa_preco_context,
    red_marked_html,
    render_dfd_sections,
    render_etp_sections,
)


class LoginTests(TestCase):
    def test_login_rejeita_usuario_nao_superuser(self):
        User = get_user_model()
        User.objects.create_user(username='comum', password='123')
        response = self.client.post(reverse('login'), {'username': 'comum', 'password': '123'})
        self.assertEqual(response.status_code, 200)
        self.assertNotIn('_auth_user_id', self.client.session)


class EtpTicTests(TestCase):
    def test_render_numera_paragrafos_por_secao(self):
        etp = EtpTic.objects.create(
            nome='Contratacao',
            numero_processo='001/2026',
            descricao_necessidade='Primeiro paragrafo.\n\nSegundo paragrafo.',
        )
        secao = render_etp_sections(etp)[1]
        self.assertEqual(secao['entradas'][0], '2.1. Primeiro paragrafo.')
        self.assertEqual(secao['entradas'][1], '2.2. Segundo paragrafo.')

    def test_editar_etp_sem_secao_abre_dados_basicos(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)

        response = self.client.get(reverse('licitacoes:etp_edit', args=[etp.pk]))

        self.assertContains(response, 'Editar ETP TIC')
        self.assertContains(response, 'name="nome"')
        self.assertNotContains(response, 'Descricao da Necessidade')

    def test_salvar_dados_basicos_etp_redireciona_para_preview(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', link='https://example.com', usa_editor_dinamico=True)

        response = self.client.post(
            reverse('licitacoes:etp_edit', args=[etp.pk]),
            {'nome': 'ETP atualizado', 'numero_processo': '002/2026', 'link': ''},
        )

        etp.refresh_from_db()
        self.assertEqual(etp.nome, 'ETP atualizado')
        self.assertRedirects(response, reverse('licitacoes:etp_preview', args=[etp.pk]), fetch_redirect_response=False)

    def test_criar_etp_novo_nasce_dinamico_sem_sessoes(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')

        response = self.client.post(
            reverse('licitacoes:etp_create'),
            {'nome': 'ETP novo', 'numero_processo': '001/2026', 'link': ''},
        )

        etp = EtpTic.objects.get(nome='ETP novo')
        self.assertTrue(etp.usa_editor_dinamico)
        self.assertEqual(etp.sessoes.count(), 0)
        self.assertRedirects(response, reverse('licitacoes:etp_detail', args=[etp.pk]), fetch_redirect_response=False)

    def test_etp_legado_nao_edita_e_abre_preview(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP legado', numero_processo='001/2026', descricao_necessidade='Necessidade.')

        response = self.client.get(reverse('licitacoes:etp_edit', args=[etp.pk]))
        preview = self.client.get(reverse('licitacoes:etp_preview', args=[etp.pk]))

        self.assertRedirects(response, reverse('licitacoes:etp_preview', args=[etp.pk]), fetch_redirect_response=False)
        self.assertContains(preview, 'ETP TIC legado disponivel somente para visualizacao.')
        self.assertContains(preview, 'Necessidade.')

    def test_etp_dinamico_cria_sessao_item_e_subitem(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)

        response_sessao = self.client.post(reverse('licitacoes:etp_sessao_create', args=[etp.pk]), {'titulo': 'Objeto'})
        sessao = etp.sessoes.get()
        response_item = self.client.post(reverse('licitacoes:etp_item_create', args=[sessao.pk]), {'texto': 'Item'})
        item = sessao.itens.get(texto='Item')
        response_subitem = self.client.post(reverse('licitacoes:etp_item_child_create', args=[sessao.pk, item.pk]), {'texto': 'Subitem'})

        rows = build_etp_item_rows(sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertRedirects(response_sessao, f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#sessao-etp-{sessao.pk}", fetch_redirect_response=False)
        self.assertRedirects(response_item, f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{item.pk}", fetch_redirect_response=False)
        self.assertRedirects(response_subitem, f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{sessao.itens.get(texto='Subitem').pk}", fetch_redirect_response=False)
        self.assertEqual(found['Item']['indice'], '1.1')
        self.assertEqual(found['Subitem']['indice'], '1.1.1')

    def test_criar_subitens_etp_por_marcador_hash_duplo(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        parent = ItemEtpTic.objects.create(sessao=sessao, texto='Memoria RAM', ordem=1)

        response = self.client.post(
            reverse('licitacoes:etp_item_child_create', args=[sessao.pk, parent.pk]),
            {
                'texto': (
                    '##Deverao ser fornecidos no minimo 384 GBytes de memoria RAM por servidor.\n'
                    '##Padrao minimo do tipo DDR-4 ECC 2666MHz ou superior.'
                )
            },
        )

        filhos = list(parent.filhos.order_by('ordem', 'id'))
        self.assertEqual(len(filhos), 2)
        self.assertEqual(filhos[0].texto, 'Deverao ser fornecidos no minimo 384 GBytes de memoria RAM por servidor.')
        self.assertEqual(filhos[1].texto, 'Padrao minimo do tipo DDR-4 ECC 2666MHz ou superior.')
        self.assertNotIn('##', filhos[0].texto)
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{filhos[0].pk}",
            fetch_redirect_response=False,
        )

    def test_criar_item_etp_com_subitens_por_marcadores_hash(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        parent = ItemEtpTic.objects.create(sessao=sessao, texto='Hardware', ordem=1)

        response = self.client.post(
            reverse('licitacoes:etp_item_child_create', args=[sessao.pk, parent.pk]),
            {
                'texto': (
                    '#SLOTS PCI\n'
                    '##Padrao PCI-Express ou superior.\n'
                    '##Disponibilizar no minimo 02 slots PCI-Express livres.'
                )
            },
        )

        item = parent.filhos.get(texto='SLOTS PCI')
        filhos = list(item.filhos.order_by('ordem', 'id'))
        self.assertEqual(item.ordem, 1)
        self.assertEqual(len(filhos), 2)
        self.assertEqual(filhos[0].texto, 'Padrao PCI-Express ou superior.')
        self.assertEqual(filhos[1].texto, 'Disponibilizar no minimo 02 slots PCI-Express livres.')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{item.pk}",
            fetch_redirect_response=False,
        )

    def test_criar_hierarquia_etp_por_marcadores_ate_cinco_niveis(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        parent = ItemEtpTic.objects.create(sessao=sessao, texto='Item base', ordem=1)

        response = self.client.post(
            reverse('licitacoes:etp_item_child_create', args=[sessao.pk, parent.pk]),
            {
                'texto': (
                    '#Nivel 1\n'
                    '##Nivel 2\n'
                    '###Nivel 3\n'
                    '####Nivel 4\n'
                    '#####Nivel 5\n'
                    '#Outro nivel 1'
                )
            },
        )

        rows = build_etp_item_rows(sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertEqual(found['Nivel 1']['indice'], '1.1.1')
        self.assertEqual(found['Nivel 2']['indice'], '1.1.1.1')
        self.assertEqual(found['Nivel 3']['indice'], '1.1.1.1.1')
        self.assertEqual(found['Nivel 4']['indice'], '1.1.1.1.1.1')
        self.assertEqual(found['Nivel 5']['indice'], '1.1.1.1.1.1.1')
        self.assertEqual(found['Outro nivel 1']['indice'], '1.1.2')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{found['Nivel 1']['item'].pk}",
            fetch_redirect_response=False,
        )

    def test_criar_hierarquia_etp_por_marcadores_em_item_raiz(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)

        response = self.client.post(
            reverse('licitacoes:etp_item_create', args=[sessao.pk]),
            {
                'texto': (
                    '#Item raiz\n'
                    '##Subitem\n'
                    '###Subitem interno\n'
                    '#Outro item raiz'
                )
            },
        )

        rows = build_etp_item_rows(sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertEqual(found['Item raiz']['indice'], '1.1')
        self.assertEqual(found['Subitem']['indice'], '1.1.1')
        self.assertEqual(found['Subitem interno']['indice'], '1.1.1.1')
        self.assertEqual(found['Outro item raiz']['indice'], '1.2')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{found['Item raiz']['item'].pk}",
            fetch_redirect_response=False,
        )

    def test_criar_etp_com_subsecao_e_inciso_por_marcadores(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)

        response = self.client.post(
            reverse('licitacoes:etp_item_create', args=[sessao.pk]),
            {
                'texto': (
                    '@Garantia da contratação\n'
                    '#Será exigida a garantia.\n'
                    '**Caução em dinheiro.\n'
                    '$$Em moeda corrente.\n'
                    '**Fiança bancária.\n'
                    '@Sustentabilidade\n'
                    '#Além dos critérios, devem ser atendidos requisitos.'
                )
            },
        )

        rows = build_etp_item_rows(sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertTrue(found['Garantia da contratação']['is_subsecao'])
        self.assertEqual(found['Garantia da contratação']['indice'], '')
        self.assertEqual(found['Será exigida a garantia.']['indice'], '1.1')
        self.assertEqual(found['Caução em dinheiro.']['enum_prefix'], 'I.')
        self.assertEqual(found['Em moeda corrente.']['enum_prefix'], 'a)')
        self.assertEqual(found['Fiança bancária.']['enum_prefix'], 'II.')
        self.assertEqual(found['Além dos critérios, devem ser atendidos requisitos.']['indice'], '1.2')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{found['Garantia da contratação']['item'].pk}",
            fetch_redirect_response=False,
        )

    def test_criar_etp_com_marcadores_destaca_todo_texto_em_vermelho(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)

        self.client.post(
            reverse('licitacoes:etp_item_create', args=[sessao.pk]),
            {
                'texto': '@Garantia\n#Será exigida garantia.',
                'modo_destaque_texto': 'todo_vermelho',
            },
        )

        textos = set(sessao.itens.values_list('texto', flat=True))
        self.assertIn('*Garantia*', textos)
        self.assertIn('*Será exigida garantia.*', textos)

    def test_criar_alineas_etp_com_marcador_isolado_no_subitem(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        subitem = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Subitem', ordem=1)

        self.client.post(
            reverse('licitacoes:etp_item_child_create', args=[sessao.pk, subitem.pk]),
            {'texto': '$$o prazo de validade;\n$$a data da emissão;'},
        )

        rows = build_etp_item_rows(sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertEqual(found['o prazo de validade;']['enum_prefix'], 'a)')
        self.assertEqual(found['a data da emissão;']['enum_prefix'], 'b)')
        self.assertEqual(found['o prazo de validade;']['item'].parent_id, subitem.id)

    def test_editar_etp_com_marcadores_substitui_item_e_renumera_irmaos(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        raiz = ItemEtpTic.objects.create(sessao=sessao, texto='Raiz', ordem=1)
        alvo = ItemEtpTic.objects.create(sessao=sessao, parent=raiz, texto='Alvo', ordem=1)
        filho_antigo = ItemEtpTic.objects.create(sessao=sessao, parent=alvo, texto='Filho antigo', ordem=1)
        posterior = ItemEtpTic.objects.create(sessao=sessao, parent=raiz, texto='Posterior', ordem=2)

        response = self.client.post(
            reverse('licitacoes:etp_item_update', args=[sessao.pk, alvo.pk]),
            {'texto': '#Novo\n##Detalhe\n#Outro'},
        )

        alvo.refresh_from_db()
        rows = build_etp_item_rows(sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertEqual(alvo.texto, 'Novo')
        self.assertEqual(found['Novo']['indice'], '1.1.1')
        self.assertEqual(found['Detalhe']['indice'], '1.1.1.1')
        self.assertEqual(found['Outro']['indice'], '1.1.2')
        self.assertEqual(found['Posterior']['indice'], '1.1.3')
        self.assertFalse(ItemEtpTic.objects.filter(pk=filho_antigo.pk).exists())
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{alvo.pk}",
            fetch_redirect_response=False,
        )

    def test_editar_etp_destaca_item_e_filhos_em_vermelho(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        filho = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Filho', ordem=1)
        neto = ItemEtpTic.objects.create(sessao=sessao, parent=filho, texto='Neto', ordem=1)

        self.client.post(
            reverse('licitacoes:etp_item_update', args=[sessao.pk, item.pk]),
            {
                'texto': 'Item alterado',
                'modo_destaque_texto': 'todo_vermelho_com_filhos',
            },
        )

        item.refresh_from_db()
        filho.refresh_from_db()
        neto.refresh_from_db()
        self.assertEqual(item.texto, '*Item alterado*')
        self.assertEqual(filho.texto, '*Filho*')
        self.assertEqual(neto.texto, '*Neto*')

    def test_editar_sessao_etp_destaca_filhos_em_vermelho(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        filho = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Filho', ordem=1)

        self.client.post(
            reverse('licitacoes:etp_sessao_update', args=[etp.pk, sessao.pk]),
            {'titulo': 'Objeto atualizado', 'filhos_em_vermelho': '1'},
        )

        item.refresh_from_db()
        filho.refresh_from_db()
        self.assertEqual(item.texto, '*Item*')
        self.assertEqual(filho.texto, '*Filho*')

    def test_mover_e_duplicar_item_etp_dinamico(self):
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        destino = ItemEtpTic.objects.create(sessao=sessao, texto='Destino', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Movido', ordem=2)
        child = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Filho', ordem=1)

        move_etp_item(item, destino, 'child')
        item.refresh_from_db()
        duplicate = duplicate_etp_item(item, destino, 'after')

        self.assertEqual(item.parent_id, destino.id)
        self.assertEqual(duplicate.texto, 'Movido')
        self.assertEqual(duplicate.filhos.get().texto, child.texto)

    def test_limpar_filhos_etp_remove_subitens_e_preserva_item(self):
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        child = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Subitem', ordem=1)
        grandchild = ItemEtpTic.objects.create(sessao=sessao, parent=child, texto='Neto', ordem=1)

        removed = clear_etp_item_children(item)

        self.assertEqual(removed, 2)
        self.assertTrue(ItemEtpTic.objects.filter(pk=item.pk).exists())
        self.assertFalse(ItemEtpTic.objects.filter(pk__in=[child.pk, grandchild.pk]).exists())

    def test_limpar_filhos_etp_pela_view_volta_para_item(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        child = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Subitem', ordem=1)

        response = self.client.post(reverse('licitacoes:etp_item_clear_children', args=[sessao.pk, item.pk]))

        self.assertFalse(ItemEtpTic.objects.filter(pk=child.pk).exists())
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#item-etp-{item.pk}",
            fetch_redirect_response=False,
        )

    def test_limpar_filhos_da_sessao_etp_preserva_sessao_e_remove_itens(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        child = ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Subitem', ordem=1)

        response = self.client.post(reverse('licitacoes:etp_sessao_clear_items', args=[etp.pk, sessao.pk]))

        self.assertTrue(SessaoEtpTic.objects.filter(pk=sessao.pk).exists())
        self.assertFalse(ItemEtpTic.objects.filter(pk__in=[item.pk, child.pk]).exists())
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:etp_detail', args=[etp.pk])}#sessao-etp-{sessao.pk}",
            fetch_redirect_response=False,
        )

    def test_duplicar_etp_dinamico_copia_sessoes_e_itens(self):
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Item', ordem=1)
        ItemEtpTic.objects.create(sessao=sessao, parent=item, texto='Subitem', ordem=1)

        duplicate = duplicate_etp(etp)

        self.assertTrue(duplicate.usa_editor_dinamico)
        self.assertEqual(duplicate.nome, 'Copia de ETP')
        new_sessao = duplicate.sessoes.get()
        new_item = new_sessao.itens.get(parent=None)
        self.assertEqual(new_item.texto, 'Item')
        self.assertEqual(new_item.filhos.get().texto, 'Subitem')

    def test_exporta_etp_dinamico_docx(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        ItemEtpTic.objects.create(sessao=sessao, texto='Texto *vermelho*', ordem=1)

        response = self.client.get(reverse('licitacoes:etp_export', args=[etp.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    def test_exporta_etp_dinamico_com_subsecao_e_inciso_docx(self):
        from io import BytesIO

        from docx import Document

        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        etp = EtpTic.objects.create(nome='ETP', numero_processo='001/2026', usa_editor_dinamico=True)
        sessao = SessaoEtpTic.objects.create(etp=etp, titulo='Objeto', ordem=1)
        subsecao = ItemEtpTic.objects.create(sessao=sessao, tipo=ItemEtpTic.Tipo.SUBSECAO, texto='Garantia da contratação', ordem=1)
        item = ItemEtpTic.objects.create(sessao=sessao, texto='Será exigida a garantia.', ordem=2)
        inciso = ItemEtpTic.objects.create(sessao=sessao, parent=item, tipo=ItemEtpTic.Tipo.INCISO, texto='Caução em dinheiro.', ordem=1)
        ItemEtpTic.objects.create(sessao=sessao, parent=inciso, tipo=ItemEtpTic.Tipo.ALINEA, texto='Em moeda corrente.', ordem=1)

        response = self.client.get(reverse('licitacoes:etp_export', args=[etp.pk]))

        document = Document(BytesIO(response.content))
        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertIn(subsecao.texto, texts)
        self.assertIn('1.1. Será exigida a garantia.', texts)
        self.assertIn('I. Caução em dinheiro.', texts)
        self.assertIn('a) Em moeda corrente.', texts)


class DfdTests(TestCase):
    def test_render_numera_apenas_secoes_aplicaveis(self):
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            informacoes_preliminares='Orgao: SEDS\n\nSetor: DIVTI',
            descricao_objeto='Primeiro paragrafo.\n\nSegundo paragrafo.',
            justificativa_necessidade='Justificativa.',
            responsaveis='Nenhum responsavel informado.',
        )

        secoes = render_dfd_sections(dfd)

        self.assertEqual(secoes[0]['entradas'], ['Orgao: SEDS', 'Setor: DIVTI'])
        self.assertEqual(secoes[1]['entradas'][0], '1.1. Primeiro paragrafo.\n\nSegundo paragrafo.')
        self.assertEqual(secoes[1]['entradas_apos_tabela'], [Dfd.OBJETO_NAO_LUXO_PADRAO])
        self.assertEqual(secoes[2]['entradas'][0], '2.1. Justificativa.')
        self.assertEqual(secoes[5]['entradas'], ['Nenhum responsavel informado.'])

    def test_render_dfd_item_1_2_normaliza_prefixo_na_exportacao(self):
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            descricao_objeto='Objeto.',
            objeto_nao_luxo='Item 1.2: O objeto desta contratação não se enquadra como sendo de bem de luxo.',
        )

        secao_objeto = render_dfd_sections(dfd)[1]

        self.assertEqual(
            secao_objeto['entradas_apos_tabela'],
            ['1.2. O objeto desta contratação não se enquadra como sendo de bem de luxo.'],
        )

    def test_preview_dfd_renderiza_responsaveis_centralizado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            responsaveis='Responsavel centralizado.',
        )

        response = self.client.get(reverse('licitacoes:dfd_preview', args=[dfd.pk]))

        self.assertContains(response, '<h2 class="h5 text-center">Responsaveis</h2>')
        self.assertContains(response, '<p class="mb-0 text-center">Responsavel centralizado.</p>')

    def test_criar_dfd_redireciona_para_edicao_secao_1(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')

        response = self.client.post(
            reverse('licitacoes:dfd_create'),
            {'nome': 'DFD', 'numero_processo': '001/2026'},
        )

        dfd = Dfd.objects.get(nome='DFD')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao=1",
            fetch_redirect_response=False,
        )

    def test_editar_dfd_sem_secao_abre_dados_basicos(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(nome='DFD', numero_processo='001/2026', descricao_objeto='Objeto.')

        response = self.client.get(reverse('licitacoes:dfd_edit', args=[dfd.pk]))

        self.assertContains(response, 'Editar DFD')
        self.assertContains(response, 'name="nome"')
        self.assertNotContains(response, 'Descricao Sucinta do Objeto')

    def test_salvar_dados_basicos_dfd_redireciona_para_preview(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(nome='DFD', numero_processo='001/2026')

        response = self.client.post(
            reverse('licitacoes:dfd_edit', args=[dfd.pk]),
            {'nome': 'DFD atualizado', 'numero_processo': '002/2026'},
        )

        dfd.refresh_from_db()
        self.assertEqual(dfd.nome, 'DFD atualizado')
        self.assertRedirects(response, reverse('licitacoes:dfd_preview', args=[dfd.pk]), fetch_redirect_response=False)

    def test_editar_dfd_avanca_para_proxima_secao(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(nome='DFD', numero_processo='001/2026')

        response = self.client.post(
            f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao=1",
            {'informacoes_preliminares': 'Orgao: SEDS', '_acao': 'proximo'},
        )

        dfd.refresh_from_db()
        self.assertEqual(dfd.secao_atual, 2)
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao=2",
            fetch_redirect_response=False,
        )

    def test_concluir_dfd_salva_responsaveis_da_secao_atual(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(nome='DFD', numero_processo='001/2026', secao_atual=6)

        response = self.client.post(
            f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao=6",
            {'responsaveis': 'Responsavel salvo.', '_acao': 'concluir'},
        )

        dfd.refresh_from_db()
        self.assertEqual(dfd.responsaveis, 'Responsavel salvo.')
        self.assertEqual(dfd.status, Dfd.Status.CONCLUIDO)
        self.assertRedirects(
            response,
            reverse('licitacoes:dfd_preview', args=[dfd.pk]),
            fetch_redirect_response=False,
        )

    def test_crud_linha_tabela_dfd(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(nome='DFD', numero_processo='001/2026')

        create_response = self.client.post(
            reverse('licitacoes:dfd_tabela_create', args=[dfd.pk]),
            {
                'especificacao': 'Notebook',
                'catmat': '123',
                'siafisico': '456',
                'unidade_medida': 'Unidade',
                'quantidade': '2',
                'valor_unitario': '100.50',
                'valor_total': '201.00',
            },
        )
        linha = dfd.itens_tabela.get()
        self.assertEqual(linha.ordem, 1)
        self.assertRedirects(
            create_response,
            f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao=2#dfd-tabela",
            fetch_redirect_response=False,
        )

        self.client.post(
            reverse('licitacoes:dfd_tabela_update', args=[dfd.pk, linha.pk]),
            {
                'especificacao': 'Desktop',
                'catmat': '123',
                'siafisico': '456',
                'unidade_medida': 'Unidade',
                'quantidade': '3',
                'valor_unitario': '150.00',
                'valor_total': '450.00',
            },
        )
        linha.refresh_from_db()
        self.assertEqual(linha.especificacao, 'Desktop')
        self.assertEqual(linha.quantidade, 3)
        self.assertEqual(str(linha.valor_total), '450.00')

        delete_response = self.client.post(reverse('licitacoes:dfd_tabela_delete', args=[dfd.pk, linha.pk]))
        self.assertFalse(DfdItemTabela.objects.filter(pk=linha.pk).exists())
        self.assertRedirects(
            delete_response,
            f"{reverse('licitacoes:dfd_edit', args=[dfd.pk])}?secao=2#dfd-tabela",
            fetch_redirect_response=False,
        )

    def test_exporta_dfd_docx_com_tabela(self):
        from io import BytesIO

        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            informacoes_preliminares='Orgao: SEDS',
            descricao_objeto='Objeto.',
            vinculacao_outro_dfd='Sem dependencia.',
            responsaveis='Responsavel centralizado.',
        )
        DfdItemTabela.objects.create(
            dfd=dfd,
            ordem=1,
            especificacao='Notebook',
            siafisico='456',
            quantidade='2.00',
            valor_unitario='100.50',
            valor_total='201.00',
        )

        response = self.client.get(reverse('licitacoes:dfd_export', args=[dfd.pk]))

        document = Document(BytesIO(response.content))
        paragraphs = '\n'.join(p.text for p in document.paragraphs)
        table_text = '\n'.join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
        self.assertIn('DFD - DFD', paragraphs)
        self.assertIn('1.1. Objeto.', paragraphs)
        self.assertIn(Dfd.OBJETO_NAO_LUXO_PADRAO, paragraphs)
        self.assertIn('4.1. Sem dependencia.', paragraphs)
        self.assertIn('Notebook', table_text)
        self.assertIn('SIAFISICO', table_text)
        self.assertIn('456', table_text)
        self.assertIn('Valor total', table_text)
        info_paragraph = next(p for p in document.paragraphs if p.text == 'Orgao: SEDS')
        object_paragraph = next(p for p in document.paragraphs if p.text == '1.1. Objeto.')
        responsaveis_paragraph = next(p for p in document.paragraphs if p.text == 'Responsavel centralizado.')
        self.assertEqual(info_paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)
        self.assertEqual(object_paragraph.alignment, WD_ALIGN_PARAGRAPH.JUSTIFY)
        self.assertEqual(responsaveis_paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
        self.assertTrue(all(run.font.name == 'Verdana' for run in object_paragraph.runs if run.text))
        self.assertTrue(all(run.font.size.pt == 10 for run in object_paragraph.runs if run.text))
        header_run = document.tables[0].rows[0].cells[0].paragraphs[0].runs[0]
        body_run = document.tables[0].rows[1].cells[1].paragraphs[0].runs[0]
        self.assertEqual(header_run.font.name, 'Verdana')
        self.assertEqual(header_run.font.size.pt, 8)
        self.assertEqual(body_run.font.name, 'Verdana')
        self.assertEqual(body_run.font.size.pt, 8)

    def test_preview_dfd_renderiza_marcacao_vermelha(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            descricao_objeto='Texto *vermelho*',
            objeto_nao_luxo='1.2. Item *destacado*',
        )
        DfdItemTabela.objects.create(
            dfd=dfd,
            ordem=1,
            especificacao='Notebook *especial*',
            quantidade='2.00',
            valor_unitario='100.50',
            valor_total='201.00',
        )

        response = self.client.get(reverse('licitacoes:dfd_preview', args=[dfd.pk]))

        self.assertContains(response, '<span class="text-danger">vermelho</span>')
        self.assertContains(response, '<span class="text-danger">destacado</span>')
        self.assertContains(response, '<span class="text-danger">especial</span>')

    def test_exporta_dfd_marcacao_vermelha_no_docx(self):
        from io import BytesIO

        from docx import Document

        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            descricao_objeto='Texto *vermelho*',
            objeto_nao_luxo='1.2. Item *destacado*',
        )
        DfdItemTabela.objects.create(
            dfd=dfd,
            ordem=1,
            especificacao='Notebook *especial*',
            quantidade='2.00',
            valor_unitario='100.50',
            valor_total='201.00',
        )

        response = self.client.get(reverse('licitacoes:dfd_export', args=[dfd.pk]))
        document = Document(BytesIO(response.content))
        runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
        runs += [run for table in document.tables for row in table.rows for cell in row.cells for paragraph in cell.paragraphs for run in paragraph.runs]
        red_texts = {run.text for run in runs if run.font.color.rgb and str(run.font.color.rgb) == 'FF0000'}

        self.assertIn('vermelho', red_texts)
        self.assertIn('destacado', red_texts)
        self.assertIn('especial', red_texts)

    def test_duplicar_dfd_copia_campos_e_tabela(self):
        dfd = Dfd.objects.create(
            nome='DFD',
            numero_processo='001/2026',
            status=Dfd.Status.CONCLUIDO,
            secao_atual=4,
            informacoes_preliminares='Orgao: SEDS',
            descricao_objeto='Objeto.',
            objeto_nao_luxo='1.2. Item nao luxo.',
            justificativa_necessidade='Necessidade.',
            estimativa_quantidade_valores='Estimativa.',
            vinculacao_outro_dfd='Sem vinculacao.',
            responsaveis='Responsavel.',
        )
        DfdItemTabela.objects.create(
            dfd=dfd,
            ordem=1,
            especificacao='Notebook',
            catmat='123',
            siafisico='456',
            unidade_medida='Unidade',
            quantidade='2.00',
            valor_unitario='100.50',
            valor_total='201.00',
        )

        duplicate = duplicate_dfd(dfd)

        self.assertEqual(duplicate.nome, 'Copia de DFD')
        self.assertEqual(duplicate.numero_processo, dfd.numero_processo)
        self.assertEqual(duplicate.status, dfd.status)
        self.assertEqual(duplicate.secao_atual, 4)
        self.assertEqual(duplicate.descricao_objeto, 'Objeto.')
        self.assertEqual(duplicate.itens_tabela.get().especificacao, 'Notebook')
        self.assertNotEqual(duplicate.itens_tabela.get().pk, dfd.itens_tabela.get().pk)

    def test_duplicar_dfd_pela_listagem_volta_para_edicao_do_duplicado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        dfd = Dfd.objects.create(nome='DFD', numero_processo='001/2026', secao_atual=3)
        DfdItemTabela.objects.create(
            dfd=dfd,
            ordem=1,
            especificacao='Notebook',
            quantidade='2.00',
            valor_unitario='100.50',
            valor_total='201.00',
        )

        response = self.client.post(reverse('licitacoes:dfd_duplicate', args=[dfd.pk]))

        duplicate = Dfd.objects.get(nome='Copia de DFD')
        self.assertEqual(duplicate.itens_tabela.get().especificacao, 'Notebook')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:dfd_edit', args=[duplicate.pk])}?secao=3",
            fetch_redirect_response=False,
        )


class RedMarkTests(TestCase):
    def test_renderiza_marcacao_vermelha_com_asteriscos_pareados_e_palavra_solteira(self):
        html = red_marked_html('Texto *vermelho* e *alerta')

        self.assertIn('<span class="text-danger">vermelho</span>', html)
        self.assertIn('<span class="text-danger">alerta</span>', html)
        self.assertNotIn('*vermelho*', html)

    def test_renderiza_marcacao_vermelha_em_palavra_solteira_e_bloco_com_espaco(self):
        html = red_marked_html('Texto *destaque\n\n* paragrafo inteiro*')

        self.assertIn('<span class="text-danger">destaque</span>', html)
        self.assertIn('<span class="text-danger"> paragrafo inteiro</span>', html)

    def test_renderiza_marcacao_vermelha_em_frase_pareada_com_um_ou_dois_asteriscos(self):
        html = red_marked_html('Texto *frase inteira marcada* e **outra frase marcada**')

        self.assertIn('<span class="text-danger">frase inteira marcada</span>', html)
        self.assertIn('<span class="text-danger">outra frase marcada</span>', html)


class TrTests(TestCase):
    def setUp(self):
        self.termo = TermoReferencia.objects.create(nome='TR', numero_processo='002/2026')
        self.sessao = SessaoTR.objects.create(termo=self.termo, titulo='Objeto', ordem=1)

    def test_numera_item_subitem_inciso_e_alinea(self):
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        sub = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem', ordem=1)
        inciso = ItemTR.objects.create(sessao=self.sessao, parent=sub, tipo=ItemTR.Tipo.INCISO, texto='Inciso', ordem=1)
        alinea = ItemTR.objects.create(sessao=self.sessao, parent=sub, tipo=ItemTR.Tipo.ALINEA, texto='Alinea', ordem=2)
        rows = build_item_rows(self.sessao)
        found = {row['item'].id: row for row in rows}
        self.assertEqual(found[item.id]['indice'], '1.1')
        self.assertEqual(found[sub.id]['indice'], '1.1.1')
        self.assertEqual(found[inciso.id]['enum_prefix'], 'I.')
        self.assertEqual(found[alinea.id]['enum_prefix'], 'a)')
        self.assertTrue(found[item.id]['pode_tabela_itens'])
        self.assertFalse(found[sub.id]['pode_tabela_itens'])

    def test_numera_inciso_criado_a_partir_de_inciso_como_irmao(self):
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        sub = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem', ordem=1)
        primeiro = ItemTR.objects.create(sessao=self.sessao, parent=sub, tipo=ItemTR.Tipo.INCISO, texto='Primeiro', ordem=1)
        parent = item_parent_for_tipo(primeiro, ItemTR.Tipo.INCISO)
        segundo = ItemTR.objects.create(sessao=self.sessao, parent=parent, tipo=ItemTR.Tipo.INCISO, texto='Segundo', ordem=2)
        rows = build_item_rows(self.sessao)
        found = {row['item'].id: row for row in rows}
        self.assertEqual(segundo.parent_id, sub.id)
        self.assertEqual(found[primeiro.id]['enum_prefix'], 'I.')
        self.assertEqual(found[segundo.id]['enum_prefix'], 'II.')

    def test_criar_tr_com_subsecao_e_inciso_por_marcadores(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')

        response = self.client.post(
            reverse('licitacoes:item_create', args=[self.sessao.pk]),
            {
                'texto': (
                    '@Garantia da contratação\n'
                    '#Será exigida a garantia.\n'
                    '**Caução em dinheiro.\n'
                    '$$Em moeda corrente.\n'
                    '**Fiança bancária.\n'
                    '@Sustentabilidade\n'
                    '#Além dos critérios, devem ser atendidos requisitos.'
                )
            },
        )

        rows = build_item_rows(self.sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertTrue(found['Garantia da contratação']['is_subsecao'])
        self.assertEqual(found['Garantia da contratação']['indice'], '')
        self.assertEqual(found['Será exigida a garantia.']['indice'], '1.1')
        self.assertEqual(found['Caução em dinheiro.']['enum_prefix'], 'I.')
        self.assertEqual(found['Em moeda corrente.']['enum_prefix'], 'a)')
        self.assertEqual(found['Fiança bancária.']['enum_prefix'], 'II.')
        self.assertEqual(found['Além dos critérios, devem ser atendidos requisitos.']['indice'], '1.2')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{found['Garantia da contratação']['item'].pk}",
            fetch_redirect_response=False,
        )

    def test_criar_tr_com_marcadores_destaca_vermelho_ignorando_subsecoes(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')

        self.client.post(
            reverse('licitacoes:item_create', args=[self.sessao.pk]),
            {
                'texto': (
                    '@Garantia da contratação\n'
                    '#Será exigida a garantia.\n'
                    '**Caução em dinheiro.\n'
                    '$$Em moeda corrente.'
                ),
                'modo_destaque_texto': 'vermelho_sem_subsecoes',
            },
        )

        textos = set(self.sessao.itens.values_list('texto', flat=True))
        self.assertIn('Garantia da contratação', textos)
        self.assertIn('*Será exigida a garantia.*', textos)
        self.assertIn('*Caução em dinheiro.*', textos)
        self.assertIn('*Em moeda corrente.*', textos)

    def test_criar_alineas_tr_com_marcador_isolado_no_subitem(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        subitem = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem', ordem=1)

        self.client.post(
            reverse('licitacoes:item_child_create', args=[self.sessao.pk, subitem.pk]),
            {'texto': '$$o prazo de validade;\n$$a data da emissão;'},
        )

        rows = build_item_rows(self.sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertEqual(found['o prazo de validade;']['enum_prefix'], 'a)')
        self.assertEqual(found['a data da emissão;']['enum_prefix'], 'b)')
        self.assertEqual(found['o prazo de validade;']['item'].parent_id, subitem.id)

    def test_editar_tr_com_marcadores_substitui_item_e_renumera_irmaos(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        raiz = ItemTR.objects.create(sessao=self.sessao, texto='Raiz', ordem=1)
        alvo = ItemTR.objects.create(sessao=self.sessao, parent=raiz, texto='Alvo', ordem=1)
        filho_antigo = ItemTR.objects.create(sessao=self.sessao, parent=alvo, texto='Filho antigo', ordem=1)
        posterior = ItemTR.objects.create(sessao=self.sessao, parent=raiz, texto='Posterior', ordem=2)

        response = self.client.post(
            reverse('licitacoes:item_update', args=[self.sessao.pk, alvo.pk]),
            {'texto': '#Novo\n##Detalhe\n#Outro'},
        )

        alvo.refresh_from_db()
        rows = build_item_rows(self.sessao)
        found = {row['item'].texto: row for row in rows}
        self.assertEqual(alvo.texto, 'Novo')
        self.assertEqual(found['Novo']['indice'], '1.1.1')
        self.assertEqual(found['Detalhe']['indice'], '1.1.1.1')
        self.assertEqual(found['Outro']['indice'], '1.1.2')
        self.assertEqual(found['Posterior']['indice'], '1.1.3')
        self.assertFalse(ItemTR.objects.filter(pk=filho_antigo.pk).exists())
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{alvo.pk}",
            fetch_redirect_response=False,
        )

    def test_editar_tr_destaca_item_e_filhos_em_vermelho(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        filho = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Filho', ordem=1)
        neto = ItemTR.objects.create(sessao=self.sessao, parent=filho, texto='Neto', ordem=1)

        self.client.post(
            reverse('licitacoes:item_update', args=[self.sessao.pk, item.pk]),
            {
                'texto': 'Item alterado',
                'modo_destaque_texto': 'todo_vermelho_com_filhos',
            },
        )

        item.refresh_from_db()
        filho.refresh_from_db()
        neto.refresh_from_db()
        self.assertEqual(item.texto, '*Item alterado*')
        self.assertEqual(filho.texto, '*Filho*')
        self.assertEqual(neto.texto, '*Neto*')

    def test_editar_sessao_tr_destaca_filhos_em_vermelho(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        filho = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Filho', ordem=1)

        self.client.post(
            reverse('licitacoes:sessao_update', args=[self.termo.pk, self.sessao.pk]),
            {'titulo': 'Objeto atualizado', 'filhos_em_vermelho': '1'},
        )

        item.refresh_from_db()
        filho.refresh_from_db()
        self.assertEqual(item.texto, '*Item*')
        self.assertEqual(filho.texto, '*Filho*')

    def test_mover_bloqueia_descendente_e_renumera(self):
        a = ItemTR.objects.create(sessao=self.sessao, texto='A', ordem=1)
        b = ItemTR.objects.create(sessao=self.sessao, texto='B', ordem=2)
        child = ItemTR.objects.create(sessao=self.sessao, parent=a, texto='Filho', ordem=1)
        with self.assertRaises(ValueError):
            move_item(a, child, 'child')
        move_item(b, a, 'child')
        b.refresh_from_db()
        self.assertEqual(b.parent_id, a.id)
        self.assertEqual(b.ordem, 2)

    def test_mover_como_primeiro_subitem_empurra_filhos_existentes(self):
        destino = ItemTR.objects.create(sessao=self.sessao, texto='Destino', ordem=1)
        primeiro = ItemTR.objects.create(sessao=self.sessao, parent=destino, texto='Primeiro filho', ordem=1)
        segundo = ItemTR.objects.create(sessao=self.sessao, parent=destino, texto='Segundo filho', ordem=2)
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item movido', ordem=2)

        move_item(item, destino, 'child', child_position=1)

        item.refresh_from_db()
        primeiro.refresh_from_db()
        segundo.refresh_from_db()
        self.assertEqual(item.parent_id, destino.id)
        self.assertEqual(item.ordem, 1)
        self.assertEqual(primeiro.ordem, 2)
        self.assertEqual(segundo.ordem, 3)

    def test_duplicar_item_copia_subitens(self):
        a = ItemTR.objects.create(sessao=self.sessao, texto='A', ordem=1)
        b = ItemTR.objects.create(sessao=self.sessao, texto='B', ordem=2)
        child = ItemTR.objects.create(sessao=self.sessao, parent=a, texto='Filho', ordem=1)
        ItemTR.objects.create(sessao=self.sessao, parent=child, texto='Neto', ordem=1)
        TabelaItemLinha.objects.create(item=a, ordem=1, descricao='Notebook', quantidade='2.00')

        duplicate = duplicate_item(a, b, 'after')
        duplicate.refresh_from_db()

        self.assertEqual(duplicate.texto, 'A')
        self.assertEqual(duplicate.parent_id, None)
        self.assertEqual(duplicate.ordem, 3)
        duplicate_child = duplicate.filhos.get()
        self.assertEqual(duplicate_child.texto, 'Filho')
        self.assertEqual(duplicate_child.filhos.get().texto, 'Neto')
        self.assertEqual(duplicate.tabela_linhas.get().descricao, 'Notebook')

    def test_limpar_filhos_tr_remove_subitens_e_preserva_item(self):
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        child = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem', ordem=1)
        grandchild = ItemTR.objects.create(sessao=self.sessao, parent=child, texto='Neto', ordem=1)

        removed = clear_item_children(item)

        self.assertEqual(removed, 2)
        self.assertTrue(ItemTR.objects.filter(pk=item.pk).exists())
        self.assertFalse(ItemTR.objects.filter(pk__in=[child.pk, grandchild.pk]).exists())

    def test_duplicar_tr_copia_sessoes_itens_e_tabelas(self):
        sessao_2 = SessaoTR.objects.create(termo=self.termo, titulo='Outra sessao', ordem=2)
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item original', ordem=1)
        child = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem original', ordem=1)
        TabelaItemLinha.objects.create(item=item, ordem=1, descricao='Notebook', quantidade='2.00')
        ItemTR.objects.create(sessao=sessao_2, texto='Item sessao 2', ordem=1)

        duplicate = duplicate_termo(self.termo)

        self.assertEqual(duplicate.nome, 'Copia de TR')
        self.assertEqual(duplicate.numero_processo, self.termo.numero_processo)
        self.assertEqual(duplicate.sessoes.count(), 2)
        new_sessao = duplicate.sessoes.get(ordem=1)
        new_item = new_sessao.itens.get(parent=None, texto='Item original')
        self.assertNotEqual(new_item.pk, item.pk)
        self.assertEqual(new_item.filhos.get().texto, child.texto)
        self.assertEqual(new_item.tabela_linhas.get().descricao, 'Notebook')

    def test_limpar_filhos_da_sessao_tr_preserva_sessao_e_remove_itens(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item', ordem=1)
        child = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem', ordem=1)

        response = self.client.post(reverse('licitacoes:sessao_clear_items', args=[self.termo.pk, self.sessao.pk]))

        self.assertTrue(SessaoTR.objects.filter(pk=self.sessao.pk).exists())
        self.assertFalse(ItemTR.objects.filter(pk__in=[item.pk, child.pk]).exists())
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#sessao-{self.sessao.pk}",
            fetch_redirect_response=False,
        )

    def test_tabela_item_create_disponivel_somente_no_item_1_1(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item_1_1 = ItemTR.objects.create(sessao=self.sessao, texto='Item 1.1', ordem=1)
        item_1_2 = ItemTR.objects.create(sessao=self.sessao, texto='Item 1.2', ordem=2)

        response = self.client.post(
            reverse('licitacoes:tabela_item_create', args=[self.sessao.pk, item_1_1.pk]),
            {
                'descricao': 'Notebook',
                'catmat_catser': '123',
                'siafisico': '456',
                'unidade_fornecimento': 'Unidade',
                'quantidade': '2.00',
            },
        )

        linha = item_1_1.tabela_linhas.get()
        self.assertEqual(linha.ordem, 1)
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#tabela-itens-{item_1_1.pk}",
            fetch_redirect_response=False,
        )

        bloqueado = self.client.get(reverse('licitacoes:tabela_item_create', args=[self.sessao.pk, item_1_2.pk]))
        self.assertEqual(bloqueado.status_code, 404)

    def test_detail_renderiza_tabela_no_item_1_1(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item 1.1', ordem=1)
        TabelaItemLinha.objects.create(item=item, ordem=1, descricao='Notebook', quantidade='2.00')

        response = self.client.get(reverse('licitacoes:tr_detail', args=[self.termo.pk]))

        self.assertContains(response, 'Itens da tabela do item 1.1')
        self.assertContains(response, 'Notebook')

    def test_detail_renderiza_marcacao_vermelha_no_item(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        ItemTR.objects.create(sessao=self.sessao, texto='Texto *vermelho*', ordem=1)

        response = self.client.get(reverse('licitacoes:tr_detail', args=[self.termo.pk]))

        self.assertContains(response, '<span class="text-danger">vermelho</span>')

    def test_exporta_marcacao_vermelha_no_docx(self):
        from io import BytesIO

        from docx import Document

        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        ItemTR.objects.create(sessao=self.sessao, texto='Texto *vermelho*', ordem=1)

        response = self.client.get(reverse('licitacoes:tr_export', args=[self.termo.pk]))
        document = Document(BytesIO(response.content))
        runs = [run for paragraph in document.paragraphs for run in paragraph.runs]
        red_run = next(run for run in runs if run.text == 'vermelho')

        self.assertEqual(str(red_run.font.color.rgb), 'FF0000')

    def test_exporta_tr_com_subsecao_e_inciso_docx(self):
        from io import BytesIO

        from docx import Document

        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        subsecao = ItemTR.objects.create(sessao=self.sessao, tipo=ItemTR.Tipo.SUBSECAO, texto='Garantia da contratação', ordem=1)
        item = ItemTR.objects.create(sessao=self.sessao, texto='Será exigida a garantia.', ordem=2)
        inciso = ItemTR.objects.create(sessao=self.sessao, parent=item, tipo=ItemTR.Tipo.INCISO, texto='Caução em dinheiro.', ordem=1)
        ItemTR.objects.create(sessao=self.sessao, parent=inciso, tipo=ItemTR.Tipo.ALINEA, texto='Em moeda corrente.', ordem=1)

        response = self.client.get(reverse('licitacoes:tr_export', args=[self.termo.pk]))

        document = Document(BytesIO(response.content))
        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertIn(subsecao.texto, texts)
        self.assertIn('1.1. Será exigida a garantia.', texts)
        self.assertIn('I. Caução em dinheiro.', texts)
        self.assertIn('a) Em moeda corrente.', texts)

    def test_duplicar_item_pela_view_volta_para_item_duplicado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item original', ordem=1)

        response = self.client.post(
            reverse('licitacoes:item_duplicate', args=[self.sessao.pk, item.pk]),
            {'target': f'sessao:{self.sessao.pk}', 'action': 'child'},
        )

        duplicate = ItemTR.objects.exclude(pk=item.pk).get(texto='Item original')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{duplicate.pk}",
            fetch_redirect_response=False,
        )

    def test_duplicar_tr_pela_listagem_volta_para_tr_duplicado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        ItemTR.objects.create(sessao=self.sessao, texto='Item original', ordem=1)

        response = self.client.post(reverse('licitacoes:tr_duplicate', args=[self.termo.pk]))

        duplicate = TermoReferencia.objects.get(nome='Copia de TR')
        self.assertEqual(duplicate.sessoes.get().itens.get().texto, 'Item original')
        self.assertRedirects(
            response,
            reverse('licitacoes:tr_detail', args=[duplicate.pk]),
            fetch_redirect_response=False,
        )

    def test_editar_item_volta_para_item_editado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Texto antigo', ordem=1)

        response = self.client.post(
            reverse('licitacoes:item_update', args=[self.sessao.pk, item.pk]),
            {'texto': 'Texto atualizado'},
        )

        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{item.pk}",
            fetch_redirect_response=False,
        )

    def test_criar_item_volta_para_item_criado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')

        response = self.client.post(
            reverse('licitacoes:item_create', args=[self.sessao.pk]),
            {'texto': 'Item novo'},
        )

        item = ItemTR.objects.get(texto='Item novo')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{item.pk}",
            fetch_redirect_response=False,
        )

    def test_criar_subitem_volta_para_subitem_criado(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        parent = ItemTR.objects.create(sessao=self.sessao, texto='Item pai', ordem=1)

        response = self.client.post(
            reverse('licitacoes:item_child_create', args=[self.sessao.pk, parent.pk]),
            {'texto': 'Subitem novo'},
        )

        item = ItemTR.objects.get(texto='Subitem novo')
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{item.pk}",
            fetch_redirect_response=False,
        )

    def test_limpar_filhos_tr_pela_view_volta_para_item(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='Item pai', ordem=1)
        child = ItemTR.objects.create(sessao=self.sessao, parent=item, texto='Subitem', ordem=1)

        response = self.client.post(reverse('licitacoes:item_clear_children', args=[self.sessao.pk, item.pk]))

        self.assertFalse(ItemTR.objects.filter(pk=child.pk).exists())
        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{item.pk}",
            fetch_redirect_response=False,
        )

    def test_excluir_item_volta_para_proximo_item(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='A', ordem=1)
        next_item = ItemTR.objects.create(sessao=self.sessao, texto='B', ordem=2)

        response = self.client.post(reverse('licitacoes:item_delete', args=[self.sessao.pk, item.pk]))

        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{next_item.pk}",
            fetch_redirect_response=False,
        )

    def test_excluir_ultimo_item_volta_para_item_anterior(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        previous_item = ItemTR.objects.create(sessao=self.sessao, texto='A', ordem=1)
        item = ItemTR.objects.create(sessao=self.sessao, texto='B', ordem=2)

        response = self.client.post(reverse('licitacoes:item_delete', args=[self.sessao.pk, item.pk]))

        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#item-{previous_item.pk}",
            fetch_redirect_response=False,
        )

    def test_excluir_unico_item_raiz_volta_para_sessao(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        item = ItemTR.objects.create(sessao=self.sessao, texto='A', ordem=1)

        response = self.client.post(reverse('licitacoes:item_delete', args=[self.sessao.pk, item.pk]))

        self.assertRedirects(
            response,
            f"{reverse('licitacoes:tr_detail', args=[self.termo.pk])}#sessao-{self.sessao.pk}",
            fetch_redirect_response=False,
        )


class PesquisaPrecoTests(TestCase):
    def setUp(self):
        User = get_user_model()
        User.objects.create_superuser(username='admin', password='123')
        self.client.login(username='admin', password='123')
        self.termo = TermoReferencia.objects.create(nome='TR Pesquisa', numero_processo='003/2026')
        self.sessao = SessaoTR.objects.create(termo=self.termo, titulo='Objeto', ordem=1)
        self.item = ItemTR.objects.create(sessao=self.sessao, texto='Item 1.1', ordem=1)
        self.linha_1 = TabelaItemLinha.objects.create(
            item=self.item,
            ordem=1,
            descricao='Notebook',
            catmat_catser='123',
            siafisico='456',
            unidade_fornecimento='Unidade',
            quantidade='2.00',
        )
        self.linha_2 = TabelaItemLinha.objects.create(
            item=self.item,
            ordem=2,
            descricao='Monitor',
            unidade_fornecimento='Unidade',
            quantidade='3.00',
        )

    def fornecedor(self, nome='Fornecedor A', cnpj='00.000.000/0001-00'):
        return Fornecedor.objects.create(
            razao_social=nome,
            cnpj=cnpj,
            telefone='(11) 1111-1111',
            contato='Contato',
            email_contato='contato@example.com',
        )

    def test_formulario_fornecedor_aceita_multiplos_emails_por_ponto_e_virgula(self):
        form = FornecedorForm(data={
            'razao_social': 'Powertec Informática',
            'cnpj': '00.000.000/0001-02',
            'telefone': '(11) 2222-2222',
            'contato': 'Contato',
            'email_contato': 'powertecinformatica2@gmail.com; powertec@danro.com.br',
        })

        self.assertTrue(form.is_valid(), form.errors)
        self.assertEqual(
            form.cleaned_data['email_contato'],
            'powertecinformatica2@gmail.com; powertec@danro.com.br',
        )

    def test_formulario_fornecedor_valida_cada_email_da_lista(self):
        form = FornecedorForm(data={
            'razao_social': 'Powertec Informática',
            'cnpj': '00.000.000/0001-02',
            'telefone': '(11) 2222-2222',
            'contato': 'Contato',
            'email_contato': 'powertecinformatica2@gmail.com; email-invalido',
        })

        self.assertFalse(form.is_valid())
        self.assertIn('email_contato', form.errors)

    def test_detail_tr_exibe_botao_pesquisa_preco(self):
        response = self.client.get(reverse('licitacoes:tr_detail', args=[self.termo.pk]))

        self.assertContains(response, 'Pesquisa de Preço')
        self.assertContains(response, reverse('licitacoes:pesquisa_preco_open', args=[self.termo.pk]))

    def test_cria_pesquisa_aquisicao_e_servico(self):
        response = self.client.post(
            reverse('licitacoes:pesquisa_preco_create', args=[self.termo.pk]),
            {
                'pesquisador_nome': 'José Eduardo',
                'pesquisador_email': 'jose@example.com',
                'pesquisador_cargo': 'Analista',
                'tipo': PesquisaPreco.Tipo.AQUISICAO,
                'vigencia_meses': '',
            },
        )

        pesquisa = self.termo.pesquisa_preco
        self.assertEqual(pesquisa.tipo, PesquisaPreco.Tipo.AQUISICAO)
        self.assertIsNone(pesquisa.vigencia_meses)
        self.assertEqual(pesquisa.pesquisador_nome, 'José Eduardo')
        self.assertRedirects(response, reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]), fetch_redirect_response=False)

        termo_servico = TermoReferencia.objects.create(nome='TR Serviço', numero_processo='004/2026')
        response = self.client.post(
            reverse('licitacoes:pesquisa_preco_create', args=[termo_servico.pk]),
            {
                'pesquisador_nome': 'Maria',
                'pesquisador_email': 'maria@example.com',
                'pesquisador_cargo': 'Assessora',
                'tipo': PesquisaPreco.Tipo.SERVICO,
                'vigencia_meses': '24',
            },
        )

        self.assertEqual(termo_servico.pesquisa_preco.vigencia_meses, 24)
        self.assertRedirects(response, reverse('licitacoes:pesquisa_preco_detail', args=[termo_servico.pk]), fetch_redirect_response=False)

    def test_exclui_pesquisa_para_comecar_outra(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)

        response = self.client.post(reverse('licitacoes:pesquisa_preco_delete', args=[self.termo.pk]))

        self.assertFalse(PesquisaPreco.objects.filter(pk=pesquisa.pk).exists())
        self.assertRedirects(response, reverse('licitacoes:tr_detail', args=[self.termo.pk]), fetch_redirect_response=False)

    def test_adiciona_fornecedor_atualiza_contato_e_calcula_dias_sem_resposta(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()

        response = self.client.post(
            reverse('licitacoes:pesquisa_preco_fornecedor_add', args=[self.termo.pk]),
            {'fornecedor': fornecedor.pk},
        )

        pesquisa_fornecedor = PesquisaPrecoFornecedor.objects.get(pesquisa=pesquisa, fornecedor=fornecedor)
        self.assertRedirects(response, reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]), fetch_redirect_response=False)

        response = self.client.post(reverse('licitacoes:pesquisa_preco_atualizar_contato', args=[self.termo.pk, pesquisa_fornecedor.pk]))

        contato = PesquisaPrecoContato.objects.get(pesquisa_fornecedor=pesquisa_fornecedor)
        self.assertEqual(contato.data_contato, timezone.localdate())
        self.assertRedirects(response, reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]), fetch_redirect_response=False)
        context = pesquisa_preco_context(pesquisa)
        self.assertEqual(context['fornecedores'][0]['dias_sem_resposta'], 0)
        self.assertEqual(context['fornecedores'][0]['row_class'], 'spi-pesquisa-row-verde')

    def test_remove_fornecedor_apenas_da_pesquisa(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pesquisa_fornecedor = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)

        response = self.client.post(reverse('licitacoes:pesquisa_preco_fornecedor_remove', args=[self.termo.pk, pesquisa_fornecedor.pk]))

        self.assertFalse(PesquisaPrecoFornecedor.objects.filter(pk=pesquisa_fornecedor.pk).exists())
        self.assertTrue(Fornecedor.objects.filter(pk=fornecedor.pk).exists())
        self.assertRedirects(response, reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]), fetch_redirect_response=False)

    def test_crud_global_exclui_fornecedor_do_sistema(self):
        fornecedor = self.fornecedor()
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)

        response = self.client.post(reverse('licitacoes:fornecedor_delete', args=[fornecedor.pk]))

        self.assertFalse(Fornecedor.objects.filter(pk=fornecedor.pk).exists())
        self.assertFalse(PesquisaPrecoFornecedor.objects.filter(fornecedor_id=fornecedor.pk).exists())
        self.assertRedirects(response, reverse('licitacoes:fornecedor_list'), fetch_redirect_response=False)

    def test_orcamento_calcula_totais_medias_e_para_contador(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor_a = self.fornecedor()
        fornecedor_b = self.fornecedor('Fornecedor B', '00.000.000/0001-01')
        pf_a = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor_a)
        pf_b = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor_b)
        PesquisaPrecoContato.objects.create(pesquisa_fornecedor=pf_a, data_contato=timezone.localdate())
        PesquisaPrecoContato.objects.create(pesquisa_fornecedor=pf_b, data_contato=timezone.localdate())

        self.client.post(
            reverse('licitacoes:pesquisa_preco_orcamento', args=[self.termo.pk, pf_a.pk]),
            {
                'data_resposta': timezone.localdate().isoformat(),
                'validade_orcamento_dias': '30',
                'documento_fornecedor': SimpleUploadedFile('orcamento-a.pdf', b'PDF A', content_type='application/pdf'),
                f'preco_item_{self.linha_1.pk}': '10.00',
                f'preco_item_{self.linha_2.pk}': '20.00',
            },
        )
        self.client.post(
            reverse('licitacoes:pesquisa_preco_orcamento', args=[self.termo.pk, pf_b.pk]),
            {
                'data_resposta': timezone.localdate().isoformat(),
                'validade_orcamento_dias': '8',
                'documento_fornecedor': SimpleUploadedFile('orcamento-b.pdf', b'PDF B', content_type='application/pdf'),
                f'preco_item_{self.linha_1.pk}': '20.00',
                f'preco_item_{self.linha_2.pk}': '40.00',
            },
        )

        self.assertEqual(PesquisaPrecoItemValor.objects.count(), 4)
        context = pesquisa_preco_context(pesquisa)
        self.assertEqual(context['medias'][0]['preco_medio'], 15)
        self.assertEqual(context['medias'][0]['valor_total_medio'], 30)
        self.assertEqual(context['total_medio'], 120)
        self.assertIsNone(context['fornecedores'][0]['dias_sem_resposta'])
        self.assertTrue(context['fornecedores'][1]['validade_alerta'])

    def test_painel_pesquisa_renderiza_fornecedor_e_quadro(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)
        PesquisaPrecoContato.objects.create(pesquisa_fornecedor=pf, data_contato=timezone.localdate())
        PesquisaPrecoItemValor.objects.create(pesquisa_fornecedor=pf, item=self.linha_1, preco_unitario='10.00')
        pf.data_resposta = timezone.localdate()
        pf.validade_orcamento_dias = 30
        pf.save()

        response = self.client.get(reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]))

        self.assertContains(response, 'Fornecedor A')
        self.assertContains(response, 'Quadro comparativo')
        self.assertContains(response, 'Notebook')

    def test_painel_pesquisa_renderiza_multiplos_emails_com_copia_individual(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        fornecedor.email_contato = 'powertecinformatica2@gmail.com; powertec@danro.com.br'
        fornecedor.save()
        PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)

        response = self.client.get(reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]))

        self.assertContains(response, 'data-email="powertecinformatica2@gmail.com"')
        self.assertContains(response, 'data-email="powertec@danro.com.br"')

    def test_formulario_orcamento_renderiza_itens(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)

        response = self.client.get(reverse('licitacoes:pesquisa_preco_orcamento', args=[self.termo.pk, pf.pk]))

        self.assertContains(response, 'Orçamento')
        self.assertContains(response, 'Documento do fornecedor')
        self.assertContains(response, f'name="preco_item_{self.linha_1.pk}"')
        self.assertContains(response, 'Notebook')

    def test_orcamento_salva_documento_do_fornecedor(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)
        arquivo = SimpleUploadedFile('orcamento.pdf', b'PDF fake', content_type='application/pdf')

        self.client.post(
            reverse('licitacoes:pesquisa_preco_orcamento', args=[self.termo.pk, pf.pk]),
            {
                'data_resposta': timezone.localdate().isoformat(),
                'validade_orcamento_dias': '30',
                'documento_fornecedor': arquivo,
                f'preco_item_{self.linha_1.pk}': '10.00',
                f'preco_item_{self.linha_2.pk}': '20.00',
            },
        )

        pf.refresh_from_db()
        self.assertTrue(pf.documento_fornecedor.name.endswith('.pdf'))

    def test_orcamento_exige_documento_do_fornecedor(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)

        response = self.client.post(
            reverse('licitacoes:pesquisa_preco_orcamento', args=[self.termo.pk, pf.pk]),
            {
                'data_resposta': timezone.localdate().isoformat(),
                'validade_orcamento_dias': '30',
                f'preco_item_{self.linha_1.pk}': '10.00',
                f'preco_item_{self.linha_2.pk}': '20.00',
            },
        )

        self.assertContains(response, 'Anexe o documento do fornecedor para salvar o orçamento.')
        self.assertFalse(PesquisaPrecoItemValor.objects.filter(pesquisa_fornecedor=pf).exists())

    def test_botao_orcamento_baixa_anexo_quando_orcamento_salvo(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)
        pf.data_resposta = timezone.localdate()
        pf.validade_orcamento_dias = 30
        pf.documento_fornecedor = 'licitacoes/pesquisa_preco/orcamentos/orcamento.pdf'
        pf.save()
        PesquisaPrecoItemValor.objects.create(pesquisa_fornecedor=pf, item=self.linha_1, preco_unitario='10.00')

        response = self.client.get(reverse('licitacoes:pesquisa_preco_detail', args=[self.termo.pk]))

        self.assertContains(response, 'href="/media/licitacoes/pesquisa_preco/orcamentos/orcamento.pdf"')
        self.assertContains(response, 'download')
        self.assertNotContains(response, reverse('licitacoes:pesquisa_preco_orcamento', args=[self.termo.pk, pf.pk]))

    def test_servico_calcula_total_para_contratacao(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.SERVICO, vigencia_meses=12)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)
        PesquisaPrecoItemValor.objects.create(pesquisa_fornecedor=pf, item=self.linha_1, preco_unitario='10.00')
        PesquisaPrecoItemValor.objects.create(pesquisa_fornecedor=pf, item=self.linha_2, preco_unitario='20.00')
        pf.data_resposta = timezone.localdate()
        pf.validade_orcamento_dias = 30
        pf.save()

        context = pesquisa_preco_context(pesquisa)

        self.assertEqual(context['fornecedores'][0]['total'], 80)
        self.assertEqual(context['fornecedores'][0]['total_contratacao'], 960)
        self.assertEqual(context['total_medio_contratacao'], 960)

    def test_exporta_xlsx(self):
        pesquisa = PesquisaPreco.objects.create(termo=self.termo, tipo=PesquisaPreco.Tipo.AQUISICAO)
        fornecedor = self.fornecedor()
        pf = PesquisaPrecoFornecedor.objects.create(pesquisa=pesquisa, fornecedor=fornecedor)
        PesquisaPrecoItemValor.objects.create(pesquisa_fornecedor=pf, item=self.linha_1, preco_unitario='10.00')
        pf.data_resposta = timezone.localdate()
        pf.validade_orcamento_dias = 30
        pf.save()

        response = self.client.get(reverse('licitacoes:pesquisa_preco_export', args=[self.termo.pk]))

        self.assertEqual(response.status_code, 200)
        self.assertEqual(response['Content-Type'], 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        self.assertIn('pesquisa_preco_', response['Content-Disposition'])
        workbook = load_workbook(filename=BytesIO(response.content), data_only=False)
        worksheet = workbook['PCs - Tab Alternativa']
        self.assertEqual(workbook.sheetnames, ['PCs - Tab Alternativa'])
        self.assertIn('B2:Q2', [str(item) for item in worksheet.merged_cells.ranges])
        self.assertEqual(worksheet['B2'].value, 'QUADRO COMPARATIVO DE PESQUISA DE PREÇOS')
        self.assertEqual(worksheet['J4'].value, fornecedor.razao_social)
