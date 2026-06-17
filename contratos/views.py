from django.contrib.auth.decorators import login_required
# Criado por José Eduardo Santana Martins e OpenAI Codex em 08/06/2026
# Objetivo: Entregar o CRUD e o fluxo operacional do Contratos V2 com checklist, competências, avaliação e pagamento.

import os
import mimetypes
import subprocess
import sys
import tempfile
import threading
import zipfile
from datetime import timedelta
from decimal import Decimal
from io import BytesIO
from pathlib import Path

from django.contrib import messages
from django.contrib.auth.mixins import LoginRequiredMixin
from django.core.files.base import ContentFile
from django.core.exceptions import ValidationError
from django.db import close_old_connections
from django.db.models import Max, Prefetch, Q, Case, When, Value, IntegerField
from django.shortcuts import get_object_or_404, redirect, render
from django.http import FileResponse, Http404, JsonResponse
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from pypdf import PdfReader, PdfWriter

def is_modal_request(request):
    """Indica quando o CRUD deve responder em formato de modal AJAX."""
    return request.headers.get('X-Requested-With') == 'XMLHttpRequest' and request.GET.get('modal') == '1'

class BlockIfCompetenciasGeradasMixin:
    """Bloqueia a edição/exclusão de itens estruturais se as competências já foram geradas."""
    def dispatch(self, request, *args, **kwargs):
        obj = None
        if hasattr(self, 'get_object'):
            try:
                # Some views might need pk, but kwargs has it
                if 'pk' in kwargs or 'slug' in kwargs:
                    obj = self.get_object()
            except:
                pass
        
        contrato = None
        if obj:
            if hasattr(obj, 'competencias'):  # it's a Contrato
                contrato = obj
            elif hasattr(obj, 'contrato'):
                contrato = obj.contrato
            elif hasattr(obj, 'modelo_checklist'):
                contrato = obj.modelo_checklist.contrato
            elif hasattr(obj, 'formulario'):
                contrato = obj.formulario.contrato
            elif hasattr(obj, 'grupo'):
                contrato = obj.grupo.formulario.contrato
        
        if not contrato:
            from .models import Contrato, ChecklistModelo, FormularioAvaliacao, GrupoAvaliacao
            if 'contrato_pk' in kwargs:
                contrato = Contrato.objects.filter(pk=kwargs['contrato_pk']).first()
            elif 'modelo_pk' in kwargs:
                modelo = ChecklistModelo.objects.filter(pk=kwargs['modelo_pk']).first()
                if modelo: contrato = modelo.contrato
            elif 'formulario_pk' in kwargs:
                form = FormularioAvaliacao.objects.filter(pk=kwargs['formulario_pk']).first()
                if form: contrato = form.contrato
            elif 'grupo_pk' in kwargs:
                grupo = GrupoAvaliacao.objects.filter(pk=kwargs['grupo_pk']).first()
                if grupo: contrato = grupo.formulario.contrato
        
        if contrato and contrato.competencias.exists():
            messages.error(request, 'Ação bloqueada: Este contrato já possui competências geradas. Nenhuma alteração estrutural é permitida.')
            return redirect('contratos:contrato_detail', pk=contrato.pk)
            
        return super().dispatch(request, *args, **kwargs)

from django.views.generic import CreateView, DeleteView, DetailView, FormView, ListView, UpdateView, View

from acls.mixins import ACLRequiredMixin

from .forms import (
    AvaliacaoCompetenciaV2Form,
    ChecklistModeloItemForm,
    ChecklistModeloForm,
    ChecklistPadraoGlobalForm,
    ChecklistPadraoGlobalItemForm,
    CompetenciaChecklistUploadForm,
    CompetenciaChecklistExtraItemForm,
    CompetenciaMedicaoLoteV2Form,
    CompetenciaOBExecucaoForm,
    CompetenciaPagamentoExecucaoV2Form,
    ContratoItemForm,
    ContratoForm,
    DocumentoImportanteContratoForm,
    EscalaNotaAvaliacaoForm,
    FaixaLiberacaoAvaliacaoForm,
    FormularioAvaliacaoForm,
    GrupoAvaliacaoForm,
    ItemAvaliacaoForm,
    EmpresaContratadaForm,
    ResponsavelEmpresaForm,
    PrazoMonitoramentoForm,
)
from .models import (
    AvaliacaoQualidadeCompetencia,
    ChecklistCompetenciaAnexo,
    ChecklistCompetenciaItem,
    ChecklistModeloItem,
    ChecklistModelo,
    ChecklistPadraoGlobal,
    ChecklistPadraoGlobalItem,
    CompetenciaPagamento,
    ContratoItem,
    Contrato,
    DocumentoImportanteContrato,
    EscalaNotaAvaliacao,
    FaixaLiberacaoAvaliacao,
    FormularioAvaliacao,
    GrupoAvaliacao,
    ItemAvaliacao,
    MedicaoItemCompetencia,
    EmpresaContratada,
    ExportacaoDocumentosCompetencia,
    ResponsavelEmpresa,
    PrazoMonitoramento,
)
from .services import (
    avaliacao_v2_esta_concluida,
    competencia_checklist_v2_esta_concluido,
    competencia_medicao_v2_esta_concluida,
    criar_avaliacao_shell_competencia_v2,
    inclusive_end_date,
    recalcular_avaliacao_v2,
    recalcular_competencia_v2,
    usuario_eh_admin_sistema,
    usuario_pode_gerir_documento_importante,
    usuario_pode_gerir_documento_importante_contrato,
    usuario_pode_preencher_avaliacao_fiscal_v2,
    usuario_pode_preencher_avaliacao_gestor_v2,
    usuario_pode_gerir_contrato_v2,
)


class ContratosAccessMixin(LoginRequiredMixin, ACLRequiredMixin):
    """Protege o módulo por login e ACL própria da versão nova."""

    recurso_slug = 'contratos'


class ContratosWriteMixin(ContratosAccessMixin):
    """Exige escrita nas operações de criação, edição e exclusão."""

    acl_nivel_minimo = 'MODIFICACAO'


def assign_owner(instance, request):
    """Mantém autoria e atualização alinhadas às operações do CRUD."""

    if hasattr(instance, 'criado_por_id') and not instance.criado_por_id:
        instance.criado_por = request.user
    if hasattr(instance, 'atualizado_por_id'):
        instance.atualizado_por = request.user
    return instance


def serializar_responsavel_v2(label, usuario):
    """Entrega os dados do responsável no formato esperado pelo modal de ramais."""

    if not usuario:
        return {
            'label': label,
            'nome': '-',
            'iniciais': '-',
            'cargo': '-',
            'setor': '-',
            'email': '-',
            'ramal': '-',
            'celular': '',
            'whatsapp': '',
            'local': '-',
            'foto_url': '',
        }

    perfil = getattr(usuario, 'perfil', None)
    foto = getattr(perfil, 'foto', None)
    foto_url = ''
    if foto and getattr(foto, 'name', ''):
        try:
            foto_url = foto.url
        except ValueError:
            foto_url = ''

    nome = getattr(perfil, 'nome_completo', None) or usuario.get_full_name() or usuario.username
    return {
        'label': label,
        'nome': nome,
        'iniciais': ''.join(parte[0].upper() for parte in nome.split()[:2] if parte) or nome[:1].upper(),
        'cargo': getattr(perfil, 'cargo', '') or '-',
        'setor': getattr(perfil, 'setor', '') or '-',
        'email': usuario.email or '-',
        'ramal': getattr(perfil, 'ramal', '') or '-',
        'celular': getattr(perfil, 'celular', '') or '',
        'whatsapp': perfil.whatsapp_url if perfil else '',
        'local': perfil.andar_bloco_display if perfil else '-',
        'foto_url': foto_url,
    }


def nome_responsavel_ou_placeholder(usuario):
    """Retorna o nome do responsável ou um marcador neutro quando o papel estiver vazio."""

    if not usuario:
        return '-'
    return usuario.get_full_name() or usuario.username


def cargo_responsavel_ou_placeholder(usuario):
    """Retorna o cargo de perfil para compor cartões e assinaturas institucionais."""

    perfil = getattr(usuario, 'perfil', None) if usuario else None
    return getattr(perfil, 'cargo', '') or '-'


def formatar_moeda_brl(valor):
    """Padroniza valores monetários em pt-BR para tabelas e documentos exportados."""

    return f'R$ {Decimal(valor or 0):,.2f}'.replace(',', 'X').replace('.', ',').replace('X', '.')


def montar_assinatura_pagamento(competencia, usuario, cargo_padrao, em_exercicio=False):
    """Centraliza a regra institucional: nome + cargo ou cargo em exercício."""

    assinatura = competencia.obter_assinatura_pagamento(usuario, cargo_padrao, em_exercicio)
    return {
        'nome': assinatura['nome'],
        'cargo': assinatura['cargo'],
    }


def redirect_contract_detail(contrato):
    return redirect('contratos:contrato_detail', pk=contrato.pk)


class ContractManagePermissionMixin:
    """Garante que gestor, criador e admin operem cadastros estruturantes do contrato."""

    contrato = None

    def ensure_manage_permission(self, request, contrato):
        if usuario_pode_gerir_contrato_v2(request.user, contrato):
            return None
        messages.error(
            request,
            'Somente o gestor do contrato, o criador do contrato ou administradores do sistema podem executar esta ação.',
        )
        return redirect_contract_detail(contrato)


class ContractOperatePermissionMixin:
    """Centraliza bloqueios operacionais por etapa da competência."""

    def deny(self, contrato, message):
        messages.error(self.request, message)
        return redirect_contract_detail(contrato)


class ContractImportantDocumentPermissionMixin:
    """Restringe os documentos importantes ao gestor atual e ao criador do contrato."""


class GlobalChecklistManagePermissionMixin:
    """Restringe a manutenção global dos padrões aos perfis institucionais do módulo."""

    def dispatch(self, request, *args, **kwargs):
        # O cadastro global não pertence a um contrato específico, então usamos a mesma noção
        # institucional já adotada no portal para gestores/admins do sistema.
        if not usuario_eh_admin_sistema(request.user):
            messages.error(
                request,
                'Somente gestores e administradores do sistema podem manter checklists padrão globais.',
            )
            return redirect('contratos:contrato_list')
        return super().dispatch(request, *args, **kwargs)

    def ensure_important_document_permission(self, request, contrato):
        if usuario_pode_gerir_documento_importante_contrato(request.user, contrato):
            return None
        messages.error(
            request,
            'Somente o gestor do contrato ou o usuário que criou o contrato podem gerenciar documentos importantes.',
        )
        return redirect_contract_detail(contrato)


class ContratoListView(ContratosAccessMixin, ListView):
    model = Contrato
    template_name = 'contratos/contrato_list.html'
    context_object_name = 'contratos'
    paginate_by = 20

    def get_queryset(self):
        queryset = (
            Contrato.objects.select_related(
                'empresa_contratada',
                'fiscal_administrativo__perfil',
                'fiscal_tecnico__perfil',
                'gestor_contrato__perfil',
            )
        )
        term = self.request.GET.get('q', '').strip()
        if term:
            queryset = queryset.filter(
                Q(numero_contrato__icontains=term)
                | Q(apelido__icontains=term)
                | Q(objeto__icontains=term)
                | Q(empresa_contratada__razao_social__icontains=term)
                | Q(empresa_contratada__cnpj__icontains=term)
                | Q(fiscal_administrativo__username__icontains=term)
                | Q(fiscal_administrativo__perfil__nome_completo__icontains=term)
                | Q(fiscal_tecnico__username__icontains=term)
                | Q(fiscal_tecnico__perfil__nome_completo__icontains=term)
                | Q(gestor_contrato__username__icontains=term)
                | Q(gestor_contrato__perfil__nome_completo__icontains=term)
            ).distinct()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['q'] = self.request.GET.get('q', '').strip()
        context['usuario_pode_gerir_checklists_padrao_globais'] = usuario_eh_admin_sistema(self.request.user)
        return context


class ChecklistPadraoGlobalListView(GlobalChecklistManagePermissionMixin, ContratosAccessMixin, ListView):
    """Entrega uma tela própria para a manutenção resumida dos checklists padrão globais."""

    model = ChecklistPadraoGlobal
    template_name = 'contratos/checklist_padrao_list.html'
    context_object_name = 'checklists_padrao'
    paginate_by = 20

    def get_queryset(self):
        return ChecklistPadraoGlobal.objects.select_related('atualizado_por__perfil').order_by('id')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['usuario_pode_gerir_checklists_padrao_globais'] = usuario_eh_admin_sistema(self.request.user)
        return context


class ChecklistPadraoGlobalDetailView(GlobalChecklistManagePermissionMixin, ContratosAccessMixin, DetailView):
    """Mostra um checklist padrão global isoladamente para manter seus itens sem poluir a tela de contratos."""

    model = ChecklistPadraoGlobal
    template_name = 'contratos/checklist_padrao_detail.html'
    context_object_name = 'checklist_padrao'

    def get_queryset(self):
        return ChecklistPadraoGlobal.objects.prefetch_related('itens').select_related('criado_por__perfil', 'atualizado_por__perfil')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['usuario_pode_gerir_checklists_padrao_globais'] = usuario_eh_admin_sistema(self.request.user)
        return context


class ContratoCreateView(ContratosWriteMixin, CreateView):
    model = Contrato
    form_class = ContratoForm
    template_name = 'contratos/contrato_form.html'

    def form_valid(self, form):
        assign_owner(form.instance, self.request)
        messages.success(self.request, 'Contrato V2 cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo contrato'
        return context


class ContratoDetailView(ContratosAccessMixin, DetailView):
    model = Contrato
    template_name = 'contratos/contrato_detail.html'
    context_object_name = 'contrato'

    def get_queryset(self):
        return Contrato.objects.select_related(
            'empresa_contratada',
            'fiscal_administrativo__perfil',
            'fiscal_administrativo_suplente__perfil',
            'fiscal_tecnico__perfil',
            'fiscal_tecnico_suplente__perfil',
            'gestor_contrato__perfil',
            'gestor_contrato_suplente__perfil',
        ).prefetch_related(
            'itens',
            Prefetch(
                'documentos_importantes',
                queryset=DocumentoImportanteContrato.objects.select_related('criado_por__perfil'),
            ),
            Prefetch('checklist_modelos', queryset=ChecklistModelo.objects.prefetch_related('itens')),
            Prefetch('formularios_avaliacao', queryset=FormularioAvaliacao.objects.prefetch_related('escalas', 'faixas_liberacao', 'grupos__itens')),
            Prefetch(
                'competencias',
                queryset=CompetenciaPagamento.objects.prefetch_related(
                    'checklist_itens__anexo',
                    'medicoes__item_contrato',
                    'avaliacao_qualidade__itens',
                ).annotate(
                    ordem_status=Case(
                        When(status='PAGA', then=Value(1)),
                        default=Value(0),
                        output_field=IntegerField(),
                    )
                ).order_by('ordem_status', '-periodo_inicio', '-id'),
            ),
            Prefetch('prazos', queryset=PrazoMonitoramento.objects.filter(concluido=False), to_attr='prazos_ativos'),
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        contrato = self.object
        context['responsaveis'] = [
            serializar_responsavel_v2('Fiscal administrativo', contrato.fiscal_administrativo),
            serializar_responsavel_v2('Fiscal administrativo suplente', contrato.fiscal_administrativo_suplente),
            serializar_responsavel_v2('Fiscal técnico', contrato.fiscal_tecnico),
            serializar_responsavel_v2('Fiscal técnico suplente', contrato.fiscal_tecnico_suplente),
            serializar_responsavel_v2('Gestor contrato', contrato.gestor_contrato),
            serializar_responsavel_v2('Gestor suplente', contrato.gestor_contrato_suplente),
        ]
        context['processos_sei'] = [
            {
                'label': 'Processo SEI (Gestão)',
                'numero': contrato.processo_sei_gestao_numero or '-',
                'url': contrato.processo_sei_gestao_url or '',
            },
            {
                'label': 'Processo SEI (Execução)',
                'numero': contrato.processo_sei_execucao_numero or '-',
                'url': contrato.processo_sei_execucao_url or '',
            },
        ]
        context['usuario_pode_gerir'] = contrato.usuario_pode_gerir(self.request.user)
        # Criar documento segue a permissão estrutural do contrato; editar/excluir seguem uma regra mais restrita.
        context['usuario_pode_criar_documentos_importantes'] = contrato.usuario_pode_gerir(self.request.user)
        context['usuario_pode_gerir_documentos_importantes'] = usuario_pode_gerir_documento_importante_contrato(
            self.request.user,
            contrato,
        )
        context['responsavel_empresa_principal'] = contrato.responsavel_empresa_principal
        documentos_importantes = list(contrato.documentos_importantes.all())
        for documento in documentos_importantes:
            documento.usuario_pode_gerir = usuario_pode_gerir_documento_importante(self.request.user, documento)
        context['documentos_importantes'] = documentos_importantes
        context['usuario_tem_acoes_documentos_importantes'] = any(
            documento.usuario_pode_gerir for documento in documentos_importantes
        )
        context['usuario_pode_checklist'] = contrato.usuario_pode_preencher_checklist(self.request.user)
        context['usuario_pode_medicao'] = contrato.usuario_pode_preencher_medicao(self.request.user)
        context['usuario_pode_avaliacao'] = contrato.usuario_pode_preencher_avaliacao(self.request.user)
        context['checklists_padrao_ativos'] = ChecklistPadraoGlobal.objects.filter(ativo=True).prefetch_related('itens')
        context['usuario_pode_carregar_checklist_padrao'] = (
            contrato.usuario_pode_gerir(self.request.user)
            and not contrato.competencias.exists()
        )

        # Cálculos da timeline de vigência
        hoje = timezone.now().date()
        inicio = contrato.data_inicio_vigencia
        meses_max = contrato.vigencia_maxima_meses or 0
        prazo_inicial = contrato.prazo_inicial_meses or 0
        
        if hoje < inicio:
            meses_decorridos = 0
        else:
            meses_decorridos = (hoje.year - inicio.year) * 12 + (hoje.month - inicio.month)
            if hoje.day >= inicio.day:
                meses_decorridos += 1

        meses_decorridos = max(0, min(meses_decorridos, meses_max))
        progresso_percentual = "{:.2f}".format((meses_decorridos / meses_max * 100) if meses_max > 0 else 0)

        timeline_marcos = []
        ano = inicio.year
        mes = inicio.month
        
        for i in range(meses_max + 1):
            is_start = (i == 0)
            is_initial = (i == prazo_inicial and i > 0)
            is_end = (i == meses_max and i > 0)
            is_current = (i == meses_decorridos)
            
            mostrar_texto = is_start or is_initial or is_end or is_current
            
            timeline_marcos.append({
                'indice': i,
                'mes_ano': f"{mes:02d}/{ano}",
                'percentual_posicao': "{:.2f}".format((i / meses_max * 100) if meses_max > 0 else 0),
                'mostrar_texto': mostrar_texto,
                'is_start': is_start,
                'is_initial': is_initial,
                'is_end': is_end,
                'is_current': is_current,
                'passado': i <= meses_decorridos
            })
            
            mes += 1
            if mes > 12:
                mes = 1
                ano += 1
                
        context['timeline'] = {
            'progresso_percentual': progresso_percentual,
            'marcos': timeline_marcos,
            'meses_decorridos': meses_decorridos,
            'meses_max': meses_max,
            'prazo_inicial': prazo_inicial,
        }
        return context


class ContratoUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, UpdateView):
    model = Contrato
    form_class = ContratoForm
    template_name = 'contratos/contrato_form.html'

    def form_valid(self, form):
        assign_owner(form.instance, self.request)
        messages.success(self.request, 'Contrato V2 atualizado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar contrato'
        return context


class ContratoDeleteView(ContratosWriteMixin, DeleteView):
    model = Contrato
    template_name = 'contratos/confirm_delete.html'
    success_url = reverse_lazy('contratos:contrato_list')

    def form_valid(self, form):
        # A exclusão do contrato precisa derrubar primeiro as estruturas mensais derivadas
        # para evitar o bloqueio por chaves protegidas das avaliações já geradas.
        self.object.excluir_com_dependencias()
        messages.success(self.request, 'Contrato V2 excluído com sucesso.')
        return redirect(self.get_success_url())


class ContratoChildCreateBase(ContratosWriteMixin):
    contrato = None

    def dispatch(self, request, *args, **kwargs):
        self.contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        return super().dispatch(request, *args, **kwargs)


class ContratoItemCreateView(BlockIfCompetenciasGeradasMixin, ContratoChildCreateBase, CreateView):
    model = ContratoItem
    form_class = ContratoItemForm
    template_name = 'contratos/contrato_item_form.html'

    def get_initial(self):
        initial = super().get_initial()
        initial['ordem'] = (self.contrato.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        return initial

    def form_valid(self, form):
        form.instance.contrato = self.contrato
        if not form.instance.ordem:
            form.instance.ordem = (self.contrato.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        elif self.contrato.itens.filter(ordem=form.instance.ordem).exists():
            form.add_error('ordem', 'Já existe um item com essa numeração neste contrato.')
            return self.form_invalid(form)
        messages.success(self.request, 'Item do contrato cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.contrato.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Novo item - {self.contrato.numero_contrato}'
        context['contrato'] = self.contrato
        return context


class ContratoItemUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, UpdateView):
    model = ContratoItem
    form_class = ContratoItemForm
    template_name = 'contratos/contrato_item_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        return ContratoItem.objects.filter(contrato=self.contrato)

    def form_valid(self, form):
        if not form.instance.ordem:
            form.instance.ordem = self.object.ordem
        conflito = self.contrato.itens.exclude(pk=self.object.pk).filter(ordem=form.instance.ordem).exists()
        if conflito:
            form.add_error('ordem', 'Já existe um item com essa numeração neste contrato.')
            return self.form_invalid(form)
        messages.success(self.request, 'Item do contrato atualizado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.contrato.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Editar item - {self.contrato.numero_contrato}'
        context['contrato'] = self.contrato
        return context


class ContratoItemDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, DeleteView):
    model = ContratoItem
    template_name = 'contratos/contrato_item_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        return ContratoItem.objects.filter(contrato=self.contrato)

    def get_success_url(self):
        messages.success(self.request, 'Item do contrato excluído com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.contrato.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['contrato'] = self.contrato
        return context


class DocumentoImportanteContratoCreateView(ContratoChildCreateBase, ContractImportantDocumentPermissionMixin, CreateView):
    """Cadastra um novo documento importante vinculado ao contrato."""

    model = DocumentoImportanteContrato
    form_class = DocumentoImportanteContratoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.contrato = self.contrato or get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        if not usuario_pode_gerir_contrato_v2(request.user, self.contrato):
            messages.error(
                request,
                'Somente o gestor do contrato, o criador do contrato ou administradores do sistema podem cadastrar documentos importantes.',
            )
            return redirect_contract_detail(self.contrato)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.contrato = self.contrato
        assign_owner(form.instance, self.request)
        messages.success(self.request, 'Documento importante cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.contrato.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo documento importante'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.contrato.pk])
        return context


class DocumentoImportanteContratoUpdateView(ContratosWriteMixin, ContractImportantDocumentPermissionMixin, UpdateView):
    """Edita os dados e o anexo de um documento importante do contrato."""

    model = DocumentoImportanteContrato
    form_class = DocumentoImportanteContratoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        if not usuario_pode_gerir_documento_importante(request.user, self.object):
            messages.error(
                request,
                'Somente o gestor do contrato, o criador do contrato, o autor do documento ou administradores do sistema podem editar este documento.',
            )
            return redirect_contract_detail(self.object.contrato)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        assign_owner(form.instance, self.request)
        messages.success(self.request, 'Documento importante atualizado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar documento importante'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.contrato_id])
        return context


class DocumentoImportanteContratoDeleteView(ContratosWriteMixin, ContractImportantDocumentPermissionMixin, DeleteView):
    """Exclui um documento importante do contrato com confirmação dedicada."""

    model = DocumentoImportanteContrato
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        if not usuario_pode_gerir_documento_importante(request.user, self.object):
            messages.error(
                request,
                'Somente o gestor do contrato, o criador do contrato, o autor do documento ou administradores do sistema podem excluir este documento.',
            )
            return redirect_contract_detail(self.object.contrato)
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Documento importante excluído com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir documento importante'
        context['descricao_objeto'] = self.object.nome
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.contrato_id])
        return context


class ChecklistModeloCreateView(BlockIfCompetenciasGeradasMixin, ContratoChildCreateBase, ContractManagePermissionMixin, CreateView):
    model = ChecklistModelo
    form_class = ChecklistModeloForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        # Esta tela valida permissão antes da cadeia normal de dispatch popular o contrato.
        self.contrato = self.contrato or get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        response = self.ensure_manage_permission(request, self.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.contrato = self.contrato
        messages.success(self.request, 'Versão de checklist cadastrada com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.contrato.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova versão de checklist'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.contrato.pk])
        return context


class ChecklistModeloUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = ChecklistModelo
    form_class = ChecklistModeloForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar versão de checklist'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.contrato_id])
        return context


class ChecklistModeloDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = ChecklistModelo
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Versão de checklist excluída com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir versão de checklist'
        context['descricao_objeto'] = self.object.nome
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.contrato_id])
        return context


class ChecklistModeloItemCreateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, CreateView):
    model = ChecklistModeloItem
    form_class = ChecklistModeloItemForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.modelo = get_object_or_404(ChecklistModelo, pk=kwargs['modelo_pk'])
        response = self.ensure_manage_permission(request, self.modelo.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_initial(self):
        initial = super().get_initial()
        initial['ordem'] = (self.modelo.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        return initial

    def form_valid(self, form):
        form.instance.modelo = self.modelo
        if not form.instance.ordem:
            form.instance.ordem = (self.modelo.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        messages.success(self.request, 'Item do checklist cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.modelo.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Novo item - {self.modelo.nome}'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.modelo.contrato_id])
        return context


class ChecklistModeloItemUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = ChecklistModeloItem
    form_class = ChecklistModeloItemForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.modelo.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.modelo.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar item do checklist'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.modelo.contrato_id])
        return context


class ChecklistModeloItemDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = ChecklistModeloItem
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.modelo.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Item do checklist excluído com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.modelo.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir item do checklist'
        context['descricao_objeto'] = self.object.titulo
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.modelo.contrato_id])
        return context


class ChecklistPadraoGlobalCreateView(GlobalChecklistManagePermissionMixin, ContratosWriteMixin, CreateView):
    """Cadastra um checklist padrão global reutilizável entre contratos."""

    model = ChecklistPadraoGlobal
    form_class = ChecklistPadraoGlobalForm
    template_name = 'contratos/entity_form.html'

    def form_valid(self, form):
        assign_owner(form.instance, self.request)
        messages.success(self.request, 'Checklist padrão global cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:checklist_padrao_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo checklist padrão'
        context['cancel_url'] = reverse('contratos:checklist_padrao_list')
        return context


class ChecklistPadraoGlobalUpdateView(GlobalChecklistManagePermissionMixin, ContratosWriteMixin, UpdateView):
    """Edita os dados estruturais do checklist padrão global."""

    model = ChecklistPadraoGlobal
    form_class = ChecklistPadraoGlobalForm
    template_name = 'contratos/entity_form.html'

    def form_valid(self, form):
        assign_owner(form.instance, self.request)
        messages.success(self.request, 'Checklist padrão global atualizado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:checklist_padrao_detail', args=[self.object.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar checklist padrão'
        context['cancel_url'] = reverse('contratos:checklist_padrao_detail', args=[self.object.pk])
        return context


class ChecklistPadraoGlobalDeleteView(GlobalChecklistManagePermissionMixin, ContratosWriteMixin, DeleteView):
    """Exclui um checklist padrão global e retorna à lista principal do módulo."""

    model = ChecklistPadraoGlobal
    template_name = 'contratos/entity_confirm_delete.html'

    def get_success_url(self):
        messages.success(self.request, 'Checklist padrão global excluído com sucesso.')
        return reverse('contratos:checklist_padrao_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir checklist padrão'
        context['descricao_objeto'] = self.object.nome
        context['cancel_url'] = reverse('contratos:checklist_padrao_list')
        return context


class ChecklistPadraoGlobalItemCreateView(GlobalChecklistManagePermissionMixin, ContratosWriteMixin, CreateView):
    """Adiciona um item à versão padrão global escolhida."""

    model = ChecklistPadraoGlobalItem
    form_class = ChecklistPadraoGlobalItemForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.checklist_padrao = get_object_or_404(ChecklistPadraoGlobal, pk=kwargs['checklist_padrao_pk'])
        return super().dispatch(request, *args, **kwargs)

    def get_initial(self):
        initial = super().get_initial()
        initial['ordem'] = (self.checklist_padrao.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        return initial

    def form_valid(self, form):
        form.instance.checklist_padrao = self.checklist_padrao
        if not form.instance.ordem:
            form.instance.ordem = (self.checklist_padrao.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        messages.success(self.request, 'Item do checklist padrão cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:checklist_padrao_detail', args=[self.checklist_padrao.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Novo item - {self.checklist_padrao.nome}'
        context['cancel_url'] = reverse('contratos:checklist_padrao_detail', args=[self.checklist_padrao.pk])
        return context


class ChecklistPadraoGlobalItemUpdateView(GlobalChecklistManagePermissionMixin, ContratosWriteMixin, UpdateView):
    """Edita um item do checklist padrão global."""

    model = ChecklistPadraoGlobalItem
    form_class = ChecklistPadraoGlobalItemForm
    template_name = 'contratos/entity_form.html'

    def get_success_url(self):
        messages.success(self.request, 'Item do checklist padrão atualizado com sucesso.')
        return reverse('contratos:checklist_padrao_detail', args=[self.object.checklist_padrao_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar item do checklist padrão'
        context['cancel_url'] = reverse('contratos:checklist_padrao_detail', args=[self.object.checklist_padrao_id])
        return context


class ChecklistPadraoGlobalItemDeleteView(GlobalChecklistManagePermissionMixin, ContratosWriteMixin, DeleteView):
    """Exclui um item do checklist padrão global."""

    model = ChecklistPadraoGlobalItem
    template_name = 'contratos/entity_confirm_delete.html'

    def get_success_url(self):
        messages.success(self.request, 'Item do checklist padrão excluído com sucesso.')
        return reverse('contratos:checklist_padrao_detail', args=[self.object.checklist_padrao_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir item do checklist padrão'
        context['descricao_objeto'] = self.object.titulo
        context['cancel_url'] = reverse('contratos:checklist_padrao_detail', args=[self.object.checklist_padrao_id])
        return context


class ChecklistPadraoCarregarView(BlockIfCompetenciasGeradasMixin, ContratosAccessMixin, ContractManagePermissionMixin, View):
    """Clona um checklist padrão ativo para dentro do contrato como nova versão ativa."""

    def post(self, request, *args, **kwargs):
        contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        response = self.ensure_manage_permission(request, contrato)
        if response:
            return response
        checklist_padrao_id = request.POST.get('checklist_padrao_id')
        if not checklist_padrao_id:
            messages.error(request, 'Selecione um checklist padrão antes de carregar.')
            return redirect_contract_detail(contrato)

        checklist_padrao = get_object_or_404(
            ChecklistPadraoGlobal.objects.prefetch_related('itens'),
            pk=checklist_padrao_id,
            ativo=True,
        )
        nova_versao = ChecklistModelo.objects.create(
            contrato=contrato,
            nome=f'{checklist_padrao.nome} (Padrão)',
            descricao=checklist_padrao.descricao,
            observacoes=checklist_padrao.observacoes,
            ativo=True,
        )
        for item in checklist_padrao.itens.order_by('ordem', 'id'):
            ChecklistModeloItem.objects.create(
                modelo=nova_versao,
                ordem=item.ordem,
                titulo=item.titulo,
                descricao=item.descricao,
                obrigatorio=item.obrigatorio,
            )
        messages.success(self.request, f'Checklist padrão "{checklist_padrao.nome}" carregado com sucesso no contrato.')
        return redirect_contract_detail(contrato)


class FormularioAvaliacaoCreateView(BlockIfCompetenciasGeradasMixin, ContratoChildCreateBase, ContractManagePermissionMixin, CreateView):
    model = FormularioAvaliacao
    form_class = FormularioAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        # Esta tela valida permissão antes da cadeia normal de dispatch popular o contrato.
        self.contrato = self.contrato or get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        response = self.ensure_manage_permission(request, self.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['instance'] = FormularioAvaliacao(contrato=self.contrato)
        return kwargs

    def form_valid(self, form):
        form.instance.contrato = self.contrato
        response = super().form_valid(form)
        messages.success(self.request, 'Formulário de avaliação cadastrado com sucesso.')
        return response

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.contrato.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo formulário de avaliação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.contrato.pk])
        return context


class FormularioAvaliacaoUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = FormularioAvaliacao
    form_class = FormularioAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        response = super().form_valid(form)
        messages.success(self.request, 'Formulário de avaliação atualizado com sucesso.')
        return response

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar formulário de avaliação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.contrato_id])
        return context


class FormularioAvaliacaoDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = FormularioAvaliacao
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def delete(self, request, *args, **kwargs):
        try:
            return super().delete(request, *args, **kwargs)
        except ValidationError as exc:
            messages.error(request, exc.message)
            return redirect_contract_detail(self.object.contrato)

    def get_success_url(self):
        messages.success(self.request, 'Formulário de avaliação excluído com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir formulário de avaliação'
        context['descricao_objeto'] = self.object.nome
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.contrato_id])
        return context


class EscalaNotaAvaliacaoCreateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, CreateView):
    model = EscalaNotaAvaliacao
    form_class = EscalaNotaAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.formulario = get_object_or_404(FormularioAvaliacao, pk=kwargs['formulario_pk'])
        response = self.ensure_manage_permission(request, self.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.formulario = self.formulario
        form.instance.ordem = (self.formulario.escalas.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        messages.success(self.request, 'Nota da escala cadastrada com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova nota da escala'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.formulario.contrato_id])
        return context


class EscalaNotaAvaliacaoUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = EscalaNotaAvaliacao
    form_class = EscalaNotaAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar nota da escala'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])
        return context


class EscalaNotaAvaliacaoDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = EscalaNotaAvaliacao
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Nota da escala excluída com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir nota da escala'
        context['descricao_objeto'] = str(self.object)
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])
        return context


class FaixaLiberacaoAvaliacaoCreateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, CreateView):
    model = FaixaLiberacaoAvaliacao
    form_class = FaixaLiberacaoAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.formulario = get_object_or_404(FormularioAvaliacao, pk=kwargs['formulario_pk'])
        response = self.ensure_manage_permission(request, self.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.formulario = self.formulario
        form.instance.ordem = (self.formulario.faixas_liberacao.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        messages.success(self.request, 'Faixa de liberação cadastrada com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova faixa de liberação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.formulario.contrato_id])
        return context


class FaixaLiberacaoAvaliacaoUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = FaixaLiberacaoAvaliacao
    form_class = FaixaLiberacaoAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar faixa de liberação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])
        return context


class FaixaLiberacaoAvaliacaoDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = FaixaLiberacaoAvaliacao
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Faixa de liberação excluída com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir faixa de liberação'
        context['descricao_objeto'] = str(self.object)
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])
        return context


class GrupoAvaliacaoCreateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, CreateView):
    model = GrupoAvaliacao
    form_class = GrupoAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.formulario = get_object_or_404(FormularioAvaliacao, pk=kwargs['formulario_pk'])
        response = self.ensure_manage_permission(request, self.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.formulario = self.formulario
        form.instance.ordem = (self.formulario.grupos.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        messages.success(self.request, 'Grupo de avaliação cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo grupo de avaliação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.formulario.contrato_id])
        return context


class GrupoAvaliacaoUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = GrupoAvaliacao
    form_class = GrupoAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar grupo de avaliação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])
        return context


class GrupoAvaliacaoDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = GrupoAvaliacao
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Grupo de avaliação excluído com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir grupo de avaliação'
        context['descricao_objeto'] = self.object.nome
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.formulario.contrato_id])
        return context


class ItemAvaliacaoCreateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, CreateView):
    model = ItemAvaliacao
    form_class = ItemAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.grupo = get_object_or_404(GrupoAvaliacao, pk=kwargs['grupo_pk'])
        response = self.ensure_manage_permission(request, self.grupo.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.grupo = self.grupo
        # A ordem do item é sempre sequencial dentro do grupo para evitar divergência entre tela e banco.
        form.instance.ordem = (self.grupo.itens.aggregate(max_ordem=Max('ordem'))['max_ordem'] or 0) + 1
        messages.success(self.request, 'Item de avaliação cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.grupo.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Novo item de avaliação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.grupo.formulario.contrato_id])
        return context


class ItemAvaliacaoUpdateView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = ItemAvaliacao
    form_class = ItemAvaliacaoForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.grupo.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.grupo.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar item de avaliação'
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.grupo.formulario.contrato_id])
        return context


class ItemAvaliacaoDeleteView(BlockIfCompetenciasGeradasMixin, ContratosWriteMixin, ContractManagePermissionMixin, DeleteView):
    model = ItemAvaliacao
    template_name = 'contratos/entity_confirm_delete.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.grupo.formulario.contrato)
        if response:
            return response
        return super().dispatch(request, *args, **kwargs)

    def get_success_url(self):
        messages.success(self.request, 'Item de avaliação excluído com sucesso.')
        return reverse('contratos:contrato_detail', args=[self.object.grupo.formulario.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Excluir item de avaliação'
        context['descricao_objeto'] = self.object.descricao[:120]
        context['cancel_url'] = reverse('contratos:contrato_detail', args=[self.object.grupo.formulario.contrato_id])
        return context


class CompetenciasGenerateView(ContratosWriteMixin, ContractManagePermissionMixin, View):
    def post(self, request, *args, **kwargs):
        contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        response = self.ensure_manage_permission(request, contrato)
        if response:
            return response
        try:
            contrato.gerar_competencias()
        except ValidationError as exc:
            messages.error(request, exc.message)
            return redirect_contract_detail(contrato)
        messages.success(request, 'Competências geradas com sucesso.')
        return redirect_contract_detail(contrato)


class CompetenciaChecklistUpdateView(ContratosWriteMixin, ContractOperatePermissionMixin, FormView):
    form_class = CompetenciaChecklistUploadForm
    template_name = 'contratos/competencia_checklist_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.competencia = get_object_or_404(
            CompetenciaPagamento.objects.select_related(
                'contrato',
                'medicao_preenchida_por__perfil',
                'aceite_provisorio_preenchida_por__perfil',
                'aceite_definitivo_preenchida_por__perfil',
                'nota_principal_preenchida_por__perfil',
                'nota_adicional_preenchida_por__perfil',
                'observacoes_finais_preenchida_por__perfil',
            ),
            pk=kwargs['competencia_pk'],
        )
        if not self.competencia.contrato.usuario_pode_preencher_checklist(request.user):
            return self.deny(self.competencia.contrato, 'Você não pode preencher o checklist desta competência.')
        if not self.competencia.medicao_concluida_em:
            return self.deny(self.competencia.contrato, 'Só é possível partir para o checklist com a medição concluída.')
        if self.competencia.exige_avaliacao and not getattr(self.competencia.avaliacao_qualidade_segura, 'concluida_em', None):
            return self.deny(self.competencia.contrato, 'Só é possível partir para o checklist após concluir a avaliação.')
        if self.competencia.status in {self.competencia.Status.PAGA, self.competencia.Status.CANCELADA}:
            return self.deny(self.competencia.contrato, 'A competência já foi encerrada.')
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['competencia'] = self.competencia
        return kwargs

    def form_valid(self, form):
        for item in form.itens:
            novo_arquivo = form.cleaned_data.get(f'arquivo_{item.pk}')
            limpar = form.cleaned_data.get(f'limpar_{item.pk}')
            existente = item.anexo_principal
            if limpar and existente:
                existente.delete()
                continue
            if not novo_arquivo:
                continue
            if existente:
                existente.arquivo = novo_arquivo
                existente.nome_exibicao = ''
                existente.save(update_fields=['arquivo', 'nome_exibicao', 'atualizado_em'])
            else:
                ChecklistCompetenciaAnexo.objects.create(item=item, arquivo=novo_arquivo, nome_exibicao='')
        recalcular_competencia_v2(self.competencia)
        messages.success(self.request, 'Checklist atualizado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:competencia_checklist', args=[self.competencia.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Checklist da competência'
        context['competencia'] = self.competencia
        context['contrato'] = self.competencia.contrato
        return context


class CompetenciaMedicaoUpdateView(ContratosWriteMixin, ContractOperatePermissionMixin, FormView):
    form_class = CompetenciaMedicaoLoteV2Form
    template_name = 'contratos/competencia_medicao_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.competencia = get_object_or_404(
            CompetenciaPagamento.objects.select_related('contrato'),
            pk=kwargs['competencia_pk'],
        )
        if not self.competencia.contrato.usuario_pode_preencher_medicao(request.user):
            return self.deny(self.competencia.contrato, 'Você não pode preencher a medição desta competência.')
        if self.competencia.status in {self.competencia.Status.PAGA, self.competencia.Status.CANCELADA}:
            return self.deny(self.competencia.contrato, 'A competência já foi encerrada.')
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['contrato'] = self.competencia.contrato
        kwargs['competencia'] = self.competencia
        return kwargs

    def form_valid(self, form):
        momento_auditoria = timezone.now()
        medicoes_atuais = {
            medicao.item_contrato_id: medicao.quantidade
            for medicao in self.competencia.medicoes.all()
        }
        medicao_alterada = self.competencia.aplicar_pro_rata != bool(form.cleaned_data.get('aplicar_pro_rata', False))
        secoes_alteradas = {
            'aceite_provisorio': False,
            'aceite_definitivo': False,
            'nota_principal': False,
            'nota_adicional': False,
            'observacoes_finais': False,
        }

        self.competencia.aplicar_pro_rata = bool(form.cleaned_data.get('aplicar_pro_rata', False))
        for item in form.itens:
            quantidade = form.cleaned_data.get(f'quantidade_{item.pk}')
            if quantidade in (None, ''):
                continue
            quantidade_atual = medicoes_atuais.get(item.pk) or Decimal('0.00')
            quantidade_nova = Decimal('0.00') if quantidade == 0 else quantidade
            if quantidade_atual != quantidade_nova:
                medicao_alterada = True
            if quantidade == 0:
                MedicaoItemCompetencia.objects.filter(competencia=self.competencia, item_contrato=item).delete()
                continue
            MedicaoItemCompetencia.objects.update_or_create(
                competencia=self.competencia,
                item_contrato=item,
                defaults={
                    'quantidade': quantidade,
                    'valor_unitario_aplicado': item.valor_unitario,
                    'observacoes': '',
                },
            )
        campos_financeiros = {
            'valor_nota_fiscal',
            'retencao_ir',
            'retencao_inss',
            'retencao_iss',
            'retencao_pis_pasep',
            'retencao_cofins',
            'valor_liberado_final',
            'valor_nota_adicional',
            'retencao_ir_adicional',
            'retencao_inss_adicional',
            'retencao_iss_adicional',
            'retencao_pis_pasep_adicional',
            'retencao_cofins_adicional',
            'valor_liquido_nota_adicional',
        }
        mapa_secao_campos = {
            'aceite_provisorio': {'data_aceite_provisorio', 'prazo_aceite_definitivo_dias'},
            'aceite_definitivo': {'data_aceite_definitivo', 'prazo_pagamento_dias'},
            'nota_principal': {
                'numero_nota_fiscal',
                'valor_nota_fiscal',
                'retencao_ir',
                'retencao_inss',
                'retencao_iss',
                'retencao_pis_pasep',
                'retencao_cofins',
                'valor_liberado_final',
            },
            'nota_adicional': {
                'nota_adicional_nao_consta',
                'numero_nota_adicional',
                'valor_nota_adicional',
                'retencao_ir_adicional',
                'retencao_inss_adicional',
                'retencao_iss_adicional',
                'retencao_pis_pasep_adicional',
                'retencao_cofins_adicional',
                'valor_liquido_nota_adicional',
            },
            'observacoes_finais': {'observacoes_medicao'},
        }
        for campo in (
            'data_aceite_provisorio',
            'prazo_aceite_definitivo_dias',
            'data_aceite_definitivo',
            'prazo_pagamento_dias',
            'numero_nota_fiscal',
            'valor_nota_fiscal',
            'retencao_ir',
            'retencao_inss',
            'retencao_iss',
            'retencao_pis_pasep',
            'retencao_cofins',
            'valor_liberado_final',
            'numero_nota_adicional',
            'nota_adicional_nao_consta',
            'valor_nota_adicional',
            'retencao_ir_adicional',
            'retencao_inss_adicional',
            'retencao_iss_adicional',
            'retencao_pis_pasep_adicional',
            'retencao_cofins_adicional',
            'valor_liquido_nota_adicional',
            'observacoes_medicao',
        ):
            valor = form.cleaned_data.get(campo)
            if campo in campos_financeiros and valor in (None, ''):
                valor = Decimal('0.00')
            if valor != getattr(self.competencia, campo):
                for secao, campos_secao in mapa_secao_campos.items():
                    if campo in campos_secao:
                        secoes_alteradas[secao] = True
                        break
            setattr(self.competencia, campo, valor)

        if self.competencia.nota_adicional_nao_consta:
            self.competencia.nota_adicional_arquivo = ''
            self.competencia.numero_nota_adicional = ''
            self.competencia.valor_nota_adicional = Decimal('0.00')
            self.competencia.retencao_ir_adicional = Decimal('0.00')
            self.competencia.retencao_inss_adicional = Decimal('0.00')
            self.competencia.retencao_iss_adicional = Decimal('0.00')
            self.competencia.retencao_pis_pasep_adicional = Decimal('0.00')
            self.competencia.retencao_cofins_adicional = Decimal('0.00')
            self.competencia.valor_liquido_nota_adicional = Decimal('0.00')

        for campo, secao in (
            ('aceite_provisorio_arquivo', 'aceite_provisorio'),
            ('aceite_definitivo_arquivo', 'aceite_definitivo'),
            ('nota_fiscal_fatura', 'nota_principal'),
            ('nota_adicional_arquivo', 'nota_adicional'),
        ):
            arquivo = form.cleaned_data.get(campo)
            if arquivo and not getattr(arquivo, '_committed', False):
                secoes_alteradas[secao] = True
                setattr(self.competencia, campo, arquivo)

        if self.competencia.data_aceite_provisorio and self.competencia.prazo_aceite_definitivo_dias:
            self.competencia.monitoramento_etapa = 'Aguardando Aceite Definitivo do Gestor'
            self.competencia.monitoramento_inicio = self.competencia.data_aceite_provisorio
            self.competencia.monitoramento_limite = self.competencia.data_aceite_provisorio + timedelta(days=self.competencia.prazo_aceite_definitivo_dias)
        if (
            self.competencia.aceite_definitivo_arquivo
            and self.competencia.data_aceite_definitivo
            and self.competencia.prazo_pagamento_dias
            and self.competencia.nota_fiscal_fatura
            and (self.competencia.numero_nota_fiscal or '').strip()
            and not self.competencia.nota_adicional_resolvida
        ):
            self.competencia.monitoramento_etapa = 'Aguardando Nota Fiscal Adicional'
            self.competencia.monitoramento_inicio = self.competencia.data_aceite_definitivo
            self.competencia.monitoramento_limite = self.competencia.data_aceite_definitivo + timedelta(days=self.competencia.prazo_pagamento_dias or 0)
        if competencia_medicao_v2_esta_concluida(self.competencia):
            self.competencia.monitoramento_etapa = 'Aguardando Conclusão do Pagamento'
            self.competencia.monitoramento_inicio = self.competencia.data_aceite_definitivo
            self.competencia.monitoramento_limite = self.competencia.data_aceite_definitivo + timedelta(days=self.competencia.prazo_pagamento_dias or 0)

        # Cada seção mantém seu último responsável e horário de salvamento para o usuário
        # visualizar claramente quem consolidou aquele bloco da medição.
        if medicao_alterada and self.competencia.medicao_tem_conteudo:
            self.competencia.medicao_preenchida_por = self.request.user
            self.competencia.medicao_preenchida_em = momento_auditoria
        elif not self.competencia.medicao_tem_conteudo:
            self.competencia.medicao_preenchida_por = None
            self.competencia.medicao_preenchida_em = None

        for prefixo, alterada, possui_conteudo in (
            ('aceite_provisorio', secoes_alteradas['aceite_provisorio'], self.competencia.aceite_provisorio_tem_conteudo),
            ('aceite_definitivo', secoes_alteradas['aceite_definitivo'], self.competencia.aceite_definitivo_tem_conteudo),
            ('nota_principal', secoes_alteradas['nota_principal'], self.competencia.nota_principal_tem_conteudo),
            ('nota_adicional', secoes_alteradas['nota_adicional'], self.competencia.nota_adicional_tem_conteudo),
            ('observacoes_finais', secoes_alteradas['observacoes_finais'], self.competencia.observacoes_finais_tem_conteudo),
        ):
            if alterada and possui_conteudo:
                setattr(self.competencia, f'{prefixo}_preenchida_por', self.request.user)
                setattr(self.competencia, f'{prefixo}_preenchida_em', momento_auditoria)
            elif not possui_conteudo:
                setattr(self.competencia, f'{prefixo}_preenchida_por', None)
                setattr(self.competencia, f'{prefixo}_preenchida_em', None)

        self.competencia.save()
        recalcular_competencia_v2(self.competencia)
        messages.success(self.request, 'Medição atualizada com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.competencia.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Medição da competência'
        context['competencia'] = self.competencia
        context['contrato'] = self.competencia.contrato
        context['itens_medicao'] = [
            {'item': item, 'field': context['form'][f'quantidade_{item.pk}']}
            for item in context['form'].itens
        ]
        context['itens_checklist_extra'] = self.competencia.checklist_itens.filter(
            categoria=ChecklistCompetenciaItem.Categoria.NOTA_ADICIONAL
        ).order_by('ordem', 'id')
        return context


class CompetenciaAvaliacaoUpdateView(ContratosWriteMixin, ContractOperatePermissionMixin, FormView):
    form_class = AvaliacaoCompetenciaV2Form
    template_name = 'contratos/competencia_avaliacao_form.html'

    def _avaliacao_possui_conteudo_salvo(self):
        """Libera o download assim que houver conteúdo suficiente para revisão e assinatura externa."""

        if not self.avaliacao:
            return False
        if (self.avaliacao.observacoes or '').strip():
            return True
        return self.avaliacao.itens.filter(
            Q(nota_fiscal_valor__isnull=False)
            | Q(nota_gestor_valor__isnull=False)
            | ~Q(justificativa_fiscal='')
            | ~Q(manifestacao_gestor_item='')
        ).exists()

    def dispatch(self, request, *args, **kwargs):
        self.competencia = get_object_or_404(
            CompetenciaPagamento.objects.select_related('contrato'),
            pk=kwargs['competencia_pk'],
        )
        if not self.competencia.contrato.usuario_pode_preencher_avaliacao(request.user):
            return self.deny(self.competencia.contrato, 'Você não pode preencher a avaliação desta competência.')
        if not self.competencia.medicao_concluida_em:
            return self.deny(self.competencia.contrato, 'Só é possível partir para a avaliação com a medição concluída.')
        if not self.competencia.exige_avaliacao:
            return self.deny(self.competencia.contrato, 'Esta competência não exige avaliação de qualidade.')
        if self.competencia.status in {self.competencia.Status.PAGA, self.competencia.Status.CANCELADA}:
            return self.deny(self.competencia.contrato, 'A competência já foi encerrada.')
        self.pode_preencher_fiscal = usuario_pode_preencher_avaliacao_fiscal_v2(request.user, self.competencia.contrato)
        self.pode_preencher_gestor = usuario_pode_preencher_avaliacao_gestor_v2(request.user, self.competencia.contrato)
        self.avaliacao = self.competencia.avaliacao_qualidade_segura
        if self.avaliacao is None and self.competencia.contrato.formulario_avaliacao_ativo:
            self.avaliacao = criar_avaliacao_shell_competencia_v2(self.competencia, self.competencia.contrato.formulario_avaliacao_ativo)
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['avaliacao'] = self.avaliacao
        kwargs['pode_preencher_fiscal'] = self.pode_preencher_fiscal
        kwargs['pode_preencher_gestor'] = self.pode_preencher_gestor
        return kwargs

    def form_valid(self, form):
        escala_mapa = {
            Decimal(item['valor']): item['legenda']
            for item in self.avaliacao.formulario_snapshot.get('escala', [])
        }
        agora = timezone.now()

        def normalizar_nota_para_persistencia(valor):
            """Evita que string vazia chegue ao DecimalField quando um papel salvar de forma independente."""

            if valor in (None, ''):
                return None
            if isinstance(valor, Decimal):
                return valor
            return Decimal(str(valor))

        for resposta in form.respostas:
            update_fields = []
            campo_nota_fiscal = f'nota_fiscal_{resposta.pk}'
            campo_justificativa_fiscal = f'justificativa_fiscal_{resposta.pk}'
            campo_nota_gestor = f'nota_gestor_{resposta.pk}'
            campo_manifestacao_gestor = f'manifestacao_gestor_item_{resposta.pk}'
            papel_informado = form.papeis_informados.get(resposta.pk, {})

            if self.pode_preencher_fiscal and papel_informado.get('fiscal'):
                nota_fiscal = normalizar_nota_para_persistencia(form.cleaned_data.get(campo_nota_fiscal))
                justificativa_fiscal = form.cleaned_data.get(campo_justificativa_fiscal) or ''
                resposta.nota_fiscal_valor = nota_fiscal
                resposta.nota_fiscal_preenchida_por = self.request.user
                resposta.nota_fiscal_preenchida_em = agora
                resposta.justificativa_fiscal = justificativa_fiscal
                resposta.justificativa_fiscal_preenchida_por = self.request.user
                resposta.justificativa_fiscal_preenchida_em = agora
                update_fields.extend([
                    'nota_fiscal_valor',
                    'nota_fiscal_preenchida_por',
                    'nota_fiscal_preenchida_em',
                    'justificativa_fiscal',
                    'justificativa_fiscal_preenchida_por',
                    'justificativa_fiscal_preenchida_em',
                ])

            if self.pode_preencher_gestor and papel_informado.get('gestor'):
                nota_gestor = normalizar_nota_para_persistencia(form.cleaned_data.get(campo_nota_gestor))
                manifestacao_gestor = form.cleaned_data.get(campo_manifestacao_gestor) or ''
                resposta.nota_gestor_valor = nota_gestor
                resposta.nota_gestor_preenchida_por = self.request.user
                resposta.nota_gestor_preenchida_em = agora
                resposta.manifestacao_gestor_item = manifestacao_gestor
                resposta.manifestacao_gestor_item_preenchida_por = self.request.user
                resposta.manifestacao_gestor_item_preenchida_em = agora
                update_fields.extend([
                    'nota_gestor_valor',
                    'nota_gestor_preenchida_por',
                    'nota_gestor_preenchida_em',
                    'manifestacao_gestor_item',
                    'manifestacao_gestor_item_preenchida_por',
                    'manifestacao_gestor_item_preenchida_em',
                ])

            resposta.nota_valor = resposta.nota_vigente
            resposta.nota_legenda = escala_mapa.get(resposta.nota_valor, '')
            update_fields.extend(['nota_valor', 'nota_legenda', 'atualizado_em'])
            resposta.save(update_fields=sorted(set(update_fields)))
        self.avaliacao.observacoes = form.cleaned_data.get('observacoes') or ''
        self.avaliacao.preenchido_por = self.request.user
        if form.cleaned_data.get('avaliacao_assinada'):
            self.competencia.avaliacao_assinada = form.cleaned_data['avaliacao_assinada']
            self.competencia.save(update_fields=['avaliacao_assinada', 'atualizado_em'])
        # Mantém a competência em avaliação pendente até o gestor preencher as manifestações exigidas.
        self.avaliacao.concluida_em = timezone.now() if avaliacao_v2_esta_concluida(self.avaliacao) else None
        self.avaliacao.save(update_fields=['observacoes', 'preenchido_por', 'concluida_em', 'atualizado_em'])
        recalcular_avaliacao_v2(self.avaliacao)
        if self.avaliacao.concluida_em:
            messages.success(self.request, 'Avaliação concluída com sucesso.')
        elif self.competencia.avaliacao_assinada:
            messages.warning(self.request, 'Avaliação salva. A etapa permanecerá pendente até que todos os itens obrigatórios sejam preenchidos.')
        else:
            messages.warning(self.request, 'Avaliação salva. Baixe o relatório, colete as assinaturas e depois anexe o PDF assinado para concluir a etapa.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.competencia.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Avaliação da competência'
        context['competencia'] = self.competencia
        context['contrato'] = self.competencia.contrato
        context['avaliacao'] = self.avaliacao
        context['avaliacao_tem_conteudo'] = self._avaliacao_possui_conteudo_salvo()
        return context


class CompetenciaAvaliacaoDownloadView(ContratosWriteMixin, ContractOperatePermissionMixin, View):
    """Entrega o relatório da avaliação em PDF para assinatura externa pelas partes."""

    def get(self, request, *args, **kwargs):
        competencia = get_object_or_404(
            CompetenciaPagamento.objects.select_related('contrato'),
            pk=kwargs['competencia_pk'],
        )
        if not competencia.contrato.usuario_pode_preencher_avaliacao(request.user):
            return self.deny(competencia.contrato, 'Você não pode baixar a avaliação desta competência.')
        avaliacao = competencia.avaliacao_qualidade_segura
        if not avaliacao:
            return self.deny(competencia.contrato, 'A competência ainda não possui avaliação gerada.')
        possui_conteudo = avaliacao.itens.filter(
            Q(nota_fiscal_valor__isnull=False)
            | Q(nota_gestor_valor__isnull=False)
            | ~Q(justificativa_fiscal='')
            | ~Q(manifestacao_gestor_item='')
        ).exists() or bool((avaliacao.observacoes or '').strip())
        if not possui_conteudo:
            return self.deny(competencia.contrato, 'Salve ao menos um item da avaliação antes de baixar o relatório.')

        with tempfile.TemporaryDirectory(prefix=f'avaliacao_competencia_{competencia.pk}_') as temp_dir:
            temp_dir_path = Path(temp_dir)
            caminho_docx = temp_dir_path / f'avaliacao_competencia_{competencia.pk}.docx'
            # Reaproveita o mesmo documento institucional da exportação para manter uma única fonte da verdade.
            doc = gerar_relatorio_avaliacao_competencia('/root/aplicacoesspi/docs/papel-timbrado-spi.docx', competencia.contrato, competencia)
            doc.save(caminho_docx)
            caminho_pdf = _converter_docx_para_pdf(caminho_docx, temp_dir_path)
            response = FileResponse(open(caminho_pdf, 'rb'), content_type='application/pdf')
            response['Content-Disposition'] = (
                f'attachment; filename="avaliacao_competencia_{competencia.periodo_inicio:%m_%Y}.pdf"'
            )
            return response


class CompetenciaChecklistExtraItemCreateView(ContratosWriteMixin, ContractManagePermissionMixin, FormView):
    """Cadastra itens documentais extras vinculados exclusivamente à nota adicional."""

    form_class = CompetenciaChecklistExtraItemForm
    template_name = 'contratos/entity_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.competencia = get_object_or_404(CompetenciaPagamento.objects.select_related('contrato'), pk=kwargs['competencia_pk'])
        response = self.ensure_manage_permission(request, self.competencia.contrato)
        if response:
            return response
        if not self.competencia.possui_nota_adicional:
            messages.error(request, 'Cadastre a nota adicional antes de criar documentos específicos para ela.')
            return redirect_contract_detail(self.competencia.contrato)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        ordem = (self.competencia.checklist_itens.aggregate(maior=Max('ordem')).get('maior') or 0) + 1
        item = ChecklistCompetenciaItem.objects.create(
            competencia=self.competencia,
            ordem=ordem,
            titulo=form.cleaned_data['titulo'],
            categoria=ChecklistCompetenciaItem.Categoria.NOTA_ADICIONAL,
            obrigatorio=True,
        )
        ChecklistCompetenciaAnexo.objects.create(item=item, arquivo=form.cleaned_data['arquivo'], nome_exibicao='')
        messages.success(self.request, 'Documento adicional do checklist cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:competencia_medicao', args=[self.competencia.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Novo documento da nota adicional - {self.competencia.periodo_inicio:%m/%Y}'
        context['cancel_url'] = reverse('contratos:competencia_medicao', args=[self.competencia.pk])
        return context


class CompetenciaOBExecuteView(ContratosWriteMixin, ContractManagePermissionMixin, UpdateView):
    model = CompetenciaPagamento
    form_class = CompetenciaOBExecucaoForm
    template_name = 'contratos/competencia_ob_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.object = self.get_object()
        response = self.ensure_manage_permission(request, self.object.contrato)
        if response:
            return response
        if self.object.status != self.object.Status.OB_PENDENTE:
            messages.error(request, 'Só pode partir para a OB após concluir medição, avaliação, checklist e download.')
            return redirect_contract_detail(self.object.contrato)
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.status = self.object.Status.PAGA
        if not self.object.data_pagamento:
            self.object.data_pagamento = timezone.localdate()
        self.object.monitoramento_etapa = ''
        self.object.monitoramento_inicio = None
        self.object.monitoramento_limite = None
        self.object.save()
        messages.success(self.request, 'Ordem bancária registrada com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', args=[self.object.contrato_id])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Anexar Ordem Bancária'
        context['competencia'] = self.object
        context['contrato'] = self.object.contrato
        return context


class CompetenciaPagamentoExecuteView(CompetenciaOBExecuteView):
    """Mantém compatibilidade com a rota antiga de pagamento durante a transição do fluxo."""


@login_required
def proximo_numero_contrato(request):
    ano = request.GET.get('ano')
    if not ano or not ano.isdigit() or len(ano) != 4:
        return JsonResponse({'numero': ''}, status=400)
    
    numero = numero_contrato_por_ano(int(ano))
    return JsonResponse({'numero': numero})



class EmpresaListView(ContratosAccessMixin, ListView):
    """Lista empresas contratadas com atalho para manutenção."""

    model = EmpresaContratada
    template_name = 'contratos/empresa_list.html'
    context_object_name = 'empresas'


class EmpresaCreateView(ContratosWriteMixin, CreateView):
    """Cadastro de empresa contratada."""

    model = EmpresaContratada
    form_class = EmpresaContratadaForm
    template_name = 'contratos/empresa_form.html'
    def render_to_response(self, context, **response_kwargs):
        if is_modal_request(self.request):
            return render(self.request, 'contratos/partials/empresa_modal_form.html', context, **response_kwargs)
        return super().render_to_response(context, **response_kwargs)

    def form_valid(self, form):
        self.object = form.save()
        if is_modal_request(self.request):
            return JsonResponse(
                {
                    'success': True,
                    'empresa': {
                        'id': self.object.pk,
                        'label': self.object.razao_social,
                        'cnpj': self.object.cnpj,
                    },
                }
            )
        messages.success(self.request, 'Empresa cadastrada com sucesso.')
        return redirect(self.get_success_url())

    def get_success_url(self):
        return reverse('contratos:responsavel_create', args=[self.object.pk])

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Nova empresa contratada'
        context['modal_mode'] = is_modal_request(self.request)
        return context


class EmpresaUpdateView(ContratosWriteMixin, UpdateView):
    """Edição dos dados da empresa contratada."""

    model = EmpresaContratada
    form_class = EmpresaContratadaForm
    template_name = 'contratos/empresa_form.html'
    success_url = reverse_lazy('contratos:empresa_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = 'Editar empresa contratada'
        return context


class EmpresaDeleteView(ContratosWriteMixin, DeleteView):
    """Exclusão simples da empresa quando ainda for permitida operacionalmente."""

    model = EmpresaContratada
    template_name = 'contratos/confirm_delete.html'
    success_url = reverse_lazy('contratos:empresa_list')


class ResponsavelEmpresaCreateView(ContratosWriteMixin, CreateView):
    """Adiciona um responsável ao cadastro da empresa."""

    model = ResponsavelEmpresa
    form_class = ResponsavelEmpresaForm
    template_name = 'contratos/empresa_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.empresa = get_object_or_404(EmpresaContratada, pk=kwargs['empresa_pk'])
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.empresa = self.empresa
        messages.success(self.request, 'Responsável cadastrado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:empresa_list')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo'] = f'Novo responsável - {self.empresa.razao_social}'
        return context


# =========================================================================
# MONITORAMENTO DE PRAZOS (CONTRATO PÓS-COMPETÊNCIAS)
# =========================================================================

class CheckCompetenciasGeradasMixin:
    """Garante que a operação sobre prazos só ocorra se o contrato tiver competências geradas."""

    def dispatch(self, request, *args, **kwargs):
        contrato = None
        if 'contrato_pk' in kwargs:
            contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        elif hasattr(self, 'get_object'):
            obj = self.get_object()
            contrato = obj.contrato
            
        if not contrato or not contrato.competencias.exists():
            messages.error(request, 'Ação não permitida: O monitoramento de prazos só é liberado após a geração de competências.')
            return redirect('contratos:contrato_detail', pk=contrato.pk if contrato else 1)
            
        return super().dispatch(request, *args, **kwargs)


class PrazoMonitoramentoCreateView(ContratosWriteMixin, CheckCompetenciasGeradasMixin, CreateView):
    """Cria um novo prazo de monitoramento associado ao contrato."""
    model = PrazoMonitoramento
    form_class = PrazoMonitoramentoForm
    template_name = 'contratos/prazo_form.html'

    def dispatch(self, request, *args, **kwargs):
        self.contrato = get_object_or_404(Contrato, pk=kwargs['contrato_pk'])
        return super().dispatch(request, *args, **kwargs)

    def form_valid(self, form):
        form.instance.contrato = self.contrato
        messages.success(self.request, 'Prazo de monitoramento adicionado com sucesso.')
        return super().form_valid(form)

    def get_success_url(self):
        return reverse('contratos:contrato_detail', kwargs={'pk': self.contrato.pk})

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['contrato'] = self.contrato
        context['titulo'] = f'Adicionar Prazo - {self.contrato.apelido}'
        return context


class PrazoMonitoramentoUpdateView(ContratosWriteMixin, CheckCompetenciasGeradasMixin, UpdateView):
    """Edita um prazo de monitoramento."""
    model = PrazoMonitoramento
    form_class = PrazoMonitoramentoForm
    template_name = 'contratos/prazo_form.html'

    def get_success_url(self):
        return reverse('contratos:contrato_detail', kwargs={'pk': self.object.contrato.pk})

    def form_valid(self, form):
        messages.success(self.request, 'Prazo de monitoramento atualizado com sucesso.')
        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['contrato'] = self.object.contrato
        context['titulo'] = f'Editar Prazo - {self.object.nome}'
        return context


class PrazoMonitoramentoDeleteView(ContratosWriteMixin, CheckCompetenciasGeradasMixin, DeleteView):
    """Exclui um prazo de monitoramento."""
    model = PrazoMonitoramento
    template_name = 'contratos/confirm_delete.html'

    def get_success_url(self):
        return reverse('contratos:contrato_detail', kwargs={'pk': self.object.contrato.pk})

    def form_valid(self, form):
        messages.success(self.request, 'Prazo de monitoramento excluído com sucesso.')
        return super().form_valid(form)


class PrazoMonitoramentoConcluirView(ContratosWriteMixin, CheckCompetenciasGeradasMixin, View):
    """Marca o prazo como concluído (retirando da renderização ativa)."""
    
    def post(self, request, *args, **kwargs):
        prazo = get_object_or_404(PrazoMonitoramento, pk=kwargs['pk'])
        prazo.concluido = True
        prazo.save()
        messages.success(request, f'Prazo "{prazo.nome}" concluído com sucesso.')
        return redirect('contratos:contrato_detail', pk=prazo.contrato.pk)


import os
import subprocess
import tempfile
from decimal import Decimal
from django.http import FileResponse, Http404
import docx
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT

def set_cell_background(cell, hex_color):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="cbd5e1"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="94a3b8"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def _estilizar_celula_tabela(celula, *, destaque=False, centralizado=False, tamanho=9):
    """Aplica o padrão visual institucional já usado nas demais tabelas dos documentos."""

    set_cell_margins(celula)
    celula.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if destaque:
        set_cell_background(celula, "143642")
    for paragrafo in celula.paragraphs:
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER if centralizado else WD_ALIGN_PARAGRAPH.LEFT
        for run in paragrafo.runs:
            run.font.bold = destaque
            run.font.color.rgb = RGBColor(255, 255, 255) if destaque else RGBColor(20, 54, 66)
            run.font.size = Pt(tamanho)


def _preencher_celula_tabela(celula, texto, *, destaque=False, centralizado=False, tamanho=9):
    """Centraliza o preenchimento das células para facilitar montagem das tabelas-resumo do atestado."""

    celula.text = texto or '-'
    _estilizar_celula_tabela(celula, destaque=destaque, centralizado=centralizado, tamanho=tamanho)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'<w:top w:val="single" w:sz="4" w:space="0" w:color="cbd5e1"/>'
        f'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="94a3b8"/>'
        f'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="e2e8f0"/>'
        f'<w:insideV w:val="none"/>'
        f'<w:left w:val="none"/>'
        f'<w:right w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def gerar_documento_modelo(doc_path, contrato, competencia):
    doc = docx.Document(doc_path)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p_title.add_run("EXPEDIENTE DE FISCALIZAÇÃO")
    p_run.bold = True
    p_run.font.size = Pt(16)
    p_run.font.color.rgb = RGBColor(20, 54, 66)
    p_title.paragraph_format.space_after = Pt(24)
    
    p_assunto = doc.add_paragraph()
    p_assunto.add_run("Assunto: ").bold = True
    p_assunto.add_run(f"Fiscalização do Contrato {contrato.numero_contrato} - {contrato.objeto} - {competencia.periodo_inicio:%m/%Y}\n")
    p_assunto.paragraph_format.space_after = Pt(12)
    
    p_just = doc.add_paragraph()
    p_just.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_just.add_run(
        f"Trata o presente de expediente de fiscalização do contrato {contrato.numero_contrato} - {contrato.objeto} "
        f"no período de {competencia.periodo_inicio:%d/%m/%Y} e {competencia.periodo_fim:%d/%m/%Y}"
    )
    p_just.paragraph_format.space_after = Pt(48)
    
    # Assinaturas centralizadas
    p_sign1 = doc.add_paragraph()
    p_sign1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome_tec = nome_responsavel_ou_placeholder(contrato.fiscal_tecnico)
    p_sign1.add_run(f"_________________________________________\n").bold = True
    p_sign1.add_run(f"{nome_tec}\n").bold = True
    p_sign1.add_run("Fiscal Técnico\n")
    p_sign1.paragraph_format.space_after = Pt(24)
    
    p_sign2 = doc.add_paragraph()
    p_sign2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome_adm = nome_responsavel_ou_placeholder(contrato.fiscal_administrativo)
    p_sign2.add_run(f"_________________________________________\n").bold = True
    p_sign2.add_run(f"{nome_adm}\n").bold = True
    p_sign2.add_run("Fiscal Administrativo\n")
    p_sign2.paragraph_format.space_after = Pt(24)
    
    p_sign3 = doc.add_paragraph()
    p_sign3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome_ges = nome_responsavel_ou_placeholder(contrato.gestor_contrato)
    p_sign3.add_run(f"_________________________________________\n").bold = True
    p_sign3.add_run(f"{nome_ges}\n").bold = True
    p_sign3.add_run("Gestor do Contrato\n")
    p_sign3.paragraph_format.space_after = Pt(24)
    
    return doc


def gerar_capa_documento(doc_path, titulo, contrato, competencia):
    doc = docx.Document(doc_path)
    
    # Adiciona espaçamento inicial para centralização vertical simples
    for _ in range(8):
        doc.add_paragraph()
        
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p_title.add_run(titulo)
    p_run.bold = True
    p_run.font.size = Pt(24)
    p_run.font.color.rgb = RGBColor(20, 54, 66)
    p_title.paragraph_format.space_after = Pt(12)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub_run = p_sub.add_run(f"Contrato nº: {contrato.numero_contrato} ({contrato.objeto})\nCompetência: {competencia.periodo_inicio:%m/%Y}")
    p_sub_run.font.size = Pt(12)
    p_sub_run.font.color.rgb = RGBColor(100, 116, 139)
    
    return doc


def gerar_relatorio_avaliacao_competencia(doc_path, contrato, competencia):
    """Monta um relatório detalhado da avaliação de qualidade para entrar no pacote consolidado."""

    avaliacao = competencia.avaliacao_qualidade_segura
    # Recalcula a avaliação antes da exportação para evitar que downloads antigos carreguem totais
    # persistidos com regra anterior.
    if avaliacao is not None:
        recalcular_avaliacao_v2(avaliacao)
        avaliacao.refresh_from_db()
        competencia.refresh_from_db()

    doc = docx.Document(doc_path)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p_title.add_run("RELATÓRIO DE AVALIAÇÃO DE QUALIDADE")
    p_run.bold = True
    p_run.font.size = Pt(16)
    p_run.font.color.rgb = RGBColor(20, 54, 66)
    p_title.paragraph_format.space_after = Pt(24)

    p_info = doc.add_paragraph()
    p_info.add_run("Contrato nº: ").bold = True
    p_info.add_run(f"{contrato.numero_contrato} ({contrato.objeto})\n")
    p_info.add_run("Competência: ").bold = True
    p_info.add_run(f"{competencia.periodo_inicio:%m/%Y} ({competencia.periodo_inicio:%d/%m/%Y} a {competencia.periodo_fim:%d/%m/%Y})\n")
    p_info.paragraph_format.space_after = Pt(14)

    h_resumo = doc.add_paragraph()
    h_resumo_run = h_resumo.add_run("1. Resultado Consolidado")
    h_resumo_run.bold = True
    h_resumo_run.font.size = Pt(13)
    h_resumo_run.font.color.rgb = RGBColor(20, 54, 66)
    h_resumo.paragraph_format.space_after = Pt(6)

    p_resumo = doc.add_paragraph()
    p_resumo.add_run("Nota final: ").bold = True
    # O relatório deve refletir exatamente a nota consolidada persistida pela regra central da avaliação.
    p_resumo.add_run(f"{avaliacao.nota_final if avaliacao else '-'}\n")
    p_resumo.add_run("Percentual de liberação sugerido: ").bold = True
    p_resumo.add_run(f"{avaliacao.percentual_liberacao_sugerido if avaliacao else '-'}%\n")
    p_resumo.add_run("Valor liberado sugerido: ").bold = True
    p_resumo.add_run(f"R$ {competencia.valor_liberado_sugerido:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'))
    p_resumo.paragraph_format.space_after = Pt(12)

    h_itens = doc.add_paragraph()
    h_itens_run = h_itens.add_run("2. Itens Avaliados")
    h_itens_run.bold = True
    h_itens_run.font.size = Pt(13)
    h_itens_run.font.color.rgb = RGBColor(20, 54, 66)
    h_itens.paragraph_format.space_after = Pt(6)

    t_itens = doc.add_table(rows=1, cols=8)
    t_itens.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_itens)

    hdr_cells = t_itens.rows[0].cells
    hdr_titles = ["Grupo", "Item", "Peso", "Nota fiscal", "Justificativa fiscal", "Nota gestor", "Manifestação gestor", "Nota final"]
    for idx, title in enumerate(hdr_titles):
        hdr_cells[idx].text = title
        set_cell_background(hdr_cells[idx], "143642")
        set_cell_margins(hdr_cells[idx])
        for p in hdr_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    respostas = list(avaliacao.itens.order_by('grupo_ordem', 'item_ordem', 'id')) if avaliacao else []
    for idx, resposta in enumerate(respostas):
        row_cells = t_itens.add_row().cells
        row_cells[0].text = resposta.grupo_nome
        row_cells[1].text = resposta.item_descricao
        row_cells[2].text = f"{resposta.item_peso_percentual}%"
        row_cells[3].text = '' if resposta.nota_fiscal_valor is None else str(resposta.nota_fiscal_valor)
        row_cells[4].text = resposta.justificativa_fiscal or '-'
        row_cells[5].text = '' if resposta.nota_gestor_valor is None else str(resposta.nota_gestor_valor)
        row_cells[6].text = resposta.manifestacao_gestor_item or '-'
        row_cells[7].text = '' if resposta.nota_valor is None else str(resposta.nota_valor)

        for cell in row_cells:
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "f8fafc")
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8.5)

    if not respostas:
        row_cells = t_itens.add_row().cells
        row_cells[0].text = "Sem avaliação detalhada disponível."
        for cell in row_cells[1:]:
            cell.text = ""
        for cell in row_cells:
            set_cell_margins(cell)

    if avaliacao and avaliacao.observacoes:
        h_obs = doc.add_paragraph()
        h_obs_run = h_obs.add_run("3. Observações Gerais")
        h_obs_run.bold = True
        h_obs_run.font.size = Pt(13)
        h_obs_run.font.color.rgb = RGBColor(20, 54, 66)
        h_obs.paragraph_format.space_before = Pt(12)
        h_obs.paragraph_format.space_after = Pt(6)

        p_obs = doc.add_paragraph()
        p_obs.add_run(avaliacao.observacoes)

    return doc


def gerar_ultima_folha_atestado(doc_path, contrato, competencia, tipo_nota='principal'):
    """Reconstrói o atestado de realização para a nota principal ou adicional."""

    doc = docx.Document(doc_path)
    empresa = contrato.empresa_contratada
    responsavel_empresa = contrato.responsavel_empresa_principal
    eh_nota_adicional = tipo_nota == 'adicional'
    numero_nota = competencia.numero_nota_adicional if eh_nota_adicional else competencia.numero_nota_fiscal
    valor_nota = competencia.valor_nota_adicional if eh_nota_adicional else competencia.valor_nota_fiscal
    retencao_ir = competencia.retencao_ir_adicional if eh_nota_adicional else competencia.retencao_ir
    retencao_inss = competencia.retencao_inss_adicional if eh_nota_adicional else competencia.retencao_inss
    retencao_iss = competencia.retencao_iss_adicional if eh_nota_adicional else competencia.retencao_iss
    retencao_pis = competencia.retencao_pis_pasep_adicional if eh_nota_adicional else competencia.retencao_pis_pasep
    retencao_cofins = competencia.retencao_cofins_adicional if eh_nota_adicional else competencia.retencao_cofins
    valor_liquido = competencia.valor_liquido_nota_adicional if eh_nota_adicional else competencia.valor_liberado_final
    vencimento_pagamento = (
        competencia.data_aceite_definitivo + timedelta(days=competencia.prazo_pagamento_dias or 0)
        if competencia.data_aceite_definitivo
        else None
    )
    fim_vigencia_atual = inclusive_end_date(contrato.data_inicio_vigencia, contrato.prazo_inicial_meses)
    processo_contrato = ' | '.join(
        filtro for filtro in [contrato.processo_sei_gestao_numero, contrato.processo_sei_execucao_numero] if filtro
    ) or '-'
    contato_empresa = ' | '.join(
        filtro for filtro in [
            responsavel_empresa.telefone if responsavel_empresa and responsavel_empresa.telefone else '',
            responsavel_empresa.email if responsavel_empresa and responsavel_empresa.email else '',
        ] if filtro
    ) or '-'

    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("ATESTADO DE REALIZAÇÃO")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    run_titulo.font.color.rgb = RGBColor(20, 54, 66)
    titulo.paragraph_format.space_after = Pt(20)

    tamanho_texto_tabela = 6.5
    tamanho_titulo_tabela = 8

    tabela_acompanhamento = doc.add_table(rows=6, cols=4)
    tabela_acompanhamento.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tabela_acompanhamento)
    cabecalho_acompanhamento = tabela_acompanhamento.rows[0].cells
    cabecalho_acompanhamento[0].merge(cabecalho_acompanhamento[-1])
    _preencher_celula_tabela(
        cabecalho_acompanhamento[0],
        'ACOMPANHAMENTO DE PAGAMENTO',
        destaque=True,
        centralizado=True,
        tamanho=tamanho_titulo_tabela,
    )
    linha_empresa = tabela_acompanhamento.rows[1].cells
    _preencher_celula_tabela(linha_empresa[0], 'EMPRESA', destaque=True, tamanho=tamanho_texto_tabela)
    linha_empresa[1].merge(linha_empresa[-1])
    _preencher_celula_tabela(linha_empresa[1], empresa.razao_social, tamanho=tamanho_texto_tabela)
    linha_processo = tabela_acompanhamento.rows[2].cells
    _preencher_celula_tabela(linha_processo[0], 'PROCESSO GESTÃO', destaque=True, tamanho=tamanho_texto_tabela)
    linha_processo[1].merge(linha_processo[-1])
    _preencher_celula_tabela(linha_processo[1], contrato.processo_sei_gestao_numero, tamanho=tamanho_texto_tabela)
    linha_execucao = tabela_acompanhamento.rows[3].cells
    _preencher_celula_tabela(linha_execucao[0], 'PROCESSO EXECUÇÃO', destaque=True, tamanho=tamanho_texto_tabela)
    linha_execucao[1].merge(linha_execucao[-1])
    _preencher_celula_tabela(linha_execucao[1], contrato.processo_sei_execucao_numero, tamanho=tamanho_texto_tabela)
    linha_contrato = tabela_acompanhamento.rows[4].cells
    _preencher_celula_tabela(linha_contrato[0], 'CONTRATO', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_contrato[1], contrato.numero_contrato, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_contrato[2], 'COMPETÊNCIA', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_contrato[3], f'{competencia.periodo_inicio:%m/%Y}', tamanho=tamanho_texto_tabela)
    linha_objeto = tabela_acompanhamento.rows[5].cells
    _preencher_celula_tabela(linha_objeto[0], 'OBJETO', destaque=True, tamanho=tamanho_texto_tabela)
    linha_objeto[1].merge(linha_objeto[-1])
    _preencher_celula_tabela(linha_objeto[1], contrato.objeto, tamanho=tamanho_texto_tabela)
    doc.add_paragraph()

    tabela_verificacao = doc.add_table(rows=5, cols=4)
    tabela_verificacao.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tabela_verificacao)
    cabecalho_verificacao = tabela_verificacao.rows[0].cells
    cabecalho_verificacao[0].merge(cabecalho_verificacao[-1])
    _preencher_celula_tabela(
        cabecalho_verificacao[0],
        'CHECKLIST DE VERIFICAÇÃO',
        destaque=True,
        centralizado=True,
        tamanho=tamanho_titulo_tabela,
    )
    linha_nf = tabela_verificacao.rows[1].cells
    _preencher_celula_tabela(linha_nf[0], 'DATA ACEITE DEFINITIVO', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_nf[1], competencia.data_aceite_definitivo.strftime('%d/%m/%Y') if competencia.data_aceite_definitivo else '-', tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_nf[2], 'NÚMERO NF', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_nf[3], numero_nota or '-', tamanho=tamanho_texto_tabela)
    linha_responsavel = tabela_verificacao.rows[2].cells
    _preencher_celula_tabela(linha_responsavel[0], 'VENCIMENTO', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_responsavel[1], vencimento_pagamento.strftime('%d/%m/%Y') if vencimento_pagamento else '-', tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_responsavel[2], 'RESPONSÁVEL DA EMPRESA', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_responsavel[3], responsavel_empresa.nome if responsavel_empresa else '-', tamanho=tamanho_texto_tabela)
    linha_contato = tabela_verificacao.rows[3].cells
    _preencher_celula_tabela(linha_contato[0], 'CONTATO', destaque=True, tamanho=tamanho_texto_tabela)
    linha_contato[1].merge(linha_contato[3])
    _preencher_celula_tabela(linha_contato[1], contato_empresa, tamanho=tamanho_texto_tabela)
    linha_contrato = tabela_verificacao.rows[4].cells
    _preencher_celula_tabela(linha_contrato[0], 'PRÓXIMO REAJUSTE', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_contrato[1], contrato.get_mes_reajuste_display() or '-', tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_contrato[2], 'QTDE. ADITIVOS', destaque=True, tamanho=tamanho_texto_tabela)
    _preencher_celula_tabela(linha_contrato[3], '0', tamanho=tamanho_texto_tabela)
    linha_vigencia = tabela_verificacao.add_row().cells
    _preencher_celula_tabela(linha_vigencia[0], 'VIGÊNCIA DO CONTRATO', destaque=True, tamanho=tamanho_texto_tabela)
    linha_vigencia[1].merge(linha_vigencia[3])
    _preencher_celula_tabela(linha_vigencia[1], fim_vigencia_atual.strftime('%d/%m/%Y') if fim_vigencia_atual else '-', tamanho=tamanho_texto_tabela)
    doc.add_paragraph()

    tabela_financeira = doc.add_table(rows=1, cols=7)
    tabela_financeira.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tabela_financeira)
    cabecalhos = [
        'Valor da nota',
        'IR',
        'INSS',
        'ISS',
        'PIS/PASEP',
        'COFINS',
        'Total líquido',
    ]
    for idx, titulo_coluna in enumerate(cabecalhos):
        celula = tabela_financeira.rows[0].cells[idx]
        celula.text = titulo_coluna
        set_cell_background(celula, "143642")
        set_cell_margins(celula)
        for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    linha_valores = tabela_financeira.add_row().cells
    linha_valores[0].text = formatar_moeda_brl(valor_nota)
    linha_valores[1].text = formatar_moeda_brl(retencao_ir)
    linha_valores[2].text = formatar_moeda_brl(retencao_inss)
    linha_valores[3].text = formatar_moeda_brl(retencao_iss)
    linha_valores[4].text = formatar_moeda_brl(retencao_pis)
    linha_valores[5].text = formatar_moeda_brl(retencao_cofins)
    linha_valores[6].text = formatar_moeda_brl(valor_liquido)
    for celula in linha_valores:
        set_cell_margins(celula)
    doc.add_paragraph()

    h_checklist = doc.add_paragraph()
    h_checklist_run = h_checklist.add_run("Checklist documental")
    h_checklist_run.bold = True
    h_checklist_run.font.size = Pt(13)
    h_checklist_run.font.color.rgb = RGBColor(20, 54, 66)
    h_checklist.paragraph_format.space_after = Pt(8)

    tabela_checklist = doc.add_table(rows=1, cols=3)
    tabela_checklist.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(tabela_checklist)
    for idx, titulo_coluna in enumerate(['Item', 'Situação', 'Data']):
        celula = tabela_checklist.rows[0].cells[idx]
        celula.text = titulo_coluna
        set_cell_background(celula, "143642")
        set_cell_margins(celula)
        for paragrafo in celula.paragraphs:
            for run in paragrafo.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(9)

    itens_checklist = list(competencia.checklist_itens.order_by('ordem', 'id'))
    for item in itens_checklist:
        linha = tabela_checklist.add_row().cells
        linha[0].text = item.titulo
        linha[1].text = 'Entregue' if item.anexo_principal else 'Pendente'
        linha[2].text = timezone.localtime(item.validado_em).strftime('%d/%m/%Y') if item.validado_em else '-'
        for celula in linha:
            set_cell_margins(celula)
    if not itens_checklist:
        linha = tabela_checklist.add_row().cells
        linha[0].text = 'Sem itens de checklist cadastrados.'
        linha[1].text = '-'
        linha[2].text = '-'
        for celula in linha:
            set_cell_margins(celula)

    doc.add_paragraph()
    signatarios = [
        montar_assinatura_pagamento(
            competencia,
            competencia.gestor_pagamento,
            'Gestor do contrato',
            competencia.gestor_pagamento_em_exercicio,
        ),
        montar_assinatura_pagamento(
            competencia,
            competencia.coordenadora_pagamento,
            'Coordenadora',
            competencia.coordenadora_em_exercicio,
        ),
        montar_assinatura_pagamento(
            competencia,
            competencia.diretora_pagamento,
            'Diretora',
            competencia.diretora_em_exercicio,
        ),
        montar_assinatura_pagamento(
            competencia,
            competencia.subsecretario_pagamento,
            'Subsecretário',
            competencia.subsecretario_em_exercicio,
        ),
    ]
    textos_assinatura = [
        "Atesto que os serviços e/ou materiais discriminados na referida(as) nota(s) fiscal(ais) foram entregues e/ou executados a contento, nos termos previstos no instrumento contratual (ou documentos equivalente), dentro do prazo previsto.",
        "Encaminhe-se Diretoria de Gestão Administrativa para chancela e aprovação",
        "Conferido e em ordem para pagamento.",
        "Autorizo o processamento do pagamento conforme informações acima.",
    ]
    for assinatura, texto_assinatura in zip(signatarios, textos_assinatura):
        texto = doc.add_paragraph()
        texto.alignment = WD_ALIGN_PARAGRAPH.CENTER
        texto.paragraph_format.space_before = Pt(14)
        texto.add_run(texto_assinatura)
        paragrafo = doc.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragrafo.paragraph_format.space_before = Pt(18)
        paragrafo.add_run("_________________________________________\n").bold = True
        paragrafo.add_run(f"{assinatura['nome']}\n").bold = True
        paragrafo.add_run(assinatura['cargo'])

    return doc


def gerar_relatorio_medicao_competencia(doc_path, contrato, competencia):
    """Gera o relatório de medição que entra no pacote consolidado em papel timbrado."""

    empresa = contrato.empresa_contratada
    doc = docx.Document(doc_path)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_run = p_title.add_run("RELATÓRIO DE MEDIÇÃO")
    p_run.bold = True
    p_run.font.size = Pt(16)
    p_run.font.color.rgb = RGBColor(20, 54, 66)
    p_title.paragraph_format.space_after = Pt(24)

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Informações Gerais")
    h1_run.bold = True
    h1_run.font.size = Pt(13)
    h1_run.font.color.rgb = RGBColor(20, 54, 66)
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(6)

    p_info = doc.add_paragraph()
    p_info.add_run("Contrato nº: ").bold = True
    p_info.add_run(f"{contrato.numero_contrato} ({contrato.objeto})\n")
    p_info.add_run("Objeto: ").bold = True
    p_info.add_run(f"{contrato.objeto}\n")
    p_info.add_run("Contratada: ").bold = True
    p_info.add_run(f"{empresa.razao_social} (CNPJ: {empresa.cnpj})\n")
    p_info.add_run("Competência: ").bold = True
    p_info.add_run(f"{competencia.periodo_inicio:%m/%Y} ({competencia.periodo_inicio:%d/%m/%Y} a {competencia.periodo_fim:%d/%m/%Y})\n")
    p_info.add_run("Responsável pela autorização: ").bold = True
    p_info.add_run(f"{competencia.autorizado_por.get_full_name() or competencia.autorizado_por.username if competencia.autorizado_por else '-'}")
    p_info.paragraph_format.space_after = Pt(14)

    h3 = doc.add_paragraph()
    h3_run = h3.add_run("2. Medição dos Serviços Prestados")
    h3_run.bold = True
    h3_run.font.size = Pt(13)
    h3_run.font.color.rgb = RGBColor(20, 54, 66)
    h3.paragraph_format.space_before = Pt(16)
    h3.paragraph_format.space_after = Pt(6)

    t_medicoes = doc.add_table(rows=1, cols=6)
    t_medicoes.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_medicoes)

    hdr_cells = t_medicoes.rows[0].cells
    hdr_titles = ["Item", "Descrição", "Qtd. Mensal", "Qtd. Medida", "Val. Unitário", "Subtotal"]
    for idx, title in enumerate(hdr_titles):
        hdr_cells[idx].text = title
        set_cell_background(hdr_cells[idx], "143642")
        set_cell_margins(hdr_cells[idx])
        for p in hdr_cells[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx in (0, 1) else WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    total_geral = Decimal('0.00')
    for idx, med in enumerate(competencia.medicoes.all()):
        row_cells = t_medicoes.add_row().cells
        row_cells[0].text = str(med.item_contrato.ordem)
        row_cells[1].text = med.item_contrato.descricao
        row_cells[2].text = f"{med.item_contrato.quantidade:.2f}".replace('.', ',')
        row_cells[3].text = f"{med.quantidade:.2f}".replace('.', ',')
        row_cells[4].text = formatar_moeda_brl(med.valor_unitario_aplicado)
        row_cells[5].text = formatar_moeda_brl(med.valor_subtotal)
        total_geral += med.valor_subtotal

        for col_idx, cell in enumerate(row_cells):
            set_cell_margins(cell)
            if idx % 2 == 1:
                set_cell_background(cell, "f8fafc")
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx in (0, 1) else WD_ALIGN_PARAGRAPH.RIGHT
                for run in p.runs:
                    run.font.size = Pt(9.5)

    row_cells = t_medicoes.add_row().cells
    row_cells[1].text = "TOTAL GERAL MEDIDO"
    row_cells[5].text = formatar_moeda_brl(total_geral)
    for col_idx, cell in enumerate(row_cells):
        set_cell_margins(cell)
        set_cell_background(cell, "f1f5f9")
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 1 else WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(10)

    sec_num = 3
    if competencia.exige_avaliacao:
        h4 = doc.add_paragraph()
        h4_run = h4.add_run(f"{sec_num}. Avaliação de Qualidade")
        h4_run.bold = True
        h4_run.font.size = Pt(13)
        h4_run.font.color.rgb = RGBColor(20, 54, 66)
        h4.paragraph_format.space_before = Pt(16)
        h4.paragraph_format.space_after = Pt(6)

        avaliacao = competencia.avaliacao_qualidade_segura
        p_eval = doc.add_paragraph()
        p_eval.add_run("Nota final: ").bold = True
        p_eval.add_run(f"{avaliacao.nota_final if avaliacao else '-'}\n")
        p_eval.add_run("Percentual de liberação sugerido: ").bold = True
        p_eval.add_run(f"{avaliacao.percentual_liberacao_sugerido if avaliacao else '-'}%\n")
        p_eval.add_run("Valor liberado sugerido: ").bold = True
        p_eval.add_run(formatar_moeda_brl(competencia.valor_liberado_sugerido))
        p_eval.paragraph_format.space_after = Pt(10)
        sec_num += 1

    if competencia.justificativa_divergencia:
        h5 = doc.add_paragraph()
        h5_run = h5.add_run(f"{sec_num}. Justificativa de Divergência")
        h5_run.bold = True
        h5_run.font.size = Pt(13)
        h5_run.font.color.rgb = RGBColor(20, 54, 66)
        h5.paragraph_format.space_before = Pt(14)
        h5.paragraph_format.space_after = Pt(6)

        p_just = doc.add_paragraph()
        p_just.add_run(competencia.justificativa_divergencia)
        p_just.paragraph_format.space_after = Pt(10)
        sec_num += 1

    h6 = doc.add_paragraph()
    h6_run = h6.add_run(f"{sec_num}. Registro do Pagamento Autorizado")
    h6_run.bold = True
    h6_run.font.size = Pt(13)
    h6_run.font.color.rgb = RGBColor(20, 54, 66)
    h6.paragraph_format.space_before = Pt(16)
    h6.paragraph_format.space_after = Pt(6)

    p_pag = doc.add_paragraph()
    p_pag.add_run("Valor a ser pago: ").bold = True
    p_pag.add_run(f"{formatar_moeda_brl(competencia.valor_liberado_final)}\n")
    p_pag.add_run("Data da autorização de pagamento: ").bold = True
    p_pag.add_run(f"{competencia.data_pagamento.strftime('%d/%m/%Y') if competencia.data_pagamento else '-'}\n")
    p_pag.paragraph_format.space_after = Pt(28)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.paragraph_format.space_before = Pt(36)
    nome_tec = nome_responsavel_ou_placeholder(contrato.fiscal_tecnico)
    data_medicao = timezone.localtime(competencia.medicao_concluida_em).strftime('%d/%m/%Y') if competencia.medicao_concluida_em else '-'
    run_name = p_sign.add_run(f"{nome_tec.upper()}\n")
    run_name.bold = True
    p_sign.add_run("Fiscal Técnico do Contrato\n")
    p_sign.add_run(f"Data: {data_medicao}\n")
    return doc


ETAPAS_EXPORTACAO_DOCUMENTOS = [
    ('memorando', 'Preparando memorando', 10),
    ('checklist', 'Gerando checklist', 25),
    ('medicao', 'Gerando medição', 45),
    ('avaliacao', 'Gerando avaliação', 60),
    ('pagamento', 'Preparando pagamento', 80),
    ('montagem_final', 'Montando PDF final', 95),
]


def _usuario_pode_acompanhar_exportacao(user, job):
    """Restringe consulta e download ao solicitante da tarefa ou a administradores do sistema."""

    return bool(
        user
        and user.is_authenticated
        and (user.pk == job.solicitado_por_id or user.is_superuser or user.is_staff)
    )


def _serializar_exportacao(job):
    """Entrega o payload usado pelo polling do frontend."""

    data = {
        'job_id': job.pk,
        'status': job.status,
        'etapa_atual': job.etapa_atual,
        'percentual': job.percentual,
        'mensagem': job.mensagem,
        'erro_detalhe': job.erro_detalhe,
        'download_url': '',
    }
    if job.status == job.Status.CONCLUIDO and job.arquivo_pdf:
        data['download_url'] = reverse('contratos:competencia_download_docs_file', args=[job.pk])
    return data


def _atualizar_exportacao_documentos(job_id, status=None, etapa_atual=None, percentual=None, mensagem=None, erro_detalhe=None, concluido=False):
    """Atualiza o progresso da exportação sem manter instâncias abertas entre threads."""

    campos = {}
    if status is not None:
        campos['status'] = status
    if etapa_atual is not None:
        campos['etapa_atual'] = etapa_atual
    if percentual is not None:
        campos['percentual'] = percentual
    if mensagem is not None:
        campos['mensagem'] = mensagem
    if erro_detalhe is not None:
        campos['erro_detalhe'] = erro_detalhe
    if concluido:
        campos['concluido_em'] = timezone.now()
    if campos:
        ExportacaoDocumentosCompetencia.objects.filter(pk=job_id).update(**campos)


def _converter_docx_para_pdf(docx_file, output_dir):
    """Converte um DOCX individual para PDF via LibreOffice headless."""

    cmd = ["/usr/bin/libreoffice", "--headless", "--convert-to", "pdf", "--outdir", str(output_dir), str(docx_file)]
    subprocess.run(cmd, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
    return Path(output_dir) / f'{Path(docx_file).stem}.pdf'


def _normalizar_nome_checklist_exportado(item, competencia):
    """Padroniza os nomes dos anexos de checklist no modo separado."""

    titulo_normalizado = (
        (item.titulo or 'Documento do checklist')
        .replace('/', '-')
        .replace('\\', '-')
        .strip()
    )
    competencia_normalizada = f'{competencia.periodo_inicio:%m-%Y}'
    contrato_normalizado = f'CRT{competencia.contrato.numero_contrato}'.replace('/', '-')
    return f'{titulo_normalizado} - {contrato_normalizado} - {competencia_normalizada}.pdf'


def _gerar_pdf_observacoes_checklist(doc_path, competencia, output_dir):
    """Transforma as observações finais da competência em um PDF simples."""

    if not (competencia.observacoes_medicao or '').strip():
        return None
    doc = docx.Document(doc_path)
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = titulo.add_run("OBSERVAÇÕES")
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)
    texto = doc.add_paragraph()
    texto.add_run(competencia.observacoes_medicao)
    caminho_docx = Path(output_dir) / f'observacoes_{competencia.pk}.docx'
    doc.save(caminho_docx)
    return _converter_docx_para_pdf(caminho_docx, output_dir)


def _salvar_artefato_exportacao(job_id, competencia, arquivo_final_path, tipo_saida):
    """Persiste o artefato final do job, seja PDF unificado ou ZIP separado."""

    extensao = '.zip' if tipo_saida == 'separado' else '.pdf'
    with open(arquivo_final_path, 'rb') as arquivo_final:
        job = ExportacaoDocumentosCompetencia.objects.get(pk=job_id)
        if job.arquivo_pdf:
            job.arquivo_pdf.delete(save=False)
        job.arquivo_pdf.save(
            f'relatorio_competencia_{competencia.periodo_inicio:%m_%Y}_{job.pk}{extensao}',
            ContentFile(arquivo_final.read()),
            save=False,
        )
        job.status = ExportacaoDocumentosCompetencia.Status.CONCLUIDO
        job.etapa_atual = 'Arquivo concluído'
        job.percentual = 100
        job.mensagem = 'Arquivo consolidado gerado com sucesso.'
        job.concluido_em = timezone.now()
        job.erro_detalhe = ''
        job.save(update_fields=['arquivo_pdf', 'status', 'etapa_atual', 'percentual', 'mensagem', 'concluido_em', 'erro_detalhe', 'atualizado_em'])


def _normalizar_pdf_para_mesclagem(pdf_path, output_dir, nome_exibicao, indice):
    """Regrava cada PDF antes da mesclagem para reduzir incompatibilidades de origem."""

    caminho_origem = Path(pdf_path)
    if not caminho_origem.exists():
        raise FileNotFoundError(f'Arquivo PDF não encontrado durante a montagem: {nome_exibicao}.')

    try:
        with open(caminho_origem, 'rb') as arquivo_origem:
            leitor = PdfReader(BytesIO(arquivo_origem.read()), strict=False)
        if not leitor.pages:
            raise ValueError('O arquivo está vazio.')

        escritor = PdfWriter()
        for pagina in leitor.pages:
            escritor.add_page(pagina)

        caminho_normalizado = Path(output_dir) / f'normalizado_{indice:02d}_{caminho_origem.stem}.pdf'
        with open(caminho_normalizado, 'wb') as arquivo_normalizado:
            escritor.write(arquivo_normalizado)
        return caminho_normalizado
    except Exception as exc:
        raise ValidationError(
            f'Não foi possível ler o PDF "{nome_exibicao}". Verifique se o arquivo está íntegro e tente reenviá-lo.'
        ) from exc


def _gerar_arquivo_pdf_consolidado(job_id):
    """Monta o pacote final da competência em PDF único ou ZIP separado."""

    job = ExportacaoDocumentosCompetencia.objects.select_related(
        'competencia__contrato__empresa_contratada',
        'solicitado_por',
    ).get(pk=job_id)
    competencia = CompetenciaPagamento.objects.select_related(
        'contrato__empresa_contratada',
        'autorizado_por',
        'gestor_pagamento',
        'coordenadora_pagamento',
        'diretora_pagamento',
        'subsecretario_pagamento',
        'contrato__fiscal_tecnico',
        'contrato__fiscal_administrativo',
        'contrato__gestor_contrato',
    ).prefetch_related(
        'checklist_itens__anexo',
        'medicoes__item_contrato',
        'avaliacao_qualidade__itens',
    ).get(pk=job.competencia_id)
    contrato = competencia.contrato
    doc_path = '/root/aplicacoesspi/docs/papel-timbrado-spi.docx'
    if not os.path.exists(doc_path):
        raise FileNotFoundError('Arquivo de papel timbrado não encontrado no servidor.')
    if not competencia.nota_fiscal_fatura:
        raise FileNotFoundError('A Nota Fiscal/Fatura é obrigatória para gerar os documentos da competência.')

    with tempfile.TemporaryDirectory(prefix=f'competencia_{competencia.pk}_') as temp_dir:
        temp_dir_path = Path(temp_dir)
        tipo_saida = job.tipo_saida or 'unificado'

        _atualizar_exportacao_documentos(job_id, status=ExportacaoDocumentosCompetencia.Status.PROCESSANDO, etapa_atual='Preparando memorando', percentual=10, mensagem='Gerando o memorando de abertura em PDF.')
        memorando_docx = temp_dir_path / f'memorando_{competencia.pk}.docx'
        gerar_documento_modelo(doc_path, contrato, competencia).save(memorando_docx)
        memorando_pdf = _converter_docx_para_pdf(memorando_docx, temp_dir_path)

        _atualizar_exportacao_documentos(job_id, etapa_atual='Gerando medição', percentual=25, mensagem='Preparando a capa da medição e o relatório em papel timbrado.')
        medicao_capa_docx = temp_dir_path / f'medicao_capa_{competencia.pk}.docx'
        gerar_capa_documento(doc_path, "MEDIÇÃO DO OBJETO", contrato, competencia).save(medicao_capa_docx)
        medicao_capa_pdf = _converter_docx_para_pdf(medicao_capa_docx, temp_dir_path)
        medicao_relatorio_docx = temp_dir_path / f'medicao_relatorio_{competencia.pk}.docx'
        gerar_relatorio_medicao_competencia(doc_path, contrato, competencia).save(medicao_relatorio_docx)
        medicao_relatorio_pdf = _converter_docx_para_pdf(medicao_relatorio_docx, temp_dir_path)

        _atualizar_exportacao_documentos(job_id, etapa_atual='Gerando checklist', percentual=50, mensagem='Preparando a capa do checklist e reunindo os documentos anexados.')
        checklist_capa_docx = temp_dir_path / f'checklist_capa_{competencia.pk}.docx'
        gerar_capa_documento(doc_path, "CHECKLIST DE PAGAMENTO", contrato, competencia).save(checklist_capa_docx)
        checklist_capa_pdf = _converter_docx_para_pdf(checklist_capa_docx, temp_dir_path)

        _atualizar_exportacao_documentos(job_id, etapa_atual='Preparando documentos financeiros', percentual=75, mensagem='Preparando notas fiscais, atestados e anexos complementares.')
        nota_fiscal_capa_docx = temp_dir_path / f'nota_fiscal_capa_{competencia.pk}.docx'
        gerar_capa_documento(doc_path, "NOTA FISCAL", contrato, competencia).save(nota_fiscal_capa_docx)
        nota_fiscal_capa_pdf = _converter_docx_para_pdf(nota_fiscal_capa_docx, temp_dir_path)
        atestado_principal_docx = temp_dir_path / f'atestado_principal_{competencia.pk}.docx'
        gerar_ultima_folha_atestado(doc_path, contrato, competencia, tipo_nota='principal').save(atestado_principal_docx)
        atestado_principal_pdf = _converter_docx_para_pdf(atestado_principal_docx, temp_dir_path)
        atestado_adicional_pdf = None
        if competencia.possui_nota_adicional:
            atestado_adicional_docx = temp_dir_path / f'atestado_adicional_{competencia.pk}.docx'
            gerar_ultima_folha_atestado(doc_path, contrato, competencia, tipo_nota='adicional').save(atestado_adicional_docx)
            atestado_adicional_pdf = _converter_docx_para_pdf(atestado_adicional_docx, temp_dir_path)
        observacoes_pdf = _gerar_pdf_observacoes_checklist(doc_path, competencia, temp_dir_path)

        _atualizar_exportacao_documentos(job_id, etapa_atual='Montando arquivo final', percentual=95, mensagem='Mesclando todas as peças na sequência configurada.')
        artefatos = [
            ('Memorando', memorando_pdf),
            ('Capa Medição', medicao_capa_pdf),
            ('Medição', medicao_relatorio_pdf),
            ('Aceite Provisório', Path(competencia.aceite_provisorio_arquivo.path) if competencia.aceite_provisorio_arquivo else None),
            ('Aceite Definitivo', Path(competencia.aceite_definitivo_arquivo.path) if competencia.aceite_definitivo_arquivo else None),
            ('Avaliação de Qualidade assinada pelas partes', Path(competencia.avaliacao_assinada.path) if competencia.avaliacao_assinada else None),
            ('Nota Fiscal', nota_fiscal_capa_pdf),
            ('NF do mês', Path(competencia.nota_fiscal_fatura.path)),
            ('Nota Fiscal adicional', Path(competencia.nota_adicional_arquivo.path) if competencia.nota_adicional_arquivo else None),
            ('Capa Checklist', checklist_capa_pdf),
        ]
        for item in competencia.checklist_itens.all().order_by('ordem', 'id'):
            if hasattr(item, 'anexo') and item.anexo and item.anexo.arquivo:
                caminho = Path(item.anexo.arquivo.path)
                if caminho.exists():
                    artefatos.append((_normalizar_nome_checklist_exportado(item, competencia), caminho))
        artefatos.append(('Atestado de Realização da NF com as retenções', atestado_principal_pdf))
        if observacoes_pdf:
            artefatos.append(('Observações', observacoes_pdf))
        if atestado_adicional_pdf:
            artefatos.append(('Atestado de realização NF Adicional com as retenções', atestado_adicional_pdf))
        artefatos = [(nome, caminho) for nome, caminho in artefatos if caminho]

        if tipo_saida == 'separado':
            caminho_zip = temp_dir_path / f'competencia_{competencia.pk}_documentos.zip'
            with zipfile.ZipFile(caminho_zip, 'w', compression=zipfile.ZIP_DEFLATED) as arquivo_zip:
                for indice, (nome_exibicao, caminho) in enumerate(artefatos, start=1):
                    extensao = Path(caminho).suffix.lower() or '.pdf'
                    nome_arquivo = f'{indice:02d} - {nome_exibicao}'
                    if 'CRT' not in nome_exibicao:
                        nome_arquivo = f'{nome_arquivo}{extensao}'
                    arquivo_zip.write(caminho, arcname=nome_arquivo)
            _salvar_artefato_exportacao(job_id, competencia, caminho_zip, tipo_saida)
        else:
            caminho_pdf_final = temp_dir_path / f'competencia_{competencia.pk}_consolidado.pdf'
            merger = PdfWriter()
            for indice, (nome_exibicao, caminho) in enumerate(artefatos, start=1):
                merger.append(str(_normalizar_pdf_para_mesclagem(caminho, temp_dir_path, nome_exibicao, indice)))
            with open(caminho_pdf_final, 'wb') as arquivo_final:
                merger.write(arquivo_final)
            merger.close()
            _salvar_artefato_exportacao(job_id, competencia, caminho_pdf_final, tipo_saida)

        competencia.download_realizado_em = timezone.now()
        if competencia.status != competencia.Status.PAGA:
            competencia.status = competencia.Status.OB_PENDENTE
        competencia.save(update_fields=['download_realizado_em', 'status', 'atualizado_em'])


def processar_exportacao_documentos_competencia(job_id):
    """Executa a tarefa de exportação e persiste o erro em caso de falha."""

    try:
        if 'test' not in sys.argv:
            close_old_connections()
        _gerar_arquivo_pdf_consolidado(job_id)
    except Exception as exc:
        _atualizar_exportacao_documentos(
            job_id,
            status=ExportacaoDocumentosCompetencia.Status.ERRO,
            etapa_atual='Falha na geração',
            percentual=100,
            mensagem='Não foi possível gerar o PDF consolidado.',
            erro_detalhe=str(exc),
            concluido=True,
        )
    finally:
        if 'test' not in sys.argv:
            close_old_connections()


class CompetenciaDownloadDocsStartView(ContratosWriteMixin, ContractManagePermissionMixin, View):
    """Inicia ou reaproveita a geração assíncrona do PDF consolidado da competência."""

    def post(self, request, *args, **kwargs):
        competencia = get_object_or_404(CompetenciaPagamento.objects.select_related('contrato'), pk=kwargs['pk'])
        response = self.ensure_manage_permission(request, competencia.contrato)
        if response:
            return JsonResponse({'detail': 'Acesso negado.'}, status=403)
        if competencia.status not in {competencia.Status.DOWNLOAD_PENDENTE, competencia.Status.OB_PENDENTE, competencia.Status.PAGA}:
            return JsonResponse({'detail': 'Os documentos só podem ser gerados após concluir medição, avaliação e checklist.'}, status=400)
        tipo_saida = request.GET.get('tipo_saida', 'unificado')

        job = ExportacaoDocumentosCompetencia.objects.filter(
            competencia=competencia,
            solicitado_por=request.user,
            tipo_saida=tipo_saida,
            status__in=[ExportacaoDocumentosCompetencia.Status.PENDENTE, ExportacaoDocumentosCompetencia.Status.PROCESSANDO],
        ).order_by('-criado_em', '-id').first()

        if job is None:
            job = ExportacaoDocumentosCompetencia.objects.create(
                competencia=competencia,
                solicitado_por=request.user,
                tipo_saida=tipo_saida,
                status=ExportacaoDocumentosCompetencia.Status.PENDENTE,
                etapa_atual='Na fila',
                percentual=0,
                mensagem='A exportação foi criada e será iniciada em instantes.',
            )
            if 'test' in sys.argv:
                processar_exportacao_documentos_competencia(job.pk)
            else:
                threading.Thread(
                    target=processar_exportacao_documentos_competencia,
                    args=(job.pk,),
                    daemon=True,
                ).start()
            job.refresh_from_db()

        return JsonResponse(_serializar_exportacao(job))


class CompetenciaDownloadDocsStatusView(ContratosAccessMixin, View):
    """Expõe o progresso da geração do PDF consolidado para o polling do frontend."""

    def get(self, request, *args, **kwargs):
        job = get_object_or_404(ExportacaoDocumentosCompetencia.objects.select_related('competencia__contrato', 'solicitado_por'), pk=kwargs['job_pk'])
        if not _usuario_pode_acompanhar_exportacao(request.user, job):
            return JsonResponse({'detail': 'Acesso negado.'}, status=403)
        return JsonResponse(_serializar_exportacao(job))


class CompetenciaDownloadDocsFileView(ContratosAccessMixin, View):
    """Entrega o arquivo final pronto a partir do job concluído."""

    def get(self, request, *args, **kwargs):
        job = get_object_or_404(ExportacaoDocumentosCompetencia.objects.select_related('competencia', 'solicitado_por'), pk=kwargs['job_pk'])
        if not _usuario_pode_acompanhar_exportacao(request.user, job):
            raise Http404('Arquivo não encontrado.')
        if job.status != ExportacaoDocumentosCompetencia.Status.CONCLUIDO or not job.arquivo_pdf:
            raise Http404('O arquivo ainda não está disponível.')
        extensao = Path(job.arquivo_pdf.name).suffix.lower()
        filename = f'relatorio_competencia_{job.competencia.periodo_inicio:%m_%Y}{extensao or ".pdf"}'
        content_type = mimetypes.guess_type(job.arquivo_pdf.path)[0] or 'application/octet-stream'
        response = FileResponse(open(job.arquivo_pdf.path, 'rb'), content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response


class CompetenciaDownloadDocsView(ContratosAccessMixin, View):
    """Mantém um atalho legado para baixar o último arquivo concluído da competência para o usuário atual."""

    def get(self, request, *args, **kwargs):
        competencia = get_object_or_404(CompetenciaPagamento, pk=kwargs['pk'])
        job = ExportacaoDocumentosCompetencia.objects.filter(
            competencia=competencia,
            solicitado_por=request.user,
            status=ExportacaoDocumentosCompetencia.Status.CONCLUIDO,
        ).order_by('-concluido_em', '-id').first()
        if not job or not job.arquivo_pdf:
            raise Http404('Nenhum PDF consolidado pronto foi encontrado para esta competência.')
        extensao = Path(job.arquivo_pdf.name).suffix.lower()
        filename = f'relatorio_competencia_{competencia.periodo_inicio:%m_%Y}{extensao or ".pdf"}'
        content_type = mimetypes.guess_type(job.arquivo_pdf.path)[0] or 'application/octet-stream'
        response = FileResponse(open(job.arquivo_pdf.path, 'rb'), content_type=content_type)
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response
